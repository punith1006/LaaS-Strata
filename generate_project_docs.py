"""
Generate 14 AI Internship Project .docx Documents
Output: ai-internship-projects/01-PixelRevive.docx ... 14-PixelPlayground.docx
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

OUT_DIR = os.path.join(os.path.dirname(__file__), "ai-internship-projects")
os.makedirs(OUT_DIR, exist_ok=True)

# ── Helpers ──────────────────────────────────────────────────────────────────
def shade(cell, hex_color):
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'))

def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h

def add_bullet(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_note_box(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(f"NOTE: {text}")
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

def add_module_table(doc, modules, blaze_idx):
    """modules: list of dicts with name, description, deliverables, tools, vram, tier"""
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Module", "Description", "Key Deliverables", "Example Tools/Models", "VRAM Budget"]
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        shade(c, "1A1A2E")
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(9)
    for idx, m in enumerate(modules):
        row = table.add_row()
        tier = "[BLAZE] Blaze" if idx == blaze_idx else "Spark"
        tier_color = "FFF3CD" if idx == blaze_idx else "F8F9FA"
        vals = [f"{m['name']}\n({tier})", m['description'], "\n".join(f"• {d}" for d in m['deliverables']),
                ", ".join(m['tools']), m['vram']]
        for i, v in enumerate(vals):
            c = row.cells[i]
            shade(c, tier_color)
            r = c.paragraphs[0].add_run(v)
            r.font.size = Pt(9)

# ── Document Builder ─────────────────────────────────────────────────────────
def build_doc(p):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)

    # Title Page
    for _ in range(6):
        doc.add_paragraph("")
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tp.add_run(p['name'])
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sp.add_run(p['tagline'])
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x4A, 0x4A, 0x6A)

    doc.add_paragraph("")
    ip = doc.add_paragraph()
    ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = ip.add_run(f"Difficulty: {p['difficulty']}  •  Team Size: {p['team_size']} members  •  Platform: LaaS GPU Desktop")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph("")
    bp = doc.add_paragraph()
    bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = bp.add_run("15-Day Vibe-Coding Internship Program")
    r.font.size = Pt(12)
    r.font.italic = True

    doc.add_page_break()

    # Problem Statement
    add_heading_styled(doc, "Problem Statement", 1)
    doc.add_paragraph(p['problem'])
    doc.add_paragraph("")
    add_note_box(doc, p['why_it_matters'])

    # Objectives
    add_heading_styled(doc, "Objectives", 1)
    for obj in p['objectives']:
        add_bullet(doc, obj)

    # GPU Tier Allocation
    add_heading_styled(doc, "GPU Tier Allocation", 1)
    doc.add_paragraph(
        "Each team receives one Blaze instance and the remaining members receive Spark instances. "
        "The Blaze instance is assigned to the module with the highest GPU compute requirement. "
        "Teams should decide among themselves which member takes the Blaze instance based on "
        "which module demands the most VRAM or training compute."
    )
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for i, h in enumerate(["Tier", "Specs", "Assigned To"]):
        c = table.rows[0].cells[i]
        shade(c, "1A1A2E")
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)
    blaze_mod = p['modules'][p['blaze_idx']]
    for tier, specs, assigned in [
        ("[BLAZE] Blaze (x1)", "4 vCPU, 8 GB RAM, 4 GB VRAM", blaze_mod['name']),
        ("Spark (x{})".format(p['team_size']-1), "2 vCPU, 4 GB RAM, 2 GB VRAM", "All other modules"),
    ]:
        row = table.add_row()
        for i, v in enumerate([tier, specs, assigned]):
            r = row.cells[i].paragraphs[0].add_run(v)
            r.font.size = Pt(10)

    # Module Breakdown
    add_heading_styled(doc, "Module Breakdown", 1)
    doc.add_paragraph(
        "Each module below represents a distinct piece of the project. Team members should "
        "self-assign modules based on interest and skill. No member is pre-assigned."
    )
    add_note_box(doc, "All suggested tools, models, and libraries are EXAMPLES ONLY. "
                 "Teams are free to choose any alternative that fits within the provisioned "
                 "resource constraints (Spark: 2 GB VRAM, Blaze: 4 GB VRAM).")
    doc.add_paragraph("")
    add_module_table(doc, p['modules'], p['blaze_idx'])

    # Detailed module descriptions
    for idx, m in enumerate(p['modules']):
        tier = "Blaze" if idx == p['blaze_idx'] else "Spark"
        add_heading_styled(doc, f"Module: {m['name']} ({tier})", 2)
        doc.add_paragraph(m['detail'])

    # Data Input
    add_heading_styled(doc, "Data Input Method", 1)
    doc.add_paragraph(p['data_input'])

    # Outcomes & Expectations
    add_heading_styled(doc, "Expected Outcomes", 1)
    for o in p['outcomes']:
        add_bullet(doc, o)

    # Demo Description
    add_heading_styled(doc, "Demo Day Presentation", 1)
    doc.add_paragraph(p['demo'])

    # Technical Guidelines
    add_heading_styled(doc, "Technical Guidelines", 1)
    for g in p['guidelines']:
        add_bullet(doc, g)

    # Timeline
    add_heading_styled(doc, "Timeline & Milestones", 1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for i, h in enumerate(["Phase", "Days", "Milestone"]):
        c = table.rows[0].cells[i]
        shade(c, "1A1A2E")
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)
    for phase, days, milestone in p['milestones']:
        row = table.add_row()
        for i, v in enumerate([phase, days, milestone]):
            r = row.cells[i].paragraphs[0].add_run(v)
            r.font.size = Pt(10)

    # Evaluation
    add_heading_styled(doc, "Evaluation Criteria", 1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for i, h in enumerate(["Criterion", "Weight", "What Judges Look For"]):
        c = table.rows[0].cells[i]
        shade(c, "1A1A2E")
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)
    for crit, weight, desc in p['evaluation']:
        row = table.add_row()
        for i, v in enumerate([crit, weight, desc]):
            r = row.cells[i].paragraphs[0].add_run(v)
            r.font.size = Pt(10)

    # Resource Reminder
    add_heading_styled(doc, "Resource Constraints Reminder", 1)
    add_bullet(doc, "Spark: 2 vCPU, 4 GB RAM, ", "GPU (Spark): ")
    add_bullet(doc, "Blaze: 4 vCPU, 8 GB RAM, ", "GPU (Blaze): ")
    add_bullet(doc, "15 GB persistent storage in /home/ubuntu", "Storage: ")
    add_bullet(doc, "Upload files via LaaS web UI or paste URLs (yt-dlp for YouTube)", "Data Input: ")
    add_bullet(doc, "No live webcam/microphone. All data must be pre-recorded.", "Important: ")
    add_bullet(doc, "Containers have internet access for pip install, apt install, model downloads.", "Network: ")
    add_bullet(doc, "/home/ubuntu persists across restarts. System packages (apt) do NOT persist.", "Persistence: ")

    # Save
    filename = f"{p['num']:02d}-{p['slug']}.docx"
    path = os.path.join(OUT_DIR, filename)
    doc.save(path)
    print(f"  [OK] {filename}")
    return path

# ── PROJECT DATA ─────────────────────────────────────────────────────────────
# fmt: off
PROJECTS = [
{
    "num": 1, "slug": "PixelRevive", "name": "PixelRevive — AI Photo Restoration Studio",
    "tagline": "Upload a damaged, old photo → get back a restored, colorized, HD version",
    "difficulty": "Standard", "team_size": 4, "blaze_idx": 0,
    "problem": "Millions of families have old, damaged photographs — torn edges, water stains, faded colors, scratches — that hold irreplaceable memories. Professional photo restoration is expensive and time-consuming. AI-powered restoration can make this accessible to everyone, turning a single damaged photo into a restored, colorized, high-resolution image in seconds.",
    "why_it_matters": "Every family has old damaged photos. The before/after transformation is universally emotional and demonstrates real AI value. This project runs 3 different GPU models in a pipeline — a rare skill in industry.",
    "objectives": [
        "Build an end-to-end photo restoration pipeline that processes damaged images through multiple AI stages",
        "Implement scratch/tear removal using AI inpainting on GPU",
        "Add automatic colorization for black-and-white photos",
        "Apply AI super-resolution to upscale restored photos to HD/4K quality",
        "Create a polished web UI with side-by-side before/after comparison and batch processing"
    ],
    "data_input": "Students upload damaged/old photos via the LaaS web UI. Files are stored in /home/ubuntu and processed by the GPU pipeline. Supported formats: JPEG, PNG, TIFF.",
    "modules": [
        {"name": "Scratch & Damage Removal", "description": "AI inpainting model that detects and fills scratches, tears, water stains, and other damage artifacts in photos.",
         "deliverables": ["Damage detection mask generator", "Inpainting model (fills masked regions)", "Before/after comparison output"],
         "tools": ["U-Net / LaMa inpainting", "OpenCV", "PyTorch"], "vram": "~800 MB",
         "detail": "This module handles the first stage of the pipeline: identifying damaged regions and intelligently filling them. The model should detect scratches, tears, water stains, and creases, then use AI inpainting to reconstruct the missing content. Consider using pre-trained inpainting models or training a small U-Net on damage datasets. The output should be a clean, damage-free version of the input photo."},
        {"name": "Photo Colorization", "description": "Automatic colorization of black-and-white or sepia-toned photos using AI.",
         "deliverables": ["Colorization model", "Color palette estimation", "Confidence scoring for color accuracy"],
         "tools": ["DeOldify", "fastai", "ControlNet colorization"], "vram": "~100 MB",
         "detail": "This module takes the cleaned photo and adds realistic color. For B&W photos, the model should estimate plausible colors based on content (sky = blue, grass = green, skin tones). Consider using DeOldify, or building a custom colorization network. The output should look natural, not oversaturated."},
        {"name": "Super-Resolution Upscaling", "description": "AI-powered upscaling from low resolution to 4K quality while preserving detail.",
         "deliverables": ["4x upscaling model", "Detail preservation check", "Output format options (JPEG/PNG)"],
         "tools": ["Real-ESRGAN", "ESRGAN via ONNX", "OpenCV dnn_superres"], "vram": "~500 MB",
         "detail": "The final processing stage: take the restored, colorized photo and upscale it to high resolution. The model should add detail without introducing artifacts. Real-ESRGAN is a strong starting point, but teams can explore any super-resolution approach. Consider offering multiple upscale factors (2x, 4x)."},
        {"name": "Web Studio & Integration", "description": "Polished web application with upload, pipeline orchestration, side-by-side comparison, and batch processing.",
         "deliverables": ["Drag-drop upload UI", "Processing pipeline orchestration", "Before/after slider comparison", "Batch processing queue", "Download restored photos"],
         "tools": ["Flask/FastAPI", "React/Next.js", "ChromaDB for metadata"], "vram": "~300 MB",
         "detail": "This module ties everything together. Build a web application where users can upload photos, run the full restoration pipeline, and compare results side-by-side with a slider. Include batch processing for multiple photos. Consider adding an AI agent that describes what was fixed ('Removed 3 scratches, colorized, upscaled to 4K')."}
    ],
    "outcomes": [
        "A working web application that restores damaged photos end-to-end",
        "Three GPU-accelerated AI models running in a pipeline",
        "Before/after comparison demonstrating clear visual improvement",
        "Batch processing capability for multiple photos",
        "Clean, professional UI that could be shown to a real customer"
    ],
    "demo": "Show a torn, faded 1970s family photo. Click 'Restore.' Watch the pipeline run: scratches disappear, color returns, image sharpens to HD. Reveal the side-by-side comparison. The audience gasps.",
    "guidelines": [
        "Start with pre-trained models — don't train from scratch unless you have a good dataset",
        "Process images at reasonable resolution during development (512×512) to iterate faster",
        "Use tmux or multiple terminal tabs to run backend + frontend simultaneously",
        "Store all models and code in /home/ubuntu — it persists across container restarts",
        "Test with at least 10 different damaged photos to ensure robustness",
        "The pipeline should handle JPEG, PNG, and ideally TIFF formats"
    ],
    "milestones": [
        ("Foundation", "Days 1–5", "All 4 modules scaffold, pre-trained models downloaded, basic pipeline working on sample images"),
        ("Integration", "Days 6–10", "Full pipeline connected end-to-end, web UI functional, batch processing working"),
        ("Polish", "Days 11–14", "UI polished, edge cases handled, 10+ test photos processed successfully"),
        ("Demo Day", "Day 15", "Live demo with audience-submitted photo, side-by-side reveal")
    ],
    "evaluation": [
        ("Working Demo", "35%", "Full pipeline runs without errors on a new photo"),
        ("AI Quality", "25%", "Restoration quality — is the output visibly better than the input?"),
        ("UI/UX", "20%", "Professional-looking interface, smooth workflow"),
        ("Technical Depth", "10%", "Can explain pipeline architecture and model choices"),
        ("Innovation", "10%", "Extra features: batch processing, AI descriptions, style options")
    ]
},
{
    "num": 2, "slug": "SnapChef", "name": "SnapChef — Fridge-to-Recipe AI",
    "tagline": "Snap a photo of your fridge → AI identifies ingredients → generates a personalized recipe",
    "difficulty": "Standard", "team_size": 4, "blaze_idx": 1,
    "problem": "People stare at their fridge every day wondering 'what can I cook?' They have ingredients but lack recipe ideas. Food waste is a massive problem — households throw away edible food because they don't know how to combine what they have. An AI that looks at your fridge and suggests recipes solves both problems.",
    "why_it_matters": "Combines computer vision with generative AI in a practical, everyday product. Food waste reduction is a real industry need. The demo — watching bounding boxes appear on fridge items and a recipe generate instantly — is immediately compelling.",
    "objectives": [
        "Build an object detection model that identifies 50+ common food items in fridge/pantry photos",
        "Generate personalized recipes from detected ingredients using a local LLM",
        "Add dietary preference filtering (vegetarian, vegan, gluten-free, allergies)",
        "Create a cook-along UI with step-by-step instructions and timers"
    ],
    "data_input": "Students upload photos of their fridge, pantry, or ingredients via the LaaS web UI. The container's GPU processes the images for ingredient detection.",
    "modules": [
        {"name": "Ingredient Detection", "description": "Object detection model trained to identify common food items in fridge/pantry photos with bounding boxes.",
         "deliverables": ["YOLOv8 fine-tuned on food items", "50+ food classes", "Confidence scores per detection"],
         "tools": ["Ultralytics YOLOv8", "OpenCV", "Roboflow (for labeling)"], "vram": "~1.5 GB",
         "detail": "Train or fine-tune a YOLOv8 model to detect common food items: vegetables, fruits, dairy, meats, condiments, beverages. Start with a pre-trained COCO model and fine-tune on food-specific datasets. The model should output bounding boxes with class labels and confidence scores. Consider handling partially visible items and cluttered fridge scenarios."},
        {"name": "Recipe Generation", "description": "LLM-powered recipe generation from detected ingredients, considering dietary preferences and cuisine style.",
         "deliverables": ["Recipe generation pipeline", "Ingredient-to-recipe prompting", "Recipe formatting (ingredients list, steps, timing)"],
         "tools": ["Qwen2-1.5B (quantized)", "Transformers", "Prompt engineering"], "vram": "~1.1 GB",
         "detail": "Use a quantized LLM to generate complete recipes from the detected ingredients. The prompt should include the ingredient list, dietary preferences, and desired cuisine. Output should be structured: recipe name, ingredient list with quantities, step-by-step instructions, estimated cooking time, and serving size. This module runs on the Blaze instance for the extra VRAM headroom."},
        {"name": "Nutrition & Preferences", "description": "Dietary filtering, nutrition scoring, allergy checking, and preference matching using embeddings.",
         "deliverables": ["Nutrition database integration", "Allergy flag system", "Dietary preference filters", "Recipe similarity search"],
         "tools": ["Sentence-Transformers", "ChromaDB", "Nutrition API/database"], "vram": "~300 MB",
         "detail": "Build a preference system where users can set dietary restrictions (vegetarian, vegan, gluten-free, nut allergy, etc.) and the system filters/scores generated recipes accordingly. Use embeddings to enable 'find similar recipes' functionality. Consider integrating a nutrition database to estimate calories and macros per recipe."},
        {"name": "Cook-Along UI", "description": "Interactive web application with ingredient checklist, step-by-step cooking instructions, and built-in timers.",
         "deliverables": ["Upload/camera interface", "Detected ingredients display", "Recipe card with steps", "Built-in cooking timers", "Preference settings panel"],
         "tools": ["React/Next.js", "TailwindCSS", "Flask/FastAPI backend"], "vram": "N/A (frontend)",
         "detail": "Build a polished web application where users upload a fridge photo, see detected ingredients with bounding boxes, select preferences, and get a generated recipe. The cook-along view should show one step at a time with a timer for time-sensitive steps. Include a 'what's in my fridge?' dashboard showing detected items."}
    ],
    "outcomes": [
        "A working app that detects fridge ingredients and generates recipes",
        "Fine-tuned object detection model for food items",
        "LLM-powered recipe generation with dietary preference support",
        "Interactive cook-along interface with timers and step navigation"
    ],
    "demo": "Open a fridge, take a photo on your phone, upload it. Watch bounding boxes appear on each ingredient. Select 'vegetarian.' A complete recipe generates instantly with step-by-step instructions. Start the timer for step 1.",
    "guidelines": [
        "Start with YOLOv8 pre-trained on COCO — it already detects some food items (banana, apple, etc.)",
        "Fine-tune on a food dataset rather than training from scratch",
        "For recipe generation, invest time in prompt engineering — the prompt quality determines output quality",
        "Build the UI early and iterate — a good demo needs a smooth interface",
        "Test with at least 5 different fridge photos from different households"
    ],
    "milestones": [
        ("Foundation", "Days 1–5", "YOLOv8 fine-tuned on food items, basic recipe generation working, UI scaffolded"),
        ("Integration", "Days 6–10", "Full pipeline: upload → detect → generate → display, preferences working"),
        ("Polish", "Days 11–14", "Cook-along view, timers, nutrition info, 5+ test photos validated"),
        ("Demo Day", "Day 15", "Live demo with real fridge photo, recipe generation, cook-along walkthrough")
    ],
    "evaluation": [
        ("Working Demo", "35%", "Full pipeline from photo to recipe without errors"),
        ("Detection Quality", "25%", "How accurately does it identify ingredients?"),
        ("Recipe Quality", "20%", "Are generated recipes actually cookable and sensible?"),
        ("UI/UX", "10%", "Clean, intuitive cook-along interface"),
        ("Innovation", "10%", "Extra features: nutrition, meal planning, shopping list")
    ]
},
{
    "num": 3, "slug": "SoundForge", "name": "SoundForge — AI Music & Beat Generator",
    "tagline": "Describe a mood in text → AI generates original music. Layer tracks, add beats, export.",
    "difficulty": "Standard", "team_size": 4, "blaze_idx": 0,
    "problem": "Content creators, podcasters, and indie developers constantly need royalty-free background music. Hiring a composer is expensive, and stock music libraries feel generic. AI music generation can produce original, mood-matched tracks on demand — giving creators unique soundtracks tailored to their content.",
    "why_it_matters": "Music AI is the most defensible AI category (per a16z). Very few student projects tackle audio generation. The demo — typing a mood and hearing original music play — is magical and unique.",
    "objectives": [
        "Generate original music clips from text descriptions using AI models on GPU",
        "Implement beat detection, BPM matching, and audio mixing capabilities",
        "Add audio effects: reverb, EQ, and neural style transfer on audio",
        "Build a multi-track studio UI with waveform visualization and MP3 export"
    ],
    "data_input": "Text descriptions typed directly into the Selkies desktop application. No external file upload needed for music generation. Audio files for mixing/effects can be uploaded.",
    "modules": [
        {"name": "Music Generation Engine", "description": "AI model that generates original music from text prompts describing mood, genre, tempo, and instruments.",
         "deliverables": ["Text-to-music generation model", "Genre/mood conditioning", "8-30 second clip generation", "Multiple variation output"],
         "tools": ["MusicGen (via transformers/diffusers)", "AudioCraft", "PyTorch"], "vram": "~2.5 GB",
         "detail": "Use MusicGen or a similar model to generate music from text prompts. The model should accept descriptions like 'chill lo-fi beat with piano' and produce original audio. This is the most GPU-intensive module — it runs on the Blaze instance. Consider offering multiple generation modes: short clips (8s), medium (15s), and loop-friendly lengths."},
        {"name": "Beat Detection & Mixing", "description": "Audio DSP pipeline for BPM detection, beat alignment, and multi-track mixing.",
         "deliverables": ["BPM detection algorithm", "Beat grid alignment", "Multi-track mixer with volume/pan", "Audio format conversion"],
         "tools": ["librosa", "pydub", "soundfile", "scipy"], "vram": "~500 MB",
         "detail": "Build a DSP pipeline that can detect BPM, align beats across tracks, and mix multiple audio sources. This enables layering: a generated melody + a drum beat + a bass line. Use librosa for spectral analysis and pydub for audio manipulation."},
        {"name": "Audio Effects Processor", "description": "Neural and traditional audio effects: reverb, EQ, compression, and style transfer.",
         "deliverables": ["Reverb effect (convolution-based)", "Parametric EQ", "Audio style transfer", "Effect chain presets"],
         "tools": ["PyTorch", "scipy.signal", "librosa", "pedalboard"], "vram": "~400 MB",
         "detail": "Implement a chain of audio effects that can be applied to generated or uploaded audio. Include both traditional DSP effects (reverb, EQ, compression) and neural effects (audio style transfer). Consider creating presets like 'lo-fi', 'cinematic', 'upbeat'."},
        {"name": "Studio UI & Export", "description": "Multi-track audio editor with waveform visualization, effect controls, and MP3/WAV export.",
         "deliverables": ["Multi-track timeline", "Waveform visualization per track", "Effect parameter controls", "Export to MP3/WAV"],
         "tools": ["React", "Web Audio API", "wavesurfer.js", "Flask/FastAPI"], "vram": "N/A (frontend)",
         "detail": "Build a web-based audio studio where users can generate music, layer tracks, apply effects, and export. Use Web Audio API for real-time playback and waveform rendering. The UI should feel like a simplified DAW (digital audio workstation)."}
    ],
    "outcomes": ["A working AI music studio that generates and mixes original tracks", "MusicGen running on GPU for text-to-music", "Multi-track mixing with effects", "Exportable audio files"],
    "demo": "Type 'energetic Tamil folk fusion with drums.' Hit generate. Music starts playing. Add a bass layer. Apply reverb. Export the mix. Play it back.",
    "guidelines": ["MusicGen generates 8-30 second clips — design your UI around this limitation", "Use librosa for all audio analysis — it's the industry standard Python audio library", "Web Audio API is powerful for browser-based audio playback and visualization", "Generate multiple variations and let users pick their favorite", "Test with at least 10 different mood/genre prompts"],
    "milestones": [("Foundation", "Days 1–5", "MusicGen generating clips, basic audio DSP working, UI scaffolded"), ("Integration", "Days 6–10", "Multi-track mixing, effects chain, export working"), ("Polish", "Days 11–14", "Studio UI polished, presets, 10+ generated tracks validated"), ("Demo Day", "Day 15", "Live generation from audience prompt, layering, export")],
    "evaluation": [("Working Demo", "35%", "Generates and plays original music from text"), ("Audio Quality", "25%", "Does the generated music sound good?"), ("Mixing/Effects", "20%", "Can layer tracks and apply effects"), ("UI/UX", "10%", "Clean studio interface"), ("Innovation", "10%", "Extra: genre blending, loop generation, mood matching")]
},
{
    "num": 4, "slug": "FormCheck", "name": "FormCheck — AI Exercise Form Analyzer",
    "tagline": "Upload an exercise video → AI tracks your pose → generates a form report with corrections",
    "difficulty": "Standard", "team_size": 4, "blaze_idx": 1,
    "problem": "Most people exercise without knowing if their form is correct. Bad form leads to injuries and ineffective workouts. Personal trainers are expensive. AI pose estimation can analyze exercise form from a simple phone video and provide instant, detailed feedback on what to fix.",
    "why_it_matters": "Pose estimation is a highly valued CV skill. The demo — watching a skeleton overlay on a squat video with angle corrections at exact timestamps — is visually compelling and immediately useful.",
    "objectives": [
        "Implement body pose tracking using MediaPipe on uploaded exercise videos",
        "Build an exercise form analysis engine that detects common form mistakes",
        "Generate detailed reports with joint angles, timestamps, and correction suggestions",
        "Create a video playback UI with skeleton overlay and angle graphs"
    ],
    "data_input": "Students upload exercise videos recorded on their phones via the LaaS web UI. The container processes frames with GPU-accelerated pose estimation.",
    "modules": [
        {"name": "Video Processing Pipeline", "description": "Frame extraction, video management, and preprocessing for pose estimation.",
         "deliverables": ["Frame extraction at configurable FPS", "Video format handling", "Frame queue for processing"],
         "tools": ["OpenCV", "ffmpeg", "numpy"], "vram": "N/A (CPU)",
         "detail": "Build a robust video processing pipeline that extracts frames from uploaded videos at a configurable rate (e.g., 10 FPS for analysis). Handle various video formats and resolutions. The pipeline feeds frames to the pose estimation module."},
        {"name": "Pose Estimation Engine", "description": "Body pose tracking using MediaPipe on uploaded videos, producing 33-point skeleton data per frame.",
         "deliverables": ["33-point body landmark detection", "Per-frame skeleton data output", "Confidence scores per landmark", "Skeleton visualization rendering"],
         "tools": ["MediaPipe Pose", "OpenCV", "PyTorch"], "vram": "~300 MB",
         "detail": "Run MediaPipe Pose on every extracted frame to get 33 body landmarks (shoulders, elbows, hips, knees, ankles, etc.). Output structured skeleton data with confidence scores. This module runs on the Blaze instance since it processes every frame of potentially long videos."},
        {"name": "Form Analysis & Scoring", "description": "Joint angle computation, correct-form comparison, rep counting, and mistake detection.",
         "deliverables": ["Joint angle calculator (knee, hip, elbow, etc.)", "Correct-form templates per exercise", "Rep counting algorithm", "Mistake detection with severity scoring"],
         "tools": ["scipy", "numpy", "scikit-learn"], "vram": "~200 MB",
         "detail": "Analyze the skeleton data to compute joint angles, compare against correct form templates for each exercise (squat, pushup, deadlift, etc.), count reps, and detect common mistakes (knees caving, back rounding, insufficient depth). Output a structured report with timestamps and severity."},
        {"name": "Report UI & Visualization", "description": "Video player with skeleton overlay, angle graphs over time, AI-generated correction tips.",
         "deliverables": ["Video player with skeleton overlay", "Joint angle timeline graphs", "Correction suggestions panel", "Exercise selection menu", "Report export"],
         "tools": ["React/Next.js", "Recharts", "Flask/FastAPI", "MediaPipe drawing utils"], "vram": "N/A (frontend)",
         "detail": "Build a web app where users upload a video, see it played back with skeleton overlay, view angle graphs over time, and get AI-generated tips like 'At 0:15, your knee angle was 72° — aim for 90° for a proper squat.' Include exercise-specific templates."}
    ],
    "outcomes": ["A working exercise form analyzer with video upload and analysis", "MediaPipe pose estimation running on GPU", "Detailed form reports with angles and corrections", "Video playback with skeleton overlay"],
    "demo": "Upload a squat video. Watch the skeleton overlay appear. See 'Knee angle: 72° (should be 90°)' highlighted at the exact timestamp. Get 5 specific corrections. The audience tries it with their own video.",
    "guidelines": ["MediaPipe Pose works well without GPU training — it's pre-trained and fast", "Focus on 3-5 exercises initially (squat, pushup, plank, lunge, deadlift)", "Process at 10 FPS for analysis — you don't need every frame", "Build correct-form templates as angle ranges, not exact values", "Test with videos from different phone angles and lighting conditions"],
    "milestones": [("Foundation", "Days 1–5", "Pose estimation working, frame extraction pipeline, basic angle computation"), ("Integration", "Days 6–10", "Full pipeline: upload → analyze → report, form comparison working"), ("Polish", "Days 11–14", "UI with skeleton overlay, angle graphs, 3+ exercises supported"), ("Demo Day", "Day 15", "Live analysis of audience-submitted exercise video")],
    "evaluation": [("Working Demo", "35%", "Full pipeline from video upload to form report"), ("Analysis Quality", "25%", "Are angle measurements and corrections accurate?"), ("Visualization", "20%", "Skeleton overlay and graphs are clear and informative"), ("UI/UX", "10%", "Clean upload and report interface"), ("Innovation", "10%", "Extra: real-time tracking, exercise library, progress history")]
},
{
    "num": 5, "slug": "HawkEye", "name": "HawkEye — AI Video Intelligence Engine",
    "tagline": "Paste a YouTube URL or upload video → AI detects objects, counts people, finds anomalies → searchable event timeline",
    "difficulty": "Standard", "team_size": 4, "blaze_idx": 1,
    "problem": "Security cameras, traffic cameras, and surveillance footage generate thousands of hours of video that nobody watches. Finding specific events (a person entering, a car speeding, an anomaly) requires manually scrubbing through hours of footage. AI can process video automatically and create a searchable timeline of detected events.",
    "why_it_matters": "Video analytics is a massive industry (security, retail, traffic). The demo — pasting a YouTube URL and getting an instant event timeline with screenshots — is immediately impressive and practical.",
    "objectives": [
        "Build a video ingestion pipeline that downloads from URLs or accepts uploads",
        "Implement per-frame object/person/vehicle detection using YOLOv8 on GPU",
        "Create an event intelligence layer with anomaly detection and object tracking",
        "Build a searchable timeline dashboard with AI-annotated screenshots"
    ],
    "data_input": "Paste a YouTube/public URL (container downloads via yt-dlp) OR upload .mp4 files via the LaaS web UI. Videos are processed frame-by-frame on GPU.",
    "modules": [
        {"name": "Video Ingestion", "description": "Download videos from URLs or accept uploads, extract frames, manage video storage.",
         "deliverables": ["yt-dlp URL downloader", "File upload handler", "Frame extraction pipeline", "Video metadata extraction"],
         "tools": ["yt-dlp", "OpenCV", "ffmpeg", "Flask"], "vram": "N/A (CPU)",
         "detail": "Build a robust video ingestion system. For URLs, use yt-dlp to download from YouTube and other platforms. For uploads, handle various formats via ffmpeg. Extract frames at configurable FPS and store metadata (duration, resolution, frame count)."},
        {"name": "Object Detection Engine", "description": "YOLOv8-based per-frame detection of people, vehicles, animals, and objects with bounding boxes.",
         "deliverables": ["YOLOv8 inference pipeline", "Multi-class detection (80+ COCO classes)", "Per-frame detection results with confidence", "Detection overlay rendering"],
         "tools": ["Ultralytics YOLOv8", "OpenCV", "PyTorch"], "vram": "~1 GB",
         "detail": "Run YOLOv8 on every extracted frame to detect objects. Store results as structured data: frame number, timestamp, detected objects with bounding boxes and confidence scores. This is the most GPU-intensive module — it runs on the Blaze instance. Consider processing at 5-10 FPS to balance speed and coverage."},
        {"name": "Event Intelligence", "description": "Anomaly detection, object tracking across frames, event classification, and pattern recognition.",
         "deliverables": ["Object tracking (ID assignment across frames)", "Anomaly detection (unusual patterns)", "Event classification (person entered, vehicle stopped, crowd formed)", "Embedding-based similarity search"],
         "tools": ["Sentence-Transformers", "ChromaDB", "scikit-learn", "numpy"], "vram": "~300 MB",
         "detail": "Build an intelligence layer on top of detection results. Track objects across frames (assign IDs), detect anomalies (sudden crowd, vehicle speeding), and classify events. Use embeddings to enable 'find similar events' search functionality."},
        {"name": "Timeline Dashboard", "description": "Searchable event timeline with video scrubber, AI annotations, detection overlays, and export.",
         "deliverables": ["Event timeline with timestamps", "Video scrubber synced to detections", "Search/filter by object type", "Screenshot gallery per event", "Export report"],
         "tools": ["React/Next.js", "Socket.IO", "Flask/FastAPI", "ChromaDB"], "vram": "N/A (frontend)",
         "detail": "Build a web dashboard showing a timeline of detected events with screenshots. Users can scrub through the video with detection overlays, search for specific objects ('show me all frames with a red car'), and export event reports. Real-time progress updates via WebSocket during processing."}
    ],
    "outcomes": ["A working video intelligence platform", "YOLOv8 detection running on GPU", "Searchable event timeline with screenshots", "Anomaly detection and object tracking"],
    "demo": "Paste a traffic camera YouTube URL. Watch the AI download and process it. See a timeline populate with events: '3 people crossing at 0:45', 'red car speeding at 1:20'. Click any event to see the annotated screenshot.",
    "guidelines": ["Process at 5-10 FPS, not 30 FPS — you'll cover more video in less time", "Limit video length to 5 minutes during development to iterate faster", "Use YOLOv8n (nano) for speed, upgrade to YOLOv8s if accuracy is insufficient", "Build the timeline UI early — it's the main interface for the demo", "Test with diverse videos: traffic, people walking, indoor scenes"],
    "milestones": [("Foundation", "Days 1–5", "Video download/upload working, YOLOv8 detecting objects, basic UI"), ("Integration", "Days 6–10", "Full pipeline: URL → detect → timeline, event classification working"), ("Polish", "Days 11–14", "Search, filtering, anomaly detection, export, 5+ test videos validated"), ("Demo Day", "Day 15", "Live processing of audience-suggested YouTube URL")],
    "evaluation": [("Working Demo", "35%", "Full pipeline from URL to searchable timeline"), ("Detection Quality", "25%", "Accurate object detection with good coverage"), ("Intelligence", "20%", "Event classification and anomaly detection are useful"), ("UI/UX", "10%", "Clean timeline with good search"), ("Innovation", "10%", "Extra: real-time processing, alerts, multi-video support")]
},
{
    "num": 6, "slug": "GameBrain", "name": "GameBrain — AI That Learns to Play Classic Games",
    "tagline": "Pick a game → watch AI train from zero to superhuman in real-time → challenge it yourself",
    "difficulty": "Challenging", "team_size": 4, "blaze_idx": 0,
    "problem": "Reinforcement learning is one of the most exciting areas of AI, but it's rarely demonstrated in an interactive, visual way. Most RL projects are abstract graphs. By training an AI to play recognizable classic games (Snake, Pong, Flappy Bird) in real-time, you create a visceral demonstration of AI learning that anyone can understand and appreciate.",
    "why_it_matters": "Watching AI learn is genuinely exciting — the audience cheers when it gets smart. RL is a rare, impressive skill. The AI vs. Human challenge mode makes demo day interactive.",
    "objectives": [
        "Implement GPU-accelerated RL training using DQN/PPO algorithms",
        "Build 3 playable game environments (Snake, Pong, Flappy Bird) with GPU rendering",
        "Create real-time training visualization showing the neural network learning",
        "Build an arena UI with AI vs. Human mode, leaderboards, and training controls"
    ],
    "data_input": "No external data input. Games and training run entirely inside the LaaS container. Human players interact via keyboard in the Selkies desktop browser.",
    "modules": [
        {"name": "RL Training Engine", "description": "Deep Q-Network and PPO training pipeline with GPU acceleration for game-playing agents.",
         "deliverables": ["DQN agent implementation", "PPO agent implementation", "Experience replay buffer", "Training loop with GPU acceleration", "Model checkpointing"],
         "tools": ["Stable-Baselines3", "PyTorch", "Gymnasium"], "vram": "~800 MB",
         "detail": "Implement RL agents using DQN and PPO algorithms. This is the most GPU-intensive module — training involves thousands of forward and backward passes. Use Stable-Baselines3 for reliable implementations and PyTorch for custom architectures. Train agents to play each game from scratch, saving checkpoints at intervals."},
        {"name": "Game Environments", "description": "GPU-accelerated implementations of Snake, Pong, and Flappy Bird as Gymnasium environments.",
         "deliverables": ["Snake environment with Gymnasium API", "Pong environment", "Flappy Bird environment", "Configurable difficulty/speed", "State representation for RL"],
         "tools": ["Gymnasium", "Pygame", "numpy", "PyTorch"], "vram": "~500 MB",
         "detail": "Build 3 classic games as Gymnasium environments. Each must implement reset(), step(), and render() methods. The state representation should be suitable for RL — either raw pixel frames or extracted features. Use Pygame for rendering and make games playable by humans too."},
        {"name": "Training Visualization", "description": "Real-time visualization of the neural network learning, reward curves, and state-action analysis.",
         "deliverables": ["Live training progress chart", "Neural network weight visualization", "State-action heatmap", "Episode reward graph", "Comparison: AI vs. random play"],
         "tools": ["matplotlib", "numpy", "Streamlit/React", "Plotly"], "vram": "~300 MB",
         "detail": "Build visualizations that show the AI learning in real-time. Display reward curves climbing, the neural network's attention patterns, and state-action heatmaps showing what the AI is 'looking at' when making decisions. This makes RL tangible and exciting."},
        {"name": "Arena UI & Challenge Mode", "description": "Game viewer with AI vs. Human mode, training controls, and leaderboard.",
         "deliverables": ["Game canvas with AI/human toggle", "Training controls (start/stop/speed)", "Leaderboard (AI score vs. human scores)", "Model selection (different training stages)", "Replay mode"],
         "tools": ["React/HTML Canvas", "WebSocket", "Flask", "Pygame"], "vram": "N/A (frontend)",
         "detail": "Build an arena where users can watch the AI play, take over controls to play themselves, and compare scores. Include a 'training lab' where users can start training, adjust hyperparameters, and watch progress. The AI vs. Human challenge is the demo day highlight."}
    ],
    "outcomes": ["Working RL agents that learn to play 3 games", "Real-time training visualization", "AI vs. Human challenge mode", "Understanding of DQN/PPO algorithms demonstrated through working code"],
    "demo": "Start training Snake AI from zero. Watch it crash into walls. Fast-forward 2 minutes. Now it plays perfectly — never dying. Challenge: audience member plays Snake. AI score: 87. Human score: 12. The room erupts.",
    "guidelines": ["Start with Snake — it's the simplest environment to get working", "Use Stable-Baselines3's DQN — it's battle-tested and well-documented", "Training takes time — implement fast-forward mode (render every 100th episode)", "Save model checkpoints at intervals so you can show 'AI at 1 min' vs 'AI at 10 min'", "Make the human play experience smooth — keyboard controls should feel responsive"],
    "milestones": [("Foundation", "Days 1–5", "Snake environment working, DQN training on GPU, basic visualization"), ("Integration", "Days 6–10", "3 games working, training visualization live, arena UI functional"), ("Polish", "Days 11–14", "AI vs. Human mode, leaderboard, replay, all 3 games trained to superhuman level"), ("Demo Day", "Day 15", "Live training demo + AI vs. audience challenge")],
    "evaluation": [("Working Demo", "35%", "AI visibly learns and plays at superhuman level"), ("Training Quality", "25%", "AI actually gets good — not just random play"), ("Visualization", "20%", "Training progress is clearly visualized"), ("Interactivity", "10%", "AI vs. Human mode works smoothly"), ("Innovation", "10%", "Extra: custom games, multi-agent, curriculum learning")]
},
{
    "num": 7, "slug": "ArtForge", "name": "ArtForge — AI Art & Style Studio",
    "tagline": "Upload any photo → apply 10+ artistic styles → adjust intensity → export",
    "difficulty": "Standard", "team_size": 4, "blaze_idx": 0,
    "problem": "People love artistic versions of their photos — profile pictures, social media posts, creative projects. But using Photoshop or commissioning artists is expensive and slow. AI neural style transfer can transform any photo into a Van Gogh, Picasso, anime, or watercolor painting in seconds, with adjustable intensity.",
    "why_it_matters": "Style transfer is visually stunning and instantly shareable. The demo — clicking through 10 styles on a selfie in rapid succession — is addictive. Every output is unique and exportable.",
    "objectives": [
        "Implement neural style transfer with 10+ pre-trained artistic styles on GPU",
        "Build adjustable intensity control for blending original and styled images",
        "Add batch processing for applying multiple styles simultaneously",
        "Create a polished studio UI with style gallery, before/after, and export"
    ],
    "data_input": "Students upload photos via the LaaS web UI. GPU processes the style transfer in the container. Results displayed in the Selkies desktop browser.",
    "modules": [
        {"name": "Style Transfer Engine", "description": "Neural style transfer models for 10+ artistic styles (Van Gogh, Picasso, anime, watercolor, pop art, etc.).",
         "deliverables": ["10+ pre-trained style models", "GPU-accelerated inference", "Style model management", "Quality settings (draft/high)"],
         "tools": ["PyTorch", "torchvision (neural style)", "kornia", "Pillow"], "vram": "~1 GB",
         "detail": "Implement fast neural style transfer using pre-trained models. Each style (Van Gogh Starry Night, Picasso, anime, watercolor, etc.) is a separate model that transforms input photos. Use torchvision's transformer networks or kornia for differentiable image processing. Optimize for speed — each style should apply in under 2 seconds."},
        {"name": "Batch Processing & Intensity", "description": "Multi-style batch processing with adjustable intensity blending between original and styled image.",
         "deliverables": ["Batch style application (all styles at once)", "Intensity slider (0-100% blend)", "Resolution options", "Progress indicator"],
         "tools": ["PyTorch", "OpenCV", "Pillow", "numpy"], "vram": "~500 MB",
         "detail": "Build a processing pipeline that can apply all 10+ styles to an image in batch mode. Implement intensity control that blends between the original and styled image (0% = original, 100% = full style). Offer resolution options for faster previews vs. high-res exports."},
        {"name": "AI Art Curation", "description": "Style recommendation system using embeddings — 'find similar styles' and gallery organization.",
         "deliverables": ["Style embedding generation", "Similarity search", "Style recommendation", "Gallery organization"],
         "tools": ["Sentence-Transformers", "ChromaDB", "CLIP"], "vram": "~300 MB",
         "detail": "Use embeddings to create a 'style fingerprint' for each artistic style. Enable 'find similar styles' search — if you like Van Gogh, you might also like this impressionist style. Organize the gallery by style family (impressionist, modern, anime, etc.)."},
        {"name": "Studio UI", "description": "Polished web application with drag-drop upload, style gallery grid, intensity slider, and export.",
         "deliverables": ["Drag-drop upload", "Style gallery with previews", "Intensity slider with live preview", "Before/after comparison", "High-res export"],
         "tools": ["React/Next.js", "TailwindCSS", "Flask/FastAPI"], "vram": "N/A (frontend)",
         "detail": "Build a beautiful studio interface. Users upload a photo and see a grid of style previews. Click any style to apply it with an intensity slider. Include before/after comparison and high-res export. The UI should feel like a professional photo editing app."}
    ],
    "outcomes": ["Working style transfer with 10+ styles", "Adjustable intensity blending", "Batch processing capability", "Professional studio UI"],
    "demo": "Upload a selfie. Click through 10 styles in rapid succession — each renders in under 2 seconds. Drag the intensity slider. Export the best one. The audience wants to try it with their own photos.",
    "guidelines": ["Pre-trained style models are available from PyTorch Hub and various GitHub repos — don't train your own", "Optimize for speed: resize input to 512×512 for preview, full resolution only for export", "Collect at least 10 diverse style models: impressionist, cubist, anime, watercolor, pop art, sketch, etc.", "Build the intensity slider early — it's the key interactive feature", "Test with diverse photos: portraits, landscapes, objects, group shots"],
    "milestones": [("Foundation", "Days 1–5", "Style transfer working with 3+ styles, basic UI, intensity control"), ("Integration", "Days 6–10", "10+ styles, batch processing, style gallery, export working"), ("Polish", "Days 11–14", "UI polished, curation/recommendation, high-res export, 10+ test photos"), ("Demo Day", "Day 15", "Live styling of audience-submitted photos")],
    "evaluation": [("Working Demo", "35%", "Multiple styles apply quickly and look good"), ("Style Quality", "25%", "Artistic transformations are visually appealing"), ("Interactivity", "20%", "Intensity slider and gallery are smooth"), ("UI/UX", "10%", "Professional studio interface"), ("Innovation", "10%", "Extra: custom styles, video style transfer, social sharing")]
},
{
    "num": 8, "slug": "EchoScribe", "name": "EchoScribe — AI Video Transcription & Study Platform",
    "tagline": "Upload a lecture video or paste YouTube URL → AI transcribes, generates chapters, flashcards, and quizzes",
    "difficulty": "Standard", "team_size": 4, "blaze_idx": 0,
    "problem": "Students watch hours of lecture videos but struggle to find specific topics, create notes, or test their understanding. Manually transcribing and organizing video content is incredibly time-consuming. AI can automatically transcribe lectures, segment them into chapters, and generate study materials — turning passive video watching into active learning.",
    "why_it_matters": "Every student immediately understands this pain point. The demo — pasting a YouTube lecture URL and getting instant chapters + flashcards — is directly useful and impressive.",
    "objectives": [
        "Implement GPU-accelerated speech-to-text transcription with timestamps using Whisper",
        "Build an intelligence layer that segments transcripts into chapters and key topics",
        "Generate study materials: flashcards, summaries, and quiz questions from transcripts",
        "Create a synced video player with transcript, chapters, and study tools"
    ],
    "data_input": "Upload video files OR paste YouTube URLs (container downloads via yt-dlp). GPU processes audio transcription and LLM generates study materials.",
    "modules": [
        {"name": "Transcription Engine", "description": "GPU-accelerated speech-to-text with word-level timestamps and optional speaker diarization.",
         "deliverables": ["Whisper-based transcription", "Word-level timestamps", "Speaker diarization (who spoke when)", "Transcript export (TXT/SRT/JSON)"],
         "tools": ["faster-whisper", "ctranslate2", "PyTorch"], "vram": "~2 GB",
         "detail": "Use faster-whisper (CTranslate2-optimized) for fast, accurate transcription on GPU. Output word-level timestamps so the UI can sync video playback to transcript position. Consider adding speaker diarization for multi-speaker lectures. This is the most GPU-intensive module — runs on Blaze."},
        {"name": "Intelligence Layer", "description": "Topic segmentation, chapter detection, key moment extraction, and semantic search over transcripts.",
         "deliverables": ["Automatic chapter segmentation", "Topic labels per chapter", "Key moment detection", "Semantic search over transcript content"],
         "tools": ["Sentence-Transformers", "ChromaDB", "scikit-learn", "numpy"], "vram": "~300 MB",
         "detail": "Analyze the transcript to automatically segment it into logical chapters based on topic shifts. Generate descriptive labels for each chapter. Use embeddings to enable semantic search — 'find where the professor explained backpropagation.' Detect key moments (definitions, examples, important statements)."},
        {"name": "Study Tools Generator", "description": "AI-powered generation of flashcards, summaries, and quiz questions from lecture transcripts.",
         "deliverables": ["Flashcard generation (question/answer pairs)", "Chapter summaries", "Quiz questions with answers", "Difficulty levels"],
         "tools": ["Qwen2-1.5B (quantized)", "Transformers", "ChromaDB"], "vram": "~1.1 GB",
         "detail": "Use a quantized LLM to generate study materials from each chapter's transcript. Create flashcards (concept → definition), chapter summaries (key points in 3-5 bullets), and quiz questions (multiple choice with explanations). Offer difficulty levels for quizzes."},
        {"name": "Synced Video Player UI", "description": "Video player synced to transcript, chapter sidebar, flashcard viewer, and quiz interface.",
         "deliverables": ["Video player with transcript sync", "Chapter navigation sidebar", "Flashcard study mode", "Quiz interface with scoring", "URL input for YouTube"],
         "tools": ["React/Next.js", "react-player", "Flask/FastAPI", "TailwindCSS"], "vram": "N/A (frontend)",
         "detail": "Build a web app where users paste a YouTube URL or upload a video. The video plays with a synced transcript that highlights the current position. A chapter sidebar lets users jump to topics. Flashcard and quiz modes use the generated study materials."}
    ],
    "outcomes": ["Working video transcription with GPU-accelerated Whisper", "Automatic chapter segmentation", "AI-generated flashcards and quizzes", "Synced video player with study tools"],
    "demo": "Paste a YouTube lecture URL. Watch it transcribe rapidly with timestamps appearing as it processes. Chapters appear automatically. Click a chapter — get AI-generated flashcards. Take a quiz. The audience asks 'can I use this for my classes?'",
    "guidelines": ["faster-whisper is 4x faster than vanilla Whisper — always use it", "Process audio only (extract from video with ffmpeg) — don't feed video frames to Whisper", "Chapter segmentation can use simple heuristics: topic shifts detected via embedding similarity", "Start with English transcription, add multi-language as a stretch goal", "Test with at least 3 different lecture videos of varying length (5 min, 15 min, 30 min)"],
    "milestones": [("Foundation", "Days 1–5", "Whisper transcription working, basic UI with video + transcript"), ("Integration", "Days 6–10", "Chapter detection, flashcard generation, quiz mode working"), ("Polish", "Days 11–14", "Synced player polished, study tools refined, 3+ test videos validated"), ("Demo Day", "Day 15", "Live transcription of audience-suggested YouTube lecture")],
    "evaluation": [("Working Demo", "35%", "Full pipeline from URL to study materials"), ("Transcription Quality", "25%", "Accurate speech-to-text with good timestamps"), ("Study Materials", "20%", "Flashcards and quizzes are actually useful"), ("UI/UX", "10%", "Synced player is smooth and intuitive"), ("Innovation", "10%", "Extra: multi-language, note-taking, spaced repetition")]
},
{
    "num": 9, "slug": "VoxLingua", "name": "VoxLingua — Voice Translator",
    "tagline": "Upload audio in Tamil → AI transcribes, translates to English, and speaks it back",
    "difficulty": "Standard", "team_size": 4, "blaze_idx": 0,
    "problem": "Language barriers prevent communication in multilingual regions like Tamil Nadu. Professional translation services are expensive and slow. Real-time AI translation that handles speech input and produces spoken output can bridge language gaps instantly — enabling conversations between Tamil, English, and Hindi speakers.",
    "why_it_matters": "Speech AI (STT + translation + TTS) in a single pipeline is impressive and directly useful. The demo — speaking Tamil and hearing English — is a 'holy shit' moment. Very few student projects tackle multilingual speech AI.",
    "objectives": [
        "Implement GPU-accelerated speech-to-text for Tamil, English, and Hindi using Whisper",
        "Build context-aware translation using a quantized LLM",
        "Add natural-sounding text-to-speech output in the target language",
        "Create a conversation UI with transcript history and audio playback"
    ],
    "data_input": "Students upload audio files (.wav/.mp3) recorded on their phones via the LaaS web UI. GPU processes transcription and translation in the container.",
    "modules": [
        {"name": "Speech-to-Text Engine", "description": "GPU-accelerated transcription for Tamil, English, and Hindi using Whisper.",
         "deliverables": ["Multi-language transcription", "Language auto-detection", "Timestamp alignment", "Confidence scoring"],
         "tools": ["faster-whisper", "ctranslate2", "PyTorch"], "vram": "~2 GB",
         "detail": "Use faster-whisper for fast, accurate multi-language transcription. The model should auto-detect the input language and transcribe with timestamps. Whisper handles Tamil, English, and Hindi natively. This is the most GPU-intensive module — runs on Blaze."},
        {"name": "Translation Engine", "description": "Context-aware translation between Tamil, English, and Hindi using a quantized LLM.",
         "deliverables": ["Bidirectional translation (Tamil↔English, Hindi↔English)", "Context preservation", "Formal/informal register options"],
         "tools": ["Qwen2-1.5B (quantized)", "Transformers", "prompt engineering"], "vram": "~1.1 GB",
         "detail": "Use a quantized LLM for translation that preserves context and meaning (not word-by-word). Handle code-switching (mixing Tamil and English in the same sentence, common in Tamil Nadu). Consider adding formal/informal register options."},
        {"name": "Text-to-Speech Output", "description": "Natural-sounding speech synthesis in the target language.",
         "deliverables": ["TTS for English and Hindi", "Natural prosody and intonation", "Speed control", "Audio file export"],
         "tools": ["Piper TTS", "espeak-ng (fallback)", "pydub", "soundfile"], "vram": "~500 MB",
         "detail": "Use Piper TTS for natural-sounding speech output. Piper supports multiple languages and voices. Generate audio files that can be played back in the UI or exported. Include speed control for accessibility."},
        {"name": "Conversation UI", "description": "Audio upload interface with transcript view, translation panel, and audio playback.",
         "deliverables": ["Audio upload/player", "Transcript display (source + translated)", "Language selector", "Conversation history", "Audio playback of translation"],
         "tools": ["React/Next.js", "Web Audio API", "Flask/FastAPI"], "vram": "N/A (frontend)",
         "detail": "Build a clean conversation interface. Users upload audio, see the transcript appear, watch the translation generate, and hear it spoken aloud. Include a conversation history mode for multi-turn exchanges. Add language selector and audio speed controls."}
    ],
    "outcomes": ["Working multilingual speech translation pipeline", "Whisper transcription for Tamil/English/Hindi", "LLM-powered context-aware translation", "Natural TTS output"],
    "demo": "Upload a Tamil voice recording. Watch the Tamil transcript appear. See the English translation generate. Hear it spoken aloud in natural English. Then reverse: speak English, hear Tamil. The audience is stunned.",
    "guidelines": ["Whisper's 'medium' model handles Tamil well, but 'base' is faster — test both", "For translation, invest time in prompt engineering — 'Translate naturally, preserving tone and context'", "Piper TTS voices vary in quality — test multiple voices and pick the best one", "Handle audio format conversion (MP3→WAV) early in the pipeline", "Test with at least 5 audio clips: pure Tamil, pure English, code-switched, noisy background, quiet"],
    "milestones": [("Foundation", "Days 1–5", "Whisper transcription working, basic translation, TTS generating audio"), ("Integration", "Days 6–10", "Full pipeline: audio → transcribe → translate → speak, UI functional"), ("Polish", "Days 11–14", "Multi-language support, conversation history, 5+ test clips validated"), ("Demo Day", "Day 15", "Live translation of audience-submitted audio")],
    "evaluation": [("Working Demo", "35%", "Full speech translation pipeline works end-to-end"), ("Translation Quality", "25%", "Translations are natural and context-aware"), ("Audio Quality", "20%", "TTS sounds natural, not robotic"), ("UI/UX", "10%", "Clean conversation interface"), ("Innovation", "10%", "Extra: real-time streaming, conversation mode, dialect support")]
},
{
    "num": 10, "slug": "MemeForge", "name": "MemeForge — Context-Aware AI Meme Generator",
    "tagline": "Upload any image → AI understands the context → generates genuinely funny captions",
    "difficulty": "Standard", "team_size": 4, "blaze_idx": 1,
    "problem": "Creating good memes requires understanding image context and pairing it with the right caption. Most meme generators just overlay text on templates — they don't understand what's IN the image. An AI that actually sees the image and generates contextually funny captions would be the first truly intelligent meme creator.",
    "why_it_matters": "It's FUN. The demo makes people laugh. 'AI actually understood the image' is surprising and delightful. Meme culture is universal and the output is instantly shareable.",
    "objectives": [
        "Implement image understanding using CLIP and YOLOv8 to detect objects and scene context",
        "Build a caption generation system using an LLM that produces contextually funny text",
        "Create a meme template engine with intelligent text placement",
        "Build a studio UI with upload, template picker, caption editor, and sharing"
    ],
    "data_input": "Students upload images via the LaaS web UI or pick from trending meme templates. GPU analyzes the image and generates captions.",
    "modules": [
        {"name": "Image Understanding", "description": "Multi-model image analysis: object detection (YOLOv8) + scene understanding (CLIP) for context extraction.",
         "deliverables": ["Object detection with labels", "Scene description generation", "Context extraction (who, what, where)", "Mood/emotion detection"],
         "tools": ["CLIP (via transformers)", "Ultralytics YOLOv8", "OpenCV", "PyTorch"], "vram": "~1 GB (peak, models run sequentially)",
         "detail": "Combine YOLOv8 for object detection and CLIP for scene understanding. Extract a rich description of what's in the image: objects, people, actions, setting, mood. This context feeds the caption generator. Example: image of a cat on a laptop → 'cat, laptop, sitting, home office, cozy'."},
        {"name": "Caption Generation", "description": "LLM-powered generation of contextually relevant, funny captions based on image understanding.",
         "deliverables": ["Multiple caption suggestions per image", "Humor style options (sarcastic, wholesome, dark, relatable)", "Template-aware caption formatting"],
         "tools": ["Qwen2-1.5B (quantized)", "Transformers", "prompt engineering"], "vram": "~1.1 GB",
         "detail": "Use a quantized LLM with carefully crafted prompts to generate funny captions. The prompt should include the image context and humor style. Generate 5-10 caption options per image and let users pick their favorite. This runs on Blaze for the LLM inference."},
        {"name": "Template Engine", "description": "Meme template matching, text placement optimization, and image composition.",
         "deliverables": ["Template library (20+ popular formats)", "Intelligent template matching", "Text placement algorithm", "Font/style options"],
         "tools": ["Sentence-Transformers", "ChromaDB", "Pillow", "ImageMagick"], "vram": "~300 MB",
         "detail": "Build a template engine that matches images to appropriate meme formats. Use embeddings to find the best template for the image context. Handle text placement (top/bottom, centered, etc.) with proper font sizing and outline effects. Include 20+ popular meme templates."},
        {"name": "Meme Studio UI", "description": "Upload interface with template picker, caption editor, text drag-drop, and share functionality.",
         "deliverables": ["Image upload + template picker", "Caption editor with suggestions", "Drag-drop text positioning", "Export and share"],
         "tools": ["React/Next.js", "HTML Canvas", "Flask/FastAPI", "TailwindCSS"], "vram": "N/A (frontend)",
         "detail": "Build a fun, polished studio interface. Upload an image or pick a template. See AI-generated caption suggestions. Edit and position text with drag-drop. Export as PNG for sharing. The UI should feel playful and encourage experimentation."}
    ],
    "outcomes": ["Working AI meme generator with image understanding", "Context-aware caption generation", "Template engine with text placement", "Fun, shareable output"],
    "demo": "Upload a photo of the professor. AI detects: 'person, desk, classroom, presenting.' Generates 5 captions. Pick the funniest one. Place text on the image. Export. The room erupts in laughter.",
    "guidelines": ["CLIP is great for zero-shot image understanding — no training needed", "For caption generation, the prompt is everything — experiment with humor styles", "Start with 10 templates, expand to 20+ if time allows", "Make the UI fun and colorful — this project's vibe is playful", "Test with diverse images: people, animals, objects, scenes, screenshots"],
    "milestones": [("Foundation", "Days 1–5", "Image understanding working, basic caption generation, template rendering"), ("Integration", "Days 6–10", "Full pipeline: upload → understand → caption → render, UI functional"), ("Polish", "Days 11–14", "20+ templates, humor styles, text positioning, 10+ test images"), ("Demo Day", "Day 15", "Live meme generation from audience-submitted photos")],
    "evaluation": [("Working Demo", "35%", "Generates funny, contextually relevant memes"), ("Humor Quality", "25%", "Captions are actually funny, not generic"), ("Template Engine", "20%", "Text placement and rendering look professional"), ("UI/UX", "10%", "Fun, polished studio interface"), ("Innovation", "10%", "Extra: trending templates, social sharing, meme history")]
},
{
    "num": 11, "slug": "Scene3D", "name": "Scene3D — Single-Image 3D Depth Scanner",
    "tagline": "Upload any photo → AI estimates depth → generates a 3D point cloud you can rotate and fly through",
    "difficulty": "Challenging", "team_size": 4, "blaze_idx": 1,
    "problem": "3D scanning typically requires expensive hardware (LiDAR, depth cameras) or complex multi-view photogrammetry. But AI depth estimation can extract 3D structure from a single regular photo — democratizing 3D content creation. Architects, game developers, and VR creators need 3D scenes but lack the equipment to capture them.",
    "why_it_matters": "The 3D viewer is unlike anything other teams will have. 'Flat photo → 3D world' is magical. Depth estimation + 3D reconstruction are rare, impressive skills.",
    "objectives": [
        "Implement AI depth estimation from single images using Depth Anything V2 on GPU",
        "Convert depth maps into 3D point clouds using Open3D",
        "Add scene enhancement (upscaling) for better depth quality",
        "Build an interactive 3D viewer with orbit/fly controls and export"
    ],
    "data_input": "Students upload photos via the LaaS web UI. GPU processes depth estimation and 3D reconstruction in the container.",
    "modules": [
        {"name": "Depth Estimation Engine", "description": "AI-powered per-pixel depth prediction from a single RGB image.",
         "deliverables": ["Depth Anything V2 inference", "Per-pixel depth map output", "Confidence scoring", "Multiple resolution support"],
         "tools": ["Depth Anything V2", "PyTorch", "huggingface-hub"], "vram": "~500 MB",
         "detail": "Use Depth Anything V2 (or MiDaS) for monocular depth estimation. The model takes a single RGB image and outputs a depth map where each pixel has a depth value. This is the foundation for 3D reconstruction. Consider running at the image's native resolution for best quality."},
        {"name": "3D Point Cloud Generator", "description": "Convert depth maps into 3D point clouds with camera parameter estimation and mesh generation.",
         "deliverables": ["Depth-to-pointcloud conversion", "Camera intrinsic estimation", "Color-mapped 3D points", "Mesh generation (optional)"],
         "tools": ["Open3D", "numpy", "trimesh", "PyTorch"], "vram": "~500 MB",
         "detail": "Take the depth map and convert it into a 3D point cloud. Each pixel becomes a 3D point with x,y,z coordinates and color from the original image. Estimate camera intrinsics (focal length) from the image. Optionally generate a mesh from the point cloud for smoother surfaces. This runs on Blaze due to the computational load."},
        {"name": "Scene Enhancement", "description": "Upscale input images for better depth quality and add semantic understanding of the scene.",
         "deliverables": ["Image upscaling (Real-ESRGAN)", "Semantic segmentation", "Scene description"],
         "tools": ["Real-ESRGAN", "OpenCV", "sentence-transformers"], "vram": "~500 MB",
         "detail": "Enhance the input image before depth estimation. Upscale low-resolution images with Real-ESRGAN for better depth quality. Optionally add semantic segmentation to label scene elements (sky, building, road, person)."},
        {"name": "3D Viewer UI", "description": "Interactive web-based 3D viewer with orbit/fly controls, depth slider, and export.",
         "deliverables": ["Three.js 3D viewer", "Orbit/fly-through camera controls", "Depth map overlay toggle", "Export to .ply/.obj", "Depth adjustment slider"],
         "tools": ["Three.js", "React", "Flask/FastAPI", "WebGL"], "vram": "N/A (frontend)",
         "detail": "Build a web-based 3D viewer using Three.js. Users can orbit around the point cloud, fly through it, toggle between the original image and depth map, and export the 3D data. Include a depth adjustment slider to fine-tune the 3D effect."}
    ],
    "outcomes": ["Working depth estimation from single images", "3D point cloud generation", "Interactive web-based 3D viewer", "Exportable 3D data"],
    "demo": "Upload a landscape photo. Watch the depth map appear. Click 'Generate 3D.' A point cloud materializes. Rotate it. Fly through it. The audience has never seen this from a single photo.",
    "guidelines": ["Depth Anything V2 Small is fast and accurate — start with it", "Open3D is easier to use than PyTorch3D for point cloud operations", "For the 3D viewer, Three.js has excellent documentation and examples", "Export formats: .ply (point cloud), .obj (mesh) are widely supported", "Test with diverse images: indoor, outdoor, portraits, landscapes, objects"],
    "milestones": [("Foundation", "Days 1–5", "Depth estimation working, basic point cloud generation, simple 3D viewer"), ("Integration", "Days 6–10", "Full pipeline: upload → depth → 3D → view, enhancement working"), ("Polish", "Days 11–14", "3D viewer polished, export, fly-through, 10+ test images validated"), ("Demo Day", "Day 15", "Live 3D reconstruction of audience-submitted photo")],
    "evaluation": [("Working Demo", "35%", "Full pipeline from photo to interactive 3D"), ("Depth Quality", "25%", "Depth maps are accurate and detailed"), ("3D Viewer", "20%", "Smooth, interactive point cloud visualization"), ("Technical Depth", "10%", "Can explain depth estimation and 3D reconstruction"), ("Innovation", "10%", "Extra: mesh generation, video depth, AR export")]
},
{
    "num": 12, "slug": "StoryQuest", "name": "StoryQuest — AI Interactive Adventure Game",
    "tagline": "An AI Dungeon Master that creates adventures with voice narration AND generates an image for every scene",
    "difficulty": "Challenging", "team_size": 4, "blaze_idx": 1,
    "problem": "Interactive fiction (text adventures) has always been compelling, but traditional versions use pre-written content. AI can generate infinite, unique stories with branching narratives. Adding AI-generated images for each scene and voice narration transforms text adventures into immersive, cinematic experiences.",
    "why_it_matters": "Interactive fiction with AI images + voice is the most immersive demo possible. The audience literally plays the game during the demo. Combines 3 AI modalities: text generation, image generation, and speech.",
    "objectives": [
        "Build an LLM-powered story engine with branching narratives and character consistency",
        "Implement Stable Diffusion scene illustration for each story moment",
        "Add voice narration using Whisper (input) and Piper TTS (output)",
        "Create an interactive game UI with illustrations, choices, and story map"
    ],
    "data_input": "Text choices typed in the desktop browser. Voice input via uploaded audio files. All AI generation (story, images, voice) happens on GPU inside the container.",
    "modules": [
        {"name": "Story Engine", "description": "LLM-powered dynamic story generation with branching narratives, character memory, and world consistency.",
         "deliverables": ["Branching narrative generation", "Character consistency tracking", "World state management", "Choice consequence system"],
         "tools": ["Qwen2-1.5B (quantized)", "Transformers", "ChromaDB (for story memory)"], "vram": "~1.1 GB",
         "detail": "Use a quantized LLM to generate story content dynamically. The engine must maintain character consistency (the dragon you met in chapter 1 should remember you in chapter 3), track world state (if you took the sword, you have it), and generate meaningful choices with consequences. Use ChromaDB to store story memory for long-term consistency."},
        {"name": "Scene Illustration", "description": "Stable Diffusion image generation for each scene, creating unique visuals from story descriptions.",
         "deliverables": ["Scene description → image generation", "Consistent art style across scenes", "Character visualization", "Scene gallery"],
         "tools": ["Stable Diffusion 1.5 (diffusers)", "xformers", "PyTorch"], "vram": "~3.5 GB",
         "detail": "Generate a unique image for each story scene using Stable Diffusion. Convert the LLM's scene description into an image prompt. Use attention slicing and xformers to fit SD 1.5 in Blaze's 4GB VRAM. Maintain a consistent art style across all scenes. This is the most GPU-intensive module — generates one image per scene (~15-30 seconds each)."},
        {"name": "Voice Narration", "description": "Voice input for choices (Whisper) and AI voice narration of story text (Piper TTS).",
         "deliverables": ["Voice input transcription", "Story text → speech synthesis", "Character voice options", "Audio playback controls"],
         "tools": ["faster-whisper", "Piper TTS", "pydub", "soundfile"], "vram": "~1 GB",
         "detail": "Add voice interaction: players can speak their choices (transcribed by Whisper) and the story is narrated aloud (Piper TTS). Consider giving different characters different TTS voices. Generate audio for each story paragraph."},
        {"name": "Game UI", "description": "Interactive text adventure interface with scene illustrations, choice buttons, story map, and audio controls.",
         "deliverables": ["Story text display with illustrations", "Choice buttons", "Story map/progress tracker", "Audio narration controls", "Save/load game state"],
         "tools": ["React/Next.js", "Flask/FastAPI", "TailwindCSS"], "vram": "N/A (frontend)",
         "detail": "Build an immersive game interface. Story text appears with the generated scene illustration above it. Choice buttons let players decide what happens next. A story map shows progress and branching paths. Audio controls for narration. Include save/load for game state."}
    ],
    "outcomes": ["Working AI adventure game with infinite stories", "AI-generated scene illustrations", "Voice narration and input", "Interactive game UI"],
    "demo": "Start a new adventure. 'You stand at the edge of a dark forest.' An AI-generated image of the forest appears. Choose 'Enter the forest.' New scene generates with image. Narration speaks the story aloud. The audience plays along, making choices.",
    "guidelines": ["SD 1.5 with attention slicing fits in 4GB VRAM — but generates slowly (~15-30s/image). Design the game around this timing", "Use a consistent style prompt prefix for all scene images (e.g., 'fantasy art style, detailed')", "Keep story memory in ChromaDB — embed key events and characters for consistency", "Start with a simple story structure (3 chapters, 2-3 choices each) and expand", "Test with at least 5 different playthroughs to ensure variety"],
    "milestones": [("Foundation", "Days 1–5", "Story engine generating branches, SD generating images, basic UI"), ("Integration", "Days 6–10", "Full pipeline: choice → story → image → narration, game flow working"), ("Polish", "Days 11–14", "Story memory, character consistency, audio, save/load, 5+ playthroughs tested"), ("Demo Day", "Day 15", "Live adventure with audience making choices")],
    "evaluation": [("Working Demo", "35%", "Full interactive adventure with images and voice"), ("Story Quality", "25%", "Narratives are engaging and choices matter"), ("Image Quality", "20%", "Scene illustrations are visually appealing and relevant"), ("Immersion", "10%", "Voice narration enhances the experience"), ("Innovation", "10%", "Extra: multiplayer, character creation, world editor")]
},
{
    "num": 13, "slug": "SignReader", "name": "SignReader — Video Sign Language Translator",
    "tagline": "Upload a signing video → AI tracks hand poses → translates gestures to text → speaks the translation",
    "difficulty": "Challenging", "team_size": 4, "blaze_idx": 2,
    "problem": "Over 466 million people worldwide have disabling hearing loss. Sign language is their primary communication, but most people don't understand it. An AI that can watch a signing video and translate the gestures into spoken text would bridge a massive communication gap and enable accessibility.",
    "why_it_matters": "Accessibility AI is deeply meaningful. The demo — watching hand tracking translate sign language into spoken words — is powerful and emotional. Very few projects tackle this problem.",
    "objectives": [
        "Implement hand pose tracking using MediaPipe Hands on uploaded signing videos",
        "Build a gesture recognition model that classifies sign sequences into words/phrases",
        "Add text-to-speech output for spoken translation",
        "Create a video player UI with tracking overlay, subtitle display, and vocabulary lookup"
    ],
    "data_input": "Students upload signing videos recorded on their phones via the LaaS web UI. GPU processes hand tracking and gesture recognition in the container.",
    "modules": [
        {"name": "Video Processing", "description": "Frame extraction and preprocessing for hand tracking from uploaded videos.",
         "deliverables": ["Frame extraction pipeline", "Hand region detection", "Frame normalization"],
         "tools": ["OpenCV", "ffmpeg", "numpy"], "vram": "N/A (CPU)",
         "detail": "Build a video processing pipeline that extracts frames, detects hand regions, and normalizes frames for the hand tracking model. Handle various video formats, resolutions, and lighting conditions."},
        {"name": "Hand Tracking Engine", "description": "MediaPipe-based 21-point hand landmark detection on every frame of uploaded video.",
         "deliverables": ["21-point hand landmark detection per frame", "Both hands tracking simultaneously", "Landmark confidence scores", "Tracking visualization overlay"],
         "tools": ["MediaPipe Hands", "OpenCV", "PyTorch"], "vram": "~300 MB",
         "detail": "Run MediaPipe Hands on every frame to detect 21 landmarks per hand (fingertips, joints, palm). Track both hands simultaneously. Output structured landmark data with confidence scores. Render tracking overlay on video frames."},
        {"name": "Gesture Recognition", "description": "LSTM-based model that classifies sequences of hand poses into sign language words and phrases.",
         "deliverables": ["Gesture sequence classifier", "Vocabulary of 20+ signs", "Phrase-level recognition", "Confidence scoring per word"],
         "tools": ["PyTorch (LSTM)", "numpy", "scikit-learn"], "vram": "~500 MB",
         "detail": "Train an LSTM or Transformer-small model to classify sequences of hand landmarks into sign language words. Start with a vocabulary of 20+ common signs (hello, thank you, please, yes, no, etc.). The model takes a sequence of landmark frames and outputs the recognized word/phrase. This runs on Blaze for the sequence model inference."},
        {"name": "Translation UI", "description": "Video player with hand tracking overlay, synced subtitle display, vocabulary lookup, and TTS playback.",
         "deliverables": ["Video player with tracking overlay", "Real-time subtitle display", "Vocabulary lookup panel", "TTS playback of translation", "Sign dictionary reference"],
         "tools": ["React/Next.js", "Flask/FastAPI", "Piper TTS", "Canvas API"], "vram": "N/A (frontend)",
         "detail": "Build a web app where users upload a signing video and see it played back with hand tracking overlay. Recognized signs appear as subtitles synced to the video timing. Include a vocabulary lookup panel where users can see the meaning of each recognized sign. Add TTS to speak the translation aloud."}
    ],
    "outcomes": ["Working sign language recognition from video", "MediaPipe hand tracking on GPU", "LSTM gesture classification", "Video player with tracking overlay and subtitles"],
    "demo": "Upload a 30-second signing video. Watch hand tracking overlay appear. See words appear as subtitles: 'Hello... my name... is... student.' Hear the spoken translation. Look up any sign in the vocabulary panel.",
    "guidelines": ["MediaPipe Hands is pre-trained and fast — no GPU training needed for tracking", "For gesture recognition, start with a small vocabulary (10 signs) and expand", "Collect your own sign data: record yourself doing each sign 10-20 times for training", "Focus on clear, well-lit videos with hands visible in frame", "Use ASL (American Sign Language) as the base — it has the most available datasets"],
    "milestones": [("Foundation", "Days 1–5", "Hand tracking working, basic gesture classification, video player UI"), ("Integration", "Days 6–10", "Full pipeline: upload → track → recognize → display, 10+ signs recognized"), ("Polish", "Days 11–14", "20+ signs, TTS playback, vocabulary panel, 5+ test videos validated"), ("Demo Day", "Day 15", "Live translation of pre-recorded signing video")],
    "evaluation": [("Working Demo", "35%", "Recognizes signs and displays translation"), ("Tracking Quality", "25%", "Hand landmarks are accurate and stable"), ("Recognition Accuracy", "20%", "Correctly classifies gestures from the vocabulary"), ("UI/UX", "10%", "Clear overlay and subtitle display"), ("Innovation", "10%", "Extra: real-time mode, sentence-level recognition, sign dictionary")]
},
{
    "num": 14, "slug": "PixelPlayground", "name": "PixelPlayground — Creative AI Toolkit",
    "tagline": "A Swiss Army knife of AI visual tools: background remover, style transfer, upscaler, colorizer, depth mapper — all in one app",
    "difficulty": "Standard+", "team_size": 5, "blaze_idx": 1,
    "problem": "Creative professionals and hobbyists need various AI image tools — but they're scattered across different websites, apps, and APIs. Background removal is on one site, style transfer on another, upscaling on a third. A unified creative toolkit that puts 6+ AI visual tools in one app would be the ultimate creative workstation.",
    "why_it_matters": "6 GPU models in one app. Each team member owns a real AI tool. The unified UI makes it look like a professional product. The demo — one-click access to 6 different AI transformations — is addictive.",
    "objectives": [
        "Implement 5 distinct AI visual tools: background removal, style transfer, super-resolution, colorization, depth estimation",
        "Build a unified drag-drop canvas that chains tools together",
        "Add AI-powered tool suggestion based on image content",
        "Create a professional creative studio UI with tool palette and batch processing"
    ],
    "data_input": "Students upload photos via the LaaS web UI. Each AI tool processes the image on GPU. Results displayed in the unified canvas.",
    "modules": [
        {"name": "Background Removal", "description": "One-click background removal using TinySAM and rembg for object cutout and background swap.",
         "deliverables": ["Automatic background removal", "Object cutout with clean edges", "Background swap (solid color, blur, custom image)", "Edge refinement"],
         "tools": ["TinySAM", "rembg", "OpenCV", "Pillow"], "vram": "~500 MB",
         "detail": "Implement one-click background removal using TinySAM (segment anything, tiny version) and rembg. The tool should cleanly separate foreground objects from backgrounds. Offer background swap options: solid color, gaussian blur, or custom image."},
        {"name": "Style Transfer & Filters", "description": "Neural style transfer with 10+ artistic styles and adjustable intensity.",
         "deliverables": ["10+ style models", "Adjustable intensity", "Real-time preview", "Style gallery"],
         "tools": ["Fast Neural Style Transfer", "torchvision", "kornia", "Pillow"], "vram": "~1 GB",
         "detail": "Offer 10+ artistic style transfers: Van Gogh, Picasso, anime, watercolor, pop art, sketch, etc. Each style runs as a GPU-accelerated neural network. Include an intensity slider for blending between original and styled. This is the most GPU-intensive module — runs on Blaze."},
        {"name": "Super-Resolution & Colorization", "description": "4x image upscaling with Real-ESRGAN and B&W photo colorization with DeOldify.",
         "deliverables": ["4x upscaling", "B&W colorization", "Photo enhancement (sharpening, denoising)", "Before/after comparison"],
         "tools": ["Real-ESRGAN", "DeOldify", "OpenCV", "Pillow"], "vram": "~600 MB",
         "detail": "Combine two tools: Real-ESRGAN for 4x super-resolution upscaling and DeOldify for automatic B&W photo colorization. Also include basic photo enhancement (sharpening, denoising). Show before/after comparison for each operation."},
        {"name": "Depth Estimation & 3D Preview", "description": "AI depth map generation with 3D parallax effect and depth-based blur.",
         "deliverables": ["Depth map visualization", "3D parallax effect", "Depth-based blur (tilt-shift)", "Depth map export"],
         "tools": ["Depth Anything V2", "Open3D", "OpenCV", "Three.js"], "vram": "~500 MB",
         "detail": "Generate depth maps from any image using Depth Anything V2. Visualize the depth map as a heatmap. Create a 3D parallax effect where the image appears to have depth when you move the mouse. Offer depth-based blur (tilt-shift effect) and depth map export."},
        {"name": "Unified UI & AI Agent", "description": "Drag-drop canvas with tool palette, batch processing, and AI tool suggestion agent.",
         "deliverables": ["Drag-drop image canvas", "Tool palette sidebar", "Tool chaining (apply multiple tools)", "Batch processing queue", "AI suggests tools based on image"],
         "tools": ["React/Next.js", "TailwindCSS", "Flask/FastAPI", "sentence-transformers"], "vram": "~300 MB",
         "detail": "Build a professional creative studio UI. Drag-drop images onto a canvas. Select tools from a palette sidebar. Chain tools together (remove background → apply style → upscale). Batch process multiple images. An AI agent analyzes uploaded images and suggests which tools would be most useful ('This photo has a cluttered background — try background removal')."}
    ],
    "outcomes": ["5 distinct AI visual tools in one app", "Unified creative canvas with tool chaining", "Batch processing capability", "Professional studio UI"],
    "demo": "Upload a photo. Click 'Remove Background' — done in 1 second. Click 'Van Gogh Style' — done in 2 seconds. Click 'Upscale 4x' — done. Click 'Depth Map' — see the 3D parallax. Each tool is a different GPU model running. The audience wants to try every tool.",
    "guidelines": ["Each tool should work independently AND as part of a chain", "Build the UI framework early — the canvas is the main interface", "Pre-download all model weights on Day 1 to avoid delays later", "Test each tool with at least 5 different images before integrating", "The AI suggestion agent is a stretch goal — get all 5 tools working first"],
    "milestones": [("Foundation", "Days 1–5", "All 5 tools working independently, basic UI scaffolded"), ("Integration", "Days 6–10", "Unified canvas, tool chaining, batch processing working"), ("Polish", "Days 11–14", "UI polished, AI suggestions, 10+ test images per tool validated"), ("Demo Day", "Day 15", "Live demo: audience uploads photo, applies all tools in sequence")],
    "evaluation": [("Working Demo", "35%", "All 5 tools work and chain together"), ("Tool Quality", "25%", "Each AI tool produces good results"), ("Integration", "20%", "Tool chaining and batch processing are smooth"), ("UI/UX", "10%", "Professional creative studio interface"), ("Innovation", "10%", "Extra: AI suggestions, custom tool creation, plugin system")]
},
]
# fmt: on

# ── Main ─────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print(f"Generating {len(PROJECTS)} project documents...")
    for p in PROJECTS:
        build_doc(p)
    print(f"\nDone! All files saved to: {OUT_DIR}")
