import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Professional corporate colors: Navy & Slate theme
NAVY = RGBColor(26, 54, 93)       # #1A365D - Primary Accent
STEEL = RGBColor(43, 108, 176)    # #2B6CB0 - Secondary Accent
CHARCOAL = RGBColor(45, 55, 72)    # #2D3748 - Section Headers
BODY_TEXT = RGBColor(74, 85, 104)  # #4A5568 - Body Text
WHITE = RGBColor(255, 255, 255)
LIGHT_GRAY = "F7FAFC"             # #F7FAFC - Alternating rows
BORDER_GRAY = "E2E8F0"            # #E2E8F0 - Table borders
LINE_COLOR = "1A365D"             # Navy for accents

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_borders(cell, color_hex="E2E8F0", width="4"):
    """Set cell borders with specified color and width."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), width)
        border.set(qn('w:color'), color_hex)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    """Set cell padding in dxa (twentieths of a point). 140 dxa = 7pt, 180 dxa = 9pt."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        margin = OxmlElement(f'w:{margin_name}')
        margin.set(qn('w:w'), str(val))
        margin.set(qn('w:type'), 'dxa')
        tcMar.append(margin)
    tcPr.append(tcMar)

def style_document(doc):
    """Apply general document styling, margins, and body text formatting."""
    # Base normal style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = BODY_TEXT
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(6)
    
    # 1-inch margins
    for section in doc.sections:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)

def add_cover_page(doc):
    """Create a modern cover page with navy accents and bold professional hierarchy."""
    # Top spacing
    for _ in range(3):
        doc.add_paragraph()
    
    # Elegant Navy accent bar (using table trick)
    accent_table = doc.add_table(rows=1, cols=1)
    accent_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    accent_cell = accent_table.rows[0].cells[0]
    accent_cell.width = Inches(6.5)
    set_cell_shading(accent_cell, LINE_COLOR)
    p = accent_cell.paragraphs[0]
    p.add_run(" ")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    
    doc.add_paragraph()
    
    # Program Code or Tag
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("L&D INTERNSHIP PROGRAMME")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = STEEL
    p.paragraph_format.space_after = Pt(12)
    
    # Main Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("30-Day Computer Science Foundations")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = NAVY
    p.paragraph_format.space_after = Pt(6)
    
    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("A First-Principles Curriculum for Bridging Academia and Engineering Excellence")
    run.font.size = Pt(14)
    run.font.italic = True
    run.font.color.rgb = CHARCOAL
    p.paragraph_format.space_after = Pt(24)
    
    # Add vertical spacing
    for _ in range(4):
        doc.add_paragraph()
        
    # Metadata Block Table (Centered)
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Prepared For:", "Computer Science B.Tech Interns"),
        ("Program Duration:", "30 Days (4 Weeks + Wrap-up)"),
        ("Key Methodology:", "First-Principles, Architectural Case Studies, Hands-on CLI & AI Tools"),
        ("Document Version:", "v1.1.0 (Restructured - June 2026)")
    ]
    
    for r_idx, (label, val) in enumerate(meta_data):
        row = meta_table.rows[r_idx]
        
        cell_lbl = row.cells[0]
        set_cell_shading(cell_lbl, "F8FAFC")
        set_cell_borders(cell_lbl, "E2E8F0", "2")
        p_lbl = cell_lbl.paragraphs[0]
        r_lbl = p_lbl.add_run(label)
        r_lbl.font.bold = True
        r_lbl.font.size = Pt(9.5)
        r_lbl.font.color.rgb = CHARCOAL
        cell_lbl.width = Inches(1.8)
        set_cell_margins(cell_lbl, 80, 80, 100, 100)
        
        cell_val = row.cells[1]
        set_cell_shading(cell_val, "FFFFFF")
        set_cell_borders(cell_val, "E2E8F0", "2")
        p_val = cell_val.paragraphs[0]
        r_val = p_val.add_run(val)
        r_val.font.size = Pt(9.5)
        r_val.font.color.rgb = BODY_TEXT
        cell_val.width = Inches(4.7)
        set_cell_margins(cell_val, 80, 80, 100, 100)
        
    doc.add_page_break()

def add_section_header(doc, text):
    """Add a professional section header with a navy left-accent indicator."""
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.allow_autofit = False
    
    # Accent indicator bar (Left)
    accent_cell = table.rows[0].cells[0]
    accent_cell.width = Inches(0.12)
    set_cell_shading(accent_cell, LINE_COLOR)
    accent_p = accent_cell.paragraphs[0]
    accent_p.add_run(" ")
    accent_p.paragraph_format.space_before = Pt(0)
    accent_p.paragraph_format.space_after = Pt(0)
    
    # Text cell (Right)
    text_cell = table.rows[0].cells[1]
    text_cell.width = Inches(6.38)
    p = text_cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    
    doc.add_paragraph() # Spacing below heading

def add_curriculum_table(doc, items_data):
    """Generate the structured grid table with merged category visual dividers."""
    table = doc.add_table(rows=1 + len(items_data), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    
    headers = ["Day", "Curated Topic & Sub-topics", "First Principles & Value Add"]
    col_widths = [0.8, 2.2, 3.5] # Total = 6.5 inches
    
    # Table Header Row
    header_row = table.rows[0]
    for i, title in enumerate(headers):
        cell = header_row.cells[i]
        set_cell_shading(cell, LINE_COLOR)
        set_cell_borders(cell, LINE_COLOR, "4")
        set_cell_margins(cell, 160, 160, 180, 180)
        p = cell.paragraphs[0]
        run = p.add_run(title)
        run.font.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = WHITE
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        cell.width = Inches(col_widths[i])
        
    day_row_count = 0 # To track alternating colors for day rows
    
    # Fill Table Rows
    for r_idx, item in enumerate(items_data):
        row = table.rows[r_idx + 1]
        
        if item["type"] == "category":
            # Visual Divider Row: Merge all columns
            merged_cell = row.cells[0].merge(row.cells[1]).merge(row.cells[2])
            set_cell_shading(merged_cell, "2B6CB0") # Steel Blue background
            set_cell_borders(merged_cell, "2B6CB0", "4")
            set_cell_margins(merged_cell, 140, 140, 180, 180)
            
            p = merged_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(item["title"])
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = WHITE
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            
            merged_cell.width = Inches(6.5)
            
        else:
            # Alternating background color for day rows
            bg_color = "FFFFFF" if day_row_count % 2 == 0 else LIGHT_GRAY
            day_row_count += 1
            
            day_str = item["day"]
            topic_title = item["title"]
            subtopics = item["subtopics"]
            value_add = item["value"]
            
            # 1. Day Column
            cell_day = row.cells[0]
            set_cell_shading(cell_day, bg_color)
            set_cell_borders(cell_day, BORDER_GRAY, "4")
            set_cell_margins(cell_day, 120, 120, 140, 140)
            p_day = cell_day.paragraphs[0]
            run_day = p_day.add_run(day_str)
            run_day.font.bold = True
            run_day.font.size = Pt(10)
            run_day.font.color.rgb = CHARCOAL
            p_day.paragraph_format.space_after = Pt(0)
            cell_day.width = Inches(col_widths[0])
            
            # 2. Topic & Subtopics Column
            cell_topic = row.cells[1]
            set_cell_shading(cell_topic, bg_color)
            set_cell_borders(cell_topic, BORDER_GRAY, "4")
            set_cell_margins(cell_topic, 120, 120, 140, 140)
            
            p_topic = cell_topic.paragraphs[0]
            run_topic = p_topic.add_run(topic_title)
            run_topic.font.bold = True
            run_topic.font.size = Pt(10)
            run_topic.font.color.rgb = NAVY
            p_topic.paragraph_format.space_after = Pt(4)
            
            for sub in subtopics:
                p_sub = cell_topic.add_paragraph()
                run_sub = p_sub.add_run(f"▪  {sub}")
                run_sub.font.size = Pt(8.5)
                run_sub.font.color.rgb = BODY_TEXT
                p_sub.paragraph_format.space_after = Pt(2)
                p_sub.paragraph_format.left_indent = Inches(0.1)
                
            cell_topic.width = Inches(col_widths[1])
            
            # 3. Value Add Column
            cell_val = row.cells[2]
            set_cell_shading(cell_val, bg_color)
            set_cell_borders(cell_val, BORDER_GRAY, "4")
            set_cell_margins(cell_val, 120, 120, 140, 140)
            p_val = cell_val.paragraphs[0]
            run_val = p_val.add_run(value_add)
            run_val.font.size = Pt(9.5)
            run_val.font.color.rgb = BODY_TEXT
            p_val.paragraph_format.space_after = Pt(0)
            cell_val.width = Inches(col_widths[2])
            
    doc.add_paragraph()

def build_document():
    doc = Document()
    style_document(doc)
    
    # 1. Title Page
    add_cover_page(doc)
    
    # 2. Executive Summary / Philosophy
    add_section_header(doc, "Program Introduction & Core Philosophy")
    
    p = doc.add_paragraph()
    r = p.add_run(
        "Modern computer science education often oscillates between raw academic theory and surface-level "
        "practical tutorials. When engineering interns join a professional technology company, they frequently "
        "suffer from a transition gap: knowing how to write syntax, but lacking an understanding of execution "
        "realities, underlying system internals, and scaling bottlenecks. "
    )
    r.font.size = Pt(10.5)
    
    p2 = doc.add_paragraph()
    r2 = p2.add_run(
        "This 30-Day L&D Curriculum is designed specifically for incoming Computer Science B.Tech interns. "
        "Its objective is not to teach end-to-end framework syntax, but to establish first-principles, build "
        "mental models, and demystify the architectural layers of software systems. Each day represents a critical "
        "conceptual pivot, structured like a progressive 'river of knowledge'. Interns advance from local execution systems "
        "(OS & Linux) to networks and databases, moving through cloud infrastructure, container orchestration, core code patterns (DSA), "
        "scalable system design, architectural case studies, physical interfaces (IoT), and modern developer workflows (GenAI, n8n, AI agents). "
        "They wrap up with collaboration standards and career counselling."
    )
    r2.font.size = Pt(10.5)
    
    doc.add_paragraph()
    
    # 3. The Curriculum Grid
    add_section_header(doc, "30-Day Daywise Schedule")
    
    # Detailed Daywise Structure with visual Category markers
    curriculum_items = [
        # PART I: OPERATING SYSTEMS & UNIX/LINUX FOUNDATIONS (DAYS 1-3)
        {"type": "category", "title": "PART I: OPERATING SYSTEMS & UNIX/LINUX FOUNDATIONS (DAYS 1-3)"},
        {"type": "day", "day": "Day 1", "title": "Operating Systems: Core Execution", "subtopics": [
            "Processes vs. Threads", "CPU Scheduling mechanisms", "Virtual Memory & Paging", "Context switching overhead", "Thread safety & race conditions"
        ], "value": "Destroys the 'black box' view of how programs execute, showing how to write thread-safe and resource-efficient code that works in harmony with the OS."},
        {"type": "day", "day": "Day 2", "title": "UNIX & Linux: CLI Fundamentals", "subtopics": [
            "Linux Filesystem Hierarchy (FHS)", "Terminal navigation & core shell", "Standard streams (stdin, stdout, stderr)", "Redirection and Piping (|)", "Text filtering (Grep, Awk, Sed)"
        ], "value": "Transitions students from GUI interfaces to command-line execution, unlocking terminal productivity and laying the groundwork for backend debugging."},
        {"type": "day", "day": "Day 3", "title": "Linux Systems & Admin Automation", "subtopics": [
            "User permissions (chmod, chown)", "Process management (ps, kill, top)", "Environment variables & configuration", "Writing automation bash scripts", "System automation via Cron"
        ], "value": "Empowers interns to automate environment setups, manage system permissions, and handle remote deployments without external supervision."},

        # PART II: NETWORKING & WEB PROTOCOLS (DAYS 4-5)
        {"type": "category", "title": "PART II: NETWORKING & WEB PROTOCOLS (DAYS 4-5)"},
        {"type": "day", "day": "Day 4", "title": "Networking: OSI & Data Flow", "subtopics": [
            "OSI Model vs. TCP/IP stack", "TCP 3-way handshake & termination", "UDP vs. TCP protocol trade-offs", "IP addressing, Subnetting & CIDR", "DNS resolution workflow"
        ], "value": "Builds a robust mental model of data transmission, which is critical for analyzing API latencies, network configurations, and packet loss."},
        {"type": "day", "day": "Day 5", "title": "Web Networking & App Layer Protocols", "subtopics": [
            "HTTP/HTTPS methods, headers, status codes", "SSL/TLS secure handshake", "RESTful API design constraints", "WebSockets vs. SSE (Server-Sent Events)"
        ], "value": "Provides an understanding of secure, stateless web communication, allowing interns to build standardized REST and real-time APIs."},

        # PART III: DATABASE MANAGEMENT SYSTEMS & STORAGE (DAYS 6-7)
        {"type": "category", "title": "PART III: DATABASE MANAGEMENT SYSTEMS & STORAGE (DAYS 6-7)"},
        {"type": "day", "day": "Day 6", "title": "Relational Databases (RDBMS) & SQL", "subtopics": [
            "ACID transactional guarantees", "SQL Joins and Subqueries", "Database Indexing (B-Trees)", "Query Execution Plan analysis", "PostgreSQL Architecture (MVCC)"
        ], "value": "Establishes database transactional consistency. Interns learn to design robust tables and optimize queries to avoid production-halting database locks."},
        {"type": "day", "day": "Day 7", "title": "NoSQL Databases & Distributed Data", "subtopics": [
            "CAP Theorem trade-offs", "Document Databases (MongoDB)", "Key-Value caching (Redis)", "Wide-Column stores (ScyllaDB / Cassandra)", "Eventual vs. Strong Consistency"
        ], "value": "Equips interns with the knowledge of when to swap traditional ACID guarantees for horizontal scale, high throughput, and fault tolerance."},

        # PART IV: CLOUD ARCHITECTURE & DEVOPS LIFECYCLE (DAYS 8-12)
        {"type": "category", "title": "PART IV: CLOUD ARCHITECTURE & DEVOPS LIFECYCLE (DAYS 8-12)"},
        {"type": "day", "day": "Day 8", "title": "Cloud Computing: Core Infrastructure", "subtopics": [
            "Virtualization vs. Bare Metal", "Cloud service models (IaaS, PaaS, SaaS)", "AWS/GCP core services (EC2/GCE, S3)", "Elasticity & cloud billing models"
        ], "value": "Shifts their developer mindset from running scripts on localhost to deploying elastic, highly available infrastructure across global server networks."},
        {"type": "day", "day": "Day 9", "title": "Cloud Security & Isolation Systems", "subtopics": [
            "Virtual Private Cloud (VPC)", "Subnets, Gateways & Route Tables", "Security Groups & Network ACLs", "IAM Roles & Principle of Least Privilege"
        ], "value": "Teaches security-first cloud architecture. Interns learn how to isolate networks and lock down resource permissions before deployment."},
        {"type": "day", "day": "Day 10", "title": "DevOps: Containerization with Docker", "subtopics": [
            "Containers vs. Virtual Machines", "Docker Engine & architecture", "Writing optimized Dockerfiles", "Docker Networking & Volume mounts", "Multi-container setups (Docker Compose)"
        ], "value": "Eliminates the classic 'works on my machine' bug, packaging code and all runtime dependencies into self-contained, predictable environments."},
        {"type": "day", "day": "Day 11", "title": "Container Orchestration with Kubernetes", "subtopics": [
            "Control Plane vs. Worker Nodes", "Pods, Services, and Deployments", "Declarative config via YAML", "Self-healing & Horizontal Autoscaling"
        ], "value": "Exposes interns to how major companies manage thousands of microservices dynamically, introducing auto-healing and zero-downtime rolling updates."},
        {"type": "day", "day": "Day 12", "title": "Infrastructure as Code & Observability", "subtopics": [
            "Terraform declarative syntax & states", "Monitoring vs. Observability", "Metrics collection (Prometheus)", "Visualizing logs & states (Grafana)", "Defining actionable alert rules"
        ], "value": "Enforces infrastructure-as-code discipline, while teaching interns how to track resource metrics and catch application issues before users report them."},

        # PART V: DATA STRUCTURES & ALGORITHM DESIGN PATTERNS (DAYS 13-14)
        {"type": "category", "title": "PART V: DATA STRUCTURES & ALGORITHM DESIGN PATTERNS (DAYS 13-14)"},
        {"type": "day", "day": "Day 13", "title": "DSA Foundations: Time-Space Complexity", "subtopics": [
            "Big-O notation & growth rates", "Memory layout of data structures", "Arrays vs. Linked Lists", "Cache locality first principles", "Stack vs. Heap allocations"
        ], "value": "Demystifies how computer memory handles data structures at the hardware level, training interns to write memory-efficient code from day one."},
        {"type": "day", "day": "Day 14", "title": "Advanced Data Structures & Indexing", "subtopics": [
            "Hash Tables & collision resolution", "Trees (BST, Trie structures)", "Graph traversal (DFS & BFS)", "Recursion vs. Iteration limits"
        ], "value": "Explains the mechanics of constant-time O(1) lookups and details the algorithms underlying search indexes, database lookups, and routing tables."},

        # PART VI: SYSTEM DESIGN & MASSIVE SCALING PRINCIPLES (DAYS 15-17)
        {"type": "category", "title": "PART VI: SYSTEM DESIGN & MASSIVE SCALING PRINCIPLES (DAYS 15-17)"},
        {"type": "day", "day": "Day 15", "title": "System Design: Scaling Foundations", "subtopics": [
            "Horizontal vs. Vertical scaling", "Load Balancer algorithms (L4 vs. L7)", "Caching strategies (Write-Through, Eviction)", "Content Delivery Networks (CDNs)"
        ], "value": "Teaches the principles of scaling, enabling interns to design services that withstand heavy user traffic spikes without system crashes."},
        {"type": "day", "day": "Day 16", "title": "System Design: Decoupled Systems", "subtopics": [
            "Monoliths to Microservices", "API Gateway routing & rate limiting", "Message Queues (RabbitMQ)", "Event Streaming (Apache Kafka)", "Sync vs. Async processing pattern"
        ], "value": "Teaches how to design distributed microservices, ensuring that a slow component or background database lock doesn't take down the entire system."},
        {"type": "day", "day": "Day 17", "title": "System Design: Fault Tolerance & HA", "subtopics": [
            "Database replication topologies", "Database Sharding & Partitioning", "Circuit Breaker pattern", "Graceful degradation & failover"
        ], "value": "Provides interns with the architectural tools to ensure software remains highly available and resilient during system outages or partial failures."},

        # PART VII: REAL-WORLD ARCHITECTURAL CASE STUDIES (DAYS 18-20)
        {"type": "category", "title": "PART VII: REAL-WORLD ARCHITECTURAL CASE STUDIES (DAYS 18-20)"},
        {"type": "day", "day": "Day 18", "title": "Case Study 1: Uber (Geospatial Dispatch)", "subtopics": [
            "Uber's dispatch system architecture", "Geospatial indexing with H3 Hexagons", "Real-time tracking via WebSockets/gRPC", "Matchmaking lock mechanics"
        ], "value": "Applies theoretical data structures and networking to a global scale business case, demonstrating how spatial geometry indexes solve ridesharing dispatch."},
        {"type": "day", "day": "Day 19", "title": "Case Study 2: Discord (1 Trillion Messages)", "subtopics": [
            "MongoDB scaling bottlenecks", "Migrating to Cassandra / ScyllaDB", "Designing optimal partitioning keys", "Write-heavy database schema optimizations"
        ], "value": "Exposes interns to actual storage internals, showing how database selection and primary key design impact read/write speeds when scaling."},
        {"type": "day", "day": "Day 20", "title": "Case Study 3: Notion (Real-time Editor)", "subtopics": [
            "Operational Transformation (OT) concepts", "Conflict-free Replicated Data Types (CRDTs)", "Block-based data models", "Real-time state synchronization"
        ], "value": "Demystifies the real-time collaboration engines used in modern editors, explaining how conflicts are resolved across multiple concurrent browser connections."},

        # PART VIII: PRODUCT ENGINEERING & SOLUTIONS ARCHITECTURE (DAY 21)
        {"type": "category", "title": "PART VIII: PRODUCT ENGINEERING & SOLUTIONS ARCHITECTURE (DAY 21)"},
        {"type": "day", "day": "Day 21", "title": "Design Thinking & Solutions Architecture", "subtopics": [
            "Analyzing customer problem statements", "Translating product requirements to tech specs", "Drafting RFCs & System Design documents", "Architectural cost vs. benefit analysis"
        ], "value": "Expands their focus from raw code to product engineering, training interns to understand business goals and draft clean technical blueprints."},

        # PART IX: INTERNET OF THINGS (IOT) & EDGE COMPUTING (DAYS 22-23)
        {"type": "category", "title": "PART IX: INTERNET OF THINGS (IOT) & EDGE COMPUTING (DAYS 22-23)"},
        {"type": "day", "day": "Day 22", "title": "IoT & Embedded Systems Foundations", "subtopics": [
            "Microcontrollers (ESP32) vs. SBCs (Raspberry Pi)", "Analog vs. Digital signals", "GPIO programming first principles", "Sensor integration and serial protocols"
        ], "value": "Connects software logic to the physical world, teaching interns how computers interface with physical sensors and environment controllers."},
        {"type": "day", "day": "Day 23", "title": "IoT Protocol Stack & Edge Ingestion", "subtopics": [
            "Publish-Subscribe pattern", "MQTT protocol vs. HTTP/CoAP", "Edge gateways & local pre-processing", "Time-series database ingestion"
        ], "value": "Teaches lightweight network protocols. Interns understand how to optimize communication in resource-constrained or low-bandwidth environments."},

        # PART X: GENERATIVE AI & AGENTIC WORKFLOW AUTOMATION (DAYS 24-27)
        {"type": "category", "title": "PART X: GENERATIVE AI & AGENTIC WORKFLOW AUTOMATION (DAYS 24-27)"},
        {"type": "day", "day": "Day 24", "title": "GenAI: LLM Developer Fundamentals", "subtopics": [
            "LLM architectures (Transformers)", "Tokenization, context windows, & cost", "Prompt engineering design patterns", "Low-code automation with n8n"
        ], "value": "Teaches interns to build functional AI solutions by programmatically linking LLMs with external APIs and automated business workflows."},
        {"type": "day", "day": "Day 25", "title": "AI Agents & Local LLM Deployments", "subtopics": [
            "Agentic autonomy & ReAct pattern", "Running local models (Ollama, Hermes/LLaMA)", "Function calling & external tool bindings", "Agent runtime security & execution sandboxing"
        ], "value": "Introduces the latest frontier in software engineering: building autonomous, local AI agents capable of choosing actions and tools."},
        {"type": "day", "day": "Day 26", "title": "Code Agents & Self-Debugging Loops", "subtopics": [
            "Autonomous code generation", "Terminal feedback loop execution", "Self-debugging & unit testing", "OpenCode Interpreter workflow analysis"
        ], "value": "Exposes interns to iterative coding agents, demonstrating how to use AI pairs to write, run, and self-correct code from errors."},
        {"type": "day", "day": "Day 27", "title": "Advanced Agent Systems & RAG", "subtopics": [
            "Semantic Search & Embeddings", "Vector Databases (Chroma/pgvector)", "Retrieval-Augmented Generation (RAG) pipelines", "Multi-agent orchestration (OpenClaw)"
        ], "value": "Combines vector database search with agent pipelines, showing how to construct secure, enterprise-grade AI systems that query private data without hallucinations."},

        # PART XI: TEAM COLLABORATION, OPEN SOURCE & CAREER LAUNCH (DAYS 28-30)
        {"type": "category", "title": "PART XI: TEAM COLLABORATION, OPEN SOURCE & CAREER LAUNCH (DAYS 28-30)"},
        {"type": "day", "day": "Day 28", "title": "Software Best Practices & Team Work", "subtopics": [
            "Git branching models & pull requests", "Merge conflict resolution strategies", "Continuous Integration (CI) test automation", "Code review etiquettes & refactoring"
        ], "value": "Prepares interns for enterprise engineering, ensuring their code contributions are structured, peer-reviewed, and automatically tested."},
        {"type": "day", "day": "Day 29", "title": "Open Source & GSOC Counseling", "subtopics": [
            "Understanding Open Source licensing", "Finding 'good first issues' on GitHub", "Google Summer of Code (GSoC) roadmap", "Building a public technical brand"
        ], "value": "Encourages community contribution, giving interns a blueprint for building a strong, globally recognized portfolio using open-source projects."},
        {"type": "day", "day": "Day 30", "title": "Internship Reflection & Career Launch", "subtopics": [
            "Presenting a first-principles project", "Drafting a 6-month post-internship roadmap", "Resume & LinkedIn review for developers", "Lifelong self-learning strategies"
        ], "value": "Synthesizes the entire 30-day journey into a structured career plan, leaving the interns with a clear trajectory for continuous growth."}
    ]
    
    add_curriculum_table(doc, curriculum_items)
    
    # 4. Weekly Milestones & Expected Outcomes
    add_section_header(doc, "Weekly Learning Outcomes & Milestones")
    
    outcomes = [
        "Week 1 (OS, Linux, Networking & Databases - Days 1–7): Interns master CLI command-line navigation and script automation. They exit this week understanding OS process execution, standard network transmission models, and relational vs. non-relational database design concepts.",
        "Week 2 (Cloud, DevOps, IaC & DSA - Days 8–14): Interns shift to cloud environments, package software using Docker, manage containers in Kubernetes, structure infrastructure with Terraform, and deep-dive into time-space complexity metrics and algorithm structures.",
        "Week 3 (System Design, Case Studies & Solutions Design - Days 15–21): Interns design scalable, decoupled architectures using load balancers and asynchronous messaging queues. They analyze massive architectural studies (Uber, Discord, Notion) and author professional RFC design documents.",
        "Week 4 (IoT, AI Agents & Career Launch - Days 22–30): Interns bridge hardware-software boundaries, deploy local LLM instances to run autonomous agentic execution loops, collaborate via standard team Git flows, and define a clear 6-month career plan."
    ]
    
    for outcome in outcomes:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(outcome)
        run.font.size = Pt(10)
        run.font.color.rgb = BODY_TEXT
        p.paragraph_format.space_after = Pt(4)
        
    doc.add_paragraph()
    
    # Save the output file
    output_path = r"c:\Users\Punith\LaaS\CS_Internship_30_Day_Curriculum.docx"
    try:
        doc.save(output_path)
        print(f"SUCCESS: Curriculum document generated at: {output_path}")
    except Exception as e:
        print(f"ERROR: Error saving document: {e}")

if __name__ == "__main__":
    build_document()
