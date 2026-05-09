#!/usr/bin/env python3
"""
Generate a professional DOCX incident report from markdown source.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
DARK_HEADER = "2C3E50"          # Dark slate for table headers
ACCENT_COLOR = "1ABC9C"         # Teal accent
RED_SEVERITY = "C0392B"         # Red for High / BLOCKED
GREEN_STATUS = "27AE60"         # Green for WORKED
LIGHT_GRAY = "F2F3F4"           # Alternating row background
CUTOUT_BG = "FEF9E7"            # Light yellow for cutoff box
CUTOUT_BORDER = "F39C12"        # Orange border for cutoff
SECTION_BORDER = "1ABC9C"       # Teal for section left border
FONT_NAME = "Calibri"

INPUT_MD = r"c:\Users\Punith\LaaS\ReadMe\KSRCE-Network-Proxy-Issue-Report.md"
OUTPUT_DOCX = r"c:\Users\Punith\LaaS\ReadMe\KSRCE-Network-Proxy-Issue-Report.docx"


def set_cell_shading(cell, color_hex):
    """Apply solid background shading to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    tcPr.append(shd)


def set_cell_border(cell, color_hex, size="12", space="0"):
    """Apply a bottom border to a cell (used for section headers)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:space="{space}" w:color="{color_hex}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)


def set_paragraph_border(paragraph, color_hex, size="12", space="4"):
    """Apply a left border to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBorders = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="{size}" w:space="{space}" w:color="{color_hex}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBorders)


def set_paragraph_shading(paragraph, color_hex):
    """Apply background shading to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    pPr.append(shd)


def set_run_color(run, color_hex):
    """Set font color for a run."""
    run.font.color.rgb = RGBColor.from_string(color_hex)


def add_run_text(paragraph, text, bold=False, italic=False, size=11, color=None, font_name=FONT_NAME):
    """Helper to add a formatted run to a paragraph."""
    run = paragraph.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        set_run_color(run, color)
    return run


def add_section_header(doc, text, level=2):
    """Add a bold section header with a teal left border."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    set_paragraph_border(p, SECTION_BORDER, size="24", space="6")
    add_run_text(p, text, bold=True, size=14, color=DARK_HEADER)
    return p


def add_normal_paragraph(doc, text, bold_parts=None, indent_left=False):
    """
    Add a normal paragraph, optionally making certain substrings bold.
    bold_parts: list of (substring, color_hex, make_bold) tuples.
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if indent_left:
        p.paragraph_format.left_indent = Inches(0.2)

    if bold_parts:
        cursor = 0
        parts = sorted(bold_parts, key=lambda x: text.find(x[0], cursor) if text.find(x[0], cursor) != -1 else len(text))
        for sub, color, is_bold in parts:
            idx = text.find(sub, cursor)
            if idx == -1:
                continue
            if idx > cursor:
                add_run_text(p, text[cursor:idx])
            add_run_text(p, sub, bold=is_bold, color=color)
            cursor = idx + len(sub)
        if cursor < len(text):
            add_run_text(p, text[cursor:])
    else:
        add_run_text(p, text)
    return p


def add_bullet_paragraph(doc, text, bold_parts=None):
    """Add a bullet paragraph."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    if bold_parts:
        cursor = 0
        for sub, color, is_bold in bold_parts:
            idx = text.find(sub, cursor)
            if idx == -1:
                continue
            if idx > cursor:
                add_run_text(p, text[cursor:idx])
            add_run_text(p, sub, bold=is_bold, color=color)
            cursor = idx + len(sub)
        if cursor < len(text):
            add_run_text(p, text[cursor:])
    else:
        add_run_text(p, text)
    return p


def add_code_block(doc, code_lines):
    """Add a shaded code block."""
    for line in code_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.line_spacing = 1.1
        set_paragraph_shading(p, "F4F6F7")
        add_run_text(p, line, size=10, font_name="Consolas")


def style_table_header(row, col_count):
    """Dark header with white text."""
    for cell in row.cells:
        set_cell_shading(cell, DARK_HEADER)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(11)
                run.font.name = FONT_NAME
                run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)


def style_table_rows(table, start_row=1):
    """Apply alternating row shading."""
    for i, row in enumerate(table.rows):
        if i < start_row:
            continue
        color = LIGHT_GRAY if (i - start_row) % 2 == 1 else "FFFFFF"
        for cell in row.cells:
            set_cell_shading(cell, color)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(10.5)
                    run.font.name = FONT_NAME
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)


def highlight_status_in_table(table):
    """Color WORKED green and BLOCKED red inside table cells."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                full_text = paragraph.text
                if "WORKED" in full_text:
                    paragraph.clear()
                    # Rebuild preserving other text
                    parts = re.split(r'(WORKED)', full_text)
                    for part in parts:
                        if part == "WORKED":
                            add_run_text(paragraph, part, bold=True, color=GREEN_STATUS)
                        else:
                            add_run_text(paragraph, part)
                elif "BLOCKED" in full_text:
                    paragraph.clear()
                    parts = re.split(r'(BLOCKED)', full_text)
                    for part in parts:
                        if part == "BLOCKED":
                            add_run_text(paragraph, part, bold=True, color=RED_SEVERITY)
                        else:
                            add_run_text(paragraph, part)
                elif "FIRST SIGN OF BLOCK" in full_text:
                    paragraph.clear()
                    parts = re.split(r'(FIRST SIGN OF BLOCK)', full_text)
                    for part in parts:
                        if part == "FIRST SIGN OF BLOCK":
                            add_run_text(paragraph, part, bold=True, color=RED_SEVERITY)
                        else:
                            add_run_text(paragraph, part)


def add_footer(doc):
    """Add a CONFIDENTIAL footer to every section."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.clear()
        add_run_text(p, "CONFIDENTIAL — INTERNAL USE ONLY", italic=True, size=9, color="7F8C8D")


def add_header_banner(doc):
    """Add a professional header banner with title and metadata."""
    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_p.paragraph_format.space_after = Pt(4)
    add_run_text(title_p, "KSRCE Network Incident — Sophos Firewall Blocking Software Downloads",
                 bold=True, size=18, color=DARK_HEADER)

    # Separator line
    sep = doc.add_paragraph()
    sep.paragraph_format.space_after = Pt(6)
    pPr = sep._p.get_or_add_pPr()
    pBorders = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="12" w:space="4" w:color="{ACCENT_COLOR}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBorders)

    # Metadata line
    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(12)
    add_run_text(meta, "Date: ", bold=True, size=11, color=DARK_HEADER)
    add_run_text(meta, "May 3, 2026    ", size=11)
    add_run_text(meta, "From: ", bold=True, size=11, color=DARK_HEADER)
    add_run_text(meta, "Punith — LaaS Platform Deployment    ", size=11)
    add_run_text(meta, "Severity: ", bold=True, size=11, color=DARK_HEADER)
    add_run_text(meta, "High    ", bold=True, size=11, color=RED_SEVERITY)
    add_run_text(meta, "Status: ", bold=True, size=11, color=DARK_HEADER)
    add_run_text(meta, "Blocked — Awaiting KSRCE IT Admin action", bold=True, size=11, color=RED_SEVERITY)

    doc.add_paragraph()  # spacer


def add_cutoff_box(doc):
    """Render the cutoff summary as a prominent bordered/shaded block."""
    box = doc.add_paragraph()
    box.paragraph_format.space_before = Pt(8)
    box.paragraph_format.space_after = Pt(12)
    box.paragraph_format.left_indent = Inches(0.2)
    box.paragraph_format.right_indent = Inches(0.2)
    set_paragraph_shading(box, CUTOUT_BG)

    # Top border
    pPr = box._p.get_or_add_pPr()
    pBorders = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="18" w:space="4" w:color="{CUTOUT_BORDER}"/>'
        f'  <w:left w:val="single" w:sz="18" w:space="4" w:color="{CUTOUT_BORDER}"/>'
        f'  <w:bottom w:val="single" w:sz="18" w:space="4" w:color="{CUTOUT_BORDER}"/>'
        f'  <w:right w:val="single" w:sz="18" w:space="4" w:color="{CUTOUT_BORDER}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBorders)

    lines = [
        ("LAST CONFIRMED WORKING:   ", True, DARK_HEADER),
        ("May 1, 5:55 PM IST  (ai2 unattended-upgrades)\n", False, DARK_HEADER),
        ("FIRST CONFIRMED FAILING:  ", True, RED_SEVERITY),
        ("May 2, 6:53 AM IST  (ai2 \"Distribution outdated\")\n", False, DARK_HEADER),
        ("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n", False, CUTOUT_BORDER),
        ("  FIREWALL POLICY CHANGE:  ", True, RED_SEVERITY),
        ("Between May 1 ~6 PM  →  May 2 ~7 AM IST (overnight)\n", False, DARK_HEADER),
        ("━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━", False, CUTOUT_BORDER),
    ]
    for text, bold, color in lines:
        add_run_text(box, text, bold=bold, color=color, size=11)


def build_affected_machines_table(doc):
    """Affected Machines table."""
    add_section_header(doc, "Affected Machines")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = Inches(1.8)
    table.columns[1].width = Inches(1.4)
    table.columns[2].width = Inches(2.4)
    table.columns[3].width = Inches(1.6)

    hdr = table.rows[0].cells
    hdr[0].text = "Machine"
    hdr[1].text = "LAN IP"
    hdr[2].text = "SSH Access"
    hdr[3].text = "GPU"
    style_table_header(table.rows[0], 4)

    data = [
        ["ai1 (aiserver1)", "20.1.1.130", "ssh ai1@103.115.236.52 -p 2223", "RTX 5090 32GB"],
        ["ai2 (aiserver2)", "20.1.1.132", "ssh ai2@103.115.236.52 -p 2224", "RTX 5090 32GB"],
        ["ai4", "20.1.1.x", "Similar SSH", "RTX 5090 32GB"],
    ]
    for row_data in data:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val
    style_table_rows(table)

    add_normal_paragraph(doc, "All machines on 20.1.1.x/16 subnet, default gateway 20.1.1.1.")


def build_evidence_table(doc):
    """Evidence for Cross-Verification table."""
    add_section_header(doc, "Evidence for Cross-Verification")
    add_normal_paragraph(doc,
        "All evidence below is from tamper-proof system logs (root-owned, kernel-level) that the IT admin can independently verify by SSHing into the machines.")

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = Inches(2.2)
    table.columns[1].width = Inches(1.2)
    table.columns[2].width = Inches(2.8)
    table.columns[3].width = Inches(1.6)

    hdr = table.rows[0].cells
    hdr[0].text = "Evidence"
    hdr[1].text = "Machine"
    hdr[2].text = "File / Command"
    hdr[3].text = "Tamper-proof?"
    style_table_header(table.rows[0], 4)

    data = [
        ["apt install/purge history with timestamps", "ai1 & ai2",
         "/var/log/apt/history.log and /var/log/apt/history.log.1.gz",
         "Yes — root-owned, system-managed"],
        ["apt-daily service run timestamps", "ai1",
         "journalctl — filter for apt-daily",
         "Yes — kernel-level systemd journal"],
        ["unattended-upgrades pass/fail log", "ai2",
         "/var/log/unattended-upgrades/unattended-upgrades.log",
         "Yes — system service log"],
        ["Current block status", "ai1 & ai2",
         "sudo apt update (live test)",
         "Reproducible right now"],
        ["Gateway identification", "ai1 & ai2",
         "ip route show default + curl -s http://20.1.1.1:8090/ | head -20",
         "Reproducible right now"],
        ["Sophos SSL cert injection", "ai1 & ai2",
         "wget --spider https://developer.download.nvidia.com/...",
         "Reproducible right now"],
    ]
    for row_data in data:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val
    style_table_rows(table)

    add_normal_paragraph(doc, "All timestamps in system logs are UTC. Add 5 hours 30 minutes for IST.",
                         bold_parts=[("All timestamps in system logs are UTC.", RED_SEVERITY, True)])


def build_timeline_table(doc, title, headers, data_rows, col_widths):
    """Generic timeline table builder."""
    add_run_text(doc.add_paragraph(), title, bold=True, size=12, color=DARK_HEADER)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    for i, w in enumerate(col_widths):
        table.columns[i].width = Inches(w)

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    style_table_header(table.rows[0], len(headers))

    for row_data in data_rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val
    style_table_rows(table)
    highlight_status_in_table(table)
    doc.add_paragraph()  # spacer


def build_timeline_ai1(doc):
    """Full Incident Timeline for ai1."""
    add_section_header(doc, "Full Incident Timeline (IST)")
    add_run_text(doc.add_paragraph(), "ai1 (aiserver1 — 20.1.1.130)", bold=True, size=12, color=DARK_HEADER)

    # Apr 30
    add_run_text(doc.add_paragraph(), "Apr 30 — All operations worked:", bold=True, size=11, color=DARK_HEADER)
    headers1 = ["IST Time", "Command", "Status"]
    data1 = [
        ["2:26 PM", "apt full-upgrade -y", "WORKED"],
        ["2:26 PM", "apt install build-essential cmake git curl wget ... (14 packages)", "WORKED"],
        ["2:36 PM", "apt install nvidia-driver-570 nvidia-utils-570", "WORKED"],
        ["3:50 PM", "apt purge nvidia* → apt autoremove", "WORKED"],
        ["3:56 PM", "apt install nvidia-driver-570 nvidia-utils-570 (2nd attempt)", "WORKED"],
        ["4:47 PM", "apt purge nvidia* → apt autoremove", "WORKED"],
        ["4:48 PM", "apt install nvidia-driver-570 nvidia-utils-570 (3rd attempt)", "WORKED"],
        ["4:57 PM", "apt purge nvidia*", "WORKED"],
        ["5:44 PM", "apt autoremove", "WORKED"],
        ["5:45 PM", "apt install nvidia-driver-570 nvidia-utils-570 (4th attempt)", "WORKED"],
        ["5:50 PM", "apt purge nvidia* → apt autoremove", "WORKED"],
        ["5:52 PM", "apt install nvidia-driver-570-open", "WORKED — last successful install"],
    ]
    build_timeline_table(doc, "", headers1, data1, [1.1, 3.8, 2.0])

    # May 1
    add_run_text(doc.add_paragraph(), "May 1 — Automated system services still worked (no human intervention):",
                 bold=True, size=11, color=DARK_HEADER)
    headers2 = ["IST Time", "Service", "Status"]
    data2 = [
        ["1:46 AM", "apt-daily.service — daily package list download", "WORKED"],
        ["11:43 AM", "apt-daily-upgrade.service — daily upgrade check", "WORKED"],
    ]
    build_timeline_table(doc, "", headers2, data2, [1.1, 3.8, 2.0])

    # May 2 onwards
    add_run_text(doc.add_paragraph(), "May 2 onwards — Everything blocked:", bold=True, size=11, color=RED_SEVERITY)
    headers3 = ["IST Time", "Command", "Status"]
    data3 = [
        ["May 3, 11:09 AM", "sudo apt update", "BLOCKED — 403 Forbidden from 20.1.1.1:8090"],
        ["May 3, 11:09 AM", "wget https://developer.download.nvidia.com/...", "BLOCKED — Sophos SSL CA certificate injection"],
        ["May 3, 11:18 AM", "curl http://archive.ubuntu.com/ubuntu/", "BLOCKED — connection hangs silently"],
    ]
    build_timeline_table(doc, "", headers3, data3, [1.4, 3.5, 2.0])


def build_timeline_ai2(doc):
    """Full Incident Timeline for ai2."""
    add_run_text(doc.add_paragraph(), "ai2 (aiserver2 — 20.1.1.132)", bold=True, size=12, color=DARK_HEADER)

    # Apr 21–30
    add_run_text(doc.add_paragraph(), "Apr 21–30 — All operations worked:", bold=True, size=11, color=DARK_HEADER)
    headers1 = ["Date (IST)", "Time", "Command", "Status"]
    data1 = [
        ["Apr 21", "4:21 PM", "OS install (grub-efi, linux-generic, openssh-server)", "WORKED"],
        ["Apr 21", "11:38 AM", "apt full-upgrade", "WORKED"],
        ["Apr 21", "12:03 PM", "apt upgrade", "WORKED"],
        ["Apr 21", "12:07 PM", "apt install apache2", "WORKED"],
        ["Apr 21", "2:32 PM", "apt install neofetch", "WORKED"],
        ["Apr 22", "12:23 PM", "unattended-upgrades (4 runs)", "WORKED"],
        ["Apr 25", "12:25 PM", "unattended-upgrades (2 runs)", "WORKED"],
        ["Apr 29", "11:48 AM", "unattended-upgrades (3 runs)", "WORKED"],
        ["Apr 29", "2:18 PM", "apt install iperf3", "WORKED"],
        ["Apr 30", "12:22 PM", "unattended-upgrades", "WORKED"],
        ["Apr 30", "2:26 PM", "apt full-upgrade -y", "WORKED"],
        ["Apr 30", "2:26 PM", "apt install build-essential cmake git curl wget ...", "WORKED"],
        ["Apr 30", "2:36 PM", "apt install nvidia-driver-570 nvidia-utils-570", "WORKED"],
    ]
    build_timeline_table(doc, "", headers1, data1, [0.9, 0.9, 3.3, 1.8])

    # May 1
    add_run_text(doc.add_paragraph(), "May 1 — Automated services still reached repos:", bold=True, size=11, color=DARK_HEADER)
    headers2 = ["IST Time", "Service", "Status"]
    data2 = [
        ["12:00 PM", "unattended-upgrades — \"No packages to upgrade\"", "WORKED — repo was reachable"],
        ["5:55 PM", "unattended-upgrades — \"No packages to upgrade\"", "WORKED — last confirmed repo access"],
    ]
    build_timeline_table(doc, "", headers2, data2, [1.1, 3.8, 2.0])

    # May 2
    add_run_text(doc.add_paragraph(), "May 2 — Block begins:", bold=True, size=11, color=RED_SEVERITY)
    headers3 = ["IST Time", "Service", "Status"]
    data3 = [
        ["6:53 AM", "unattended-upgrades — \"Distribution outdated\" warning", "FIRST SIGN OF BLOCK"],
        ["11:55 AM", "unattended-upgrades — upgraded 4 security packages", "Partial — security.ubuntu.com only"],
    ]
    build_timeline_table(doc, "", headers3, data3, [1.1, 3.8, 2.0])


def main():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    # Default font for the document
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(11)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

    # =================== HEADER BANNER ===================
    add_header_banner(doc)

    # =================== WHAT'S HAPPENING ===================
    add_section_header(doc, "What's Happening")
    add_normal_paragraph(doc,
        "The Sophos/Cyberoam firewall at the KSRCE campus gateway (20.1.1.1) is blocking all outbound access to software repositories and download servers from the AI lab servers. This is not just Ubuntu's package manager — all software downloads are affected, including NVIDIA drivers, CUDA toolkit, Docker, and any wget/curl to external software sources.")

    bullets_wh = [
        ("apt install / apt update → returns 403 Forbidden from the proxy at 20.1.1.1:8090", RED_SEVERITY, True),
        ("wget to NVIDIA servers → fails with Sophos SSL certificate injection (the firewall replaces the real SSL cert with its own, which our servers reject as untrusted)", RED_SEVERITY, True),
        ("curl to Ubuntu repos (HTTP) → connection hangs silently after TCP connect — no response", RED_SEVERITY, True),
        ("curl to Google → works perfectly fine — proving the internet itself is available, only software sources are selectively blocked", GREEN_STATUS, True),
    ]
    for btext, color, is_bold in bullets_wh:
        add_bullet_paragraph(doc, btext)

    add_normal_paragraph(doc,
        "The firewall is applying a web category filter (Category 68) that classifies software repositories as restricted content.")

    # =================== IMPACT ===================
    add_section_header(doc, "Impact")
    add_normal_paragraph(doc,
        "This has completely halted the KSRCE deployment. We cannot:",
        bold_parts=[("This has completely halted the KSRCE deployment.", RED_SEVERITY, True)])

    impact_items = [
        "Install or reinstall NVIDIA GPU drivers (ai1 and ai2 both need this)",
        "Install CUDA Toolkit for GPU compute",
        "Install Docker for container orchestration",
        "Install any system packages required for the LaaS platform",
        "Receive security updates — servers are exposed to known vulnerabilities",
    ]
    for item in impact_items:
        add_bullet_paragraph(doc, item)

    add_normal_paragraph(doc,
        "This is a repeated pattern of network-level hindrance that has nothing to do with our application or the machine setup itself. The servers are correctly configured — there are no local proxy settings, no misconfigured DNS, no firewall rules on the machines. The block is happening at the network gateway level, entirely outside our control.")
    add_normal_paragraph(doc,
        "Every day this remains unresolved is a day the LaaS deployment cannot progress.",
        bold_parts=[("Every day this remains unresolved is a day the LaaS deployment cannot progress.", RED_SEVERITY, True)])

    # =================== WHEN THIS STARTED ===================
    add_section_header(doc, "When This Started")
    add_normal_paragraph(doc,
        "The firewall change was made overnight between May 1 evening and May 2 morning (IST). We know this precisely because automated system services on both machines independently confirm the cutoff:")

    add_bullet_paragraph(doc,
        "Last confirmed working: May 1, 5:55 PM IST — ai2's unattended-upgrades successfully reached Ubuntu repos")
    add_bullet_paragraph(doc,
        "First confirmed failing: May 2, 6:53 AM IST — ai2's unattended-upgrades reported \"Distribution outdated\" (couldn't reach repos)")

    add_normal_paragraph(doc,
        "Prior to this, both machines had been downloading and installing packages without any issues for over a week.")

    # =================== AFFECTED MACHINES ===================
    build_affected_machines_table(doc)

    # =================== WHAT WE NEED ===================
    add_section_header(doc, "What We Need")
    add_normal_paragraph(doc,
        "Please whitelist the following domains in the Sophos web filter policy:",
        bold_parts=[("Please whitelist the following domains in the Sophos web filter policy:", DARK_HEADER, True)])

    add_run_text(doc.add_paragraph(), "Essential — Ubuntu Repositories:", bold=True, size=11, color=DARK_HEADER)
    domains_ubuntu = [
        "archive.ubuntu.com (and *.archive.ubuntu.com, in.archive.ubuntu.com)",
        "security.ubuntu.com",
        "ppa.launchpad.net / ppa.launchpadcontent.net",
        "keyserver.ubuntu.com",
    ]
    for d in domains_ubuntu:
        add_bullet_paragraph(doc, d)

    add_run_text(doc.add_paragraph(), "Essential — NVIDIA GPU Software:", bold=True, size=11, color=DARK_HEADER)
    for d in ["developer.download.nvidia.com", "developer.nvidia.com"]:
        add_bullet_paragraph(doc, d)

    add_run_text(doc.add_paragraph(), "Essential — Docker:", bold=True, size=11, color=DARK_HEADER)
    add_bullet_paragraph(doc, "download.docker.com")

    add_run_text(doc.add_paragraph(), "Essential — Development Tools:", bold=True, size=11, color=DARK_HEADER)
    for d in ["github.com / *.github.com / *.githubusercontent.com", "nvidia.github.io"]:
        add_bullet_paragraph(doc, d)

    add_normal_paragraph(doc,
        "Alternative: If domain whitelisting isn't feasible, consider creating a policy exception for source IPs 20.1.1.130 and 20.1.1.132, or placing the AI lab servers in a separate firewall zone.",
        bold_parts=[("Alternative:", DARK_HEADER, True)])

    # =================== HOW TO VERIFY ===================
    add_section_header(doc, "How to Verify the Fix")
    add_normal_paragraph(doc, "After whitelisting, run on any affected machine:")
    add_code_block(doc, [
        "sudo apt update",
        "wget --spider https://developer.download.nvidia.com/compute/cuda/repos/ubuntu2204/x86_64/cuda-keyring_1.1-1_all.deb",
    ])
    add_normal_paragraph(doc,
        "If apt update completes without 403 errors and wget doesn't show a Sophos certificate error, the fix is confirmed.")

    # =================== EVIDENCE ===================
    build_evidence_table(doc)

    # =================== TIMELINE ===================
    build_timeline_ai1(doc)
    build_timeline_ai2(doc)

    # =================== THE CUTOFF ===================
    add_section_header(doc, "The Cutoff")
    add_cutoff_box(doc)

    # =================== CLOSING ===================
    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.LEFT
    closing.paragraph_format.space_before = Pt(12)
    add_run_text(closing, "For questions or a remote walkthrough of the evidence, reach out to the LaaS deployment team.",
                 italic=True, size=10, color="7F8C8D")

    # =================== FOOTER ===================
    add_footer(doc)

    # Save
    doc.save(OUTPUT_DOCX)
    print(f"Report generated successfully: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
