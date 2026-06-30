import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Professional corporate colors: Navy & Charcoal theme
NAVY = RGBColor(26, 54, 93)       # #1A365D - Primary Accent
STEEL = RGBColor(43, 108, 176)    # #2B6CB0 - Secondary Accent
CHARCOAL = RGBColor(45, 55, 72)    # #2D3748 - Section Headers
BODY_TEXT = RGBColor(74, 85, 104)  # #4A5568 - Body Text
WHITE = RGBColor(255, 255, 255)
LIGHT_GRAY = "F7FAFC"             # #F7FAFC - Alternating rows
BORDER_GRAY = "E2E8F0"            # #E2E8F0 - Table borders
LINE_COLOR = "1A365D"             # Navy for accents

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_borders(cell, color_hex="E2E8F0", width="4"):
    """Set cell borders with specified color and width."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), width)
        border.set(qn('w:color'), color_hex)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    """Set cell padding in dxa (twentieths of a point). 140 dxa = 7pt, 180 dxa = 9pt."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        margin = OxmlElement(f'w:{margin_name}')
        margin.set(qn('w:w'), str(val))
        margin.set(qn('w:type'), 'dxa')
        tcMar.append(margin)
    tcPr.append(tcMar)

def style_document(doc):
    """Apply general document styling, margins, and body text formatting."""
    # Base normal style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = BODY_TEXT
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(6)
    
    # 1-inch margins
    for section in doc.sections:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)

def add_cover_page(doc):
    """Create a modern cover page with navy accents and bold professional hierarchy."""
    # Top spacing
    for _ in range(3):
        doc.add_paragraph()
    
    # Elegant Navy accent bar (using table trick)
    accent_table = doc.add_table(rows=1, cols=1)
    accent_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    accent_cell = accent_table.rows[0].cells[0]
    accent_cell.width = Inches(6.5)
    set_cell_shading(accent_cell, LINE_COLOR)
    p = accent_cell.paragraphs[0]
    p.add_run(" ")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    
    doc.add_paragraph()
    
    # Program Code or Tag
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("L&D INTERNSHIP COHORT - 15-DAY PROGRAM")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = STEEL
    p.paragraph_format.space_after = Pt(12)
    
    # Main Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AI-Powered Full-Stack Product Scaffolding")
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = NAVY
    p.paragraph_format.space_after = Pt(6)
    
    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("A 15-Day Vibe-Coding Curriculum for Teams from AI, CSE, IoT, and Cybersecurity")
    run.font.size = Pt(14)
    run.font.italic = True
    run.font.color.rgb = CHARCOAL
    p.paragraph_format.space_after = Pt(24)
    
    # Add vertical spacing
    for _ in range(4):
        doc.add_paragraph()
        
    # Metadata Block Table (Centered)
    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Prepared For:", "B.Tech Cohort (57 Interns, 14 Teams of 4, 1 Team of 5)"),
        ("Program Duration:", "15 Days (Days 1–10: Lectures, Days 11–15: Open Stage Coaching)"),
        ("Session Structure:", "Lecture Days: 1.0 hr Core Concept + 0.5 hr Doubt Resolution"),
        ("Expected Outcome:", "Functional Product POC incorporating UI, DB, and AI Agents"),
        ("Document Version:", "v1.4.0 (Final GPU Allocations - June 2026)")
    ]
    
    for r_idx, (label, val) in enumerate(meta_data):
        row = meta_table.rows[r_idx]
        
        cell_lbl = row.cells[0]
        set_cell_shading(cell_lbl, "F8FAFC")
        set_cell_borders(cell_lbl, "E2E8F0", "2")
        p_lbl = cell_lbl.paragraphs[0]
        r_lbl = p_lbl.add_run(label)
        r_lbl.font.bold = True
        r_lbl.font.size = Pt(9.5)
        r_lbl.font.color.rgb = CHARCOAL
        cell_lbl.width = Inches(2.0)
        set_cell_margins(cell_lbl, 80, 80, 100, 100)
        
        cell_val = row.cells[1]
        set_cell_shading(cell_val, "FFFFFF")
        set_cell_borders(cell_val, "E2E8F0", "2")
        p_val = cell_val.paragraphs[0]
        r_val = p_val.add_run(val)
        r_val.font.size = Pt(9.5)
        r_val.font.color.rgb = BODY_TEXT
        cell_val.width = Inches(4.5)
        set_cell_margins(cell_val, 80, 80, 100, 100)
        
    doc.add_page_break()

def add_section_header(doc, text):
    """Add a professional section header with a navy left-accent indicator."""
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.allow_autofit = False
    
    # Accent indicator bar (Left)
    accent_cell = table.rows[0].cells[0]
    accent_cell.width = Inches(0.12)
    set_cell_shading(accent_cell, LINE_COLOR)
    accent_p = accent_cell.paragraphs[0]
    accent_p.add_run(" ")
    accent_p.paragraph_format.space_before = Pt(0)
    accent_p.paragraph_format.space_after = Pt(0)
    
    # Text cell (Right)
    text_cell = table.rows[0].cells[1]
    text_cell.width = Inches(6.38)
    p = text_cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    
    doc.add_paragraph() # Spacing below heading

def add_curriculum_table(doc, items_data):
    """Generate the structured grid table with merged category visual dividers."""
    table = doc.add_table(rows=1 + len(items_data), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    
    headers = ["Day & Duration", "Curated Topic & Sub-topics", "First Principles & Value Add"]
    col_widths = [1.0, 2.2, 3.3] # Total = 6.5 inches
    
    # Table Header Row
    header_row = table.rows[0]
    for i, title in enumerate(headers):
        cell = header_row.cells[i]
        set_cell_shading(cell, LINE_COLOR)
        set_cell_borders(cell, LINE_COLOR, "4")
        set_cell_margins(cell, 160, 160, 180, 180)
        p = cell.paragraphs[0]
        run = p.add_run(title)
        run.font.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = WHITE
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        cell.width = Inches(col_widths[i])
        
    day_row_count = 0 # To track alternating colors for day rows
    
    # Fill Table Rows
    for r_idx, item in enumerate(items_data):
        row = table.rows[r_idx + 1]
        
        if item["type"] == "category":
            # Visual Divider Row: Merge all columns
            merged_cell = row.cells[0].merge(row.cells[1]).merge(row.cells[2])
            set_cell_shading(merged_cell, "2B6CB0") # Steel Blue background
            set_cell_borders(merged_cell, "2B6CB0", "4")
            set_cell_margins(merged_cell, 140, 140, 180, 180)
            
            p = merged_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(item["title"])
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = WHITE
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            
            merged_cell.width = Inches(6.5)
            
        else:
            # Alternating background color for day rows
            bg_color = "FFFFFF" if day_row_count % 2 == 0 else LIGHT_GRAY
            day_row_count += 1
            
            day_str = item["day"]
            topic_title = item["title"]
            subtopics = item["subtopics"]
            value_add = item["value"]
            
            # 1. Day Column
            cell_day = row.cells[0]
            set_cell_shading(cell_day, bg_color)
            set_cell_borders(cell_day, BORDER_GRAY, "4")
            set_cell_margins(cell_day, 120, 120, 140, 140)
            p_day = cell_day.paragraphs[0]
            run_day = p_day.add_run(day_str)
            run_day.font.bold = True
            run_day.font.size = Pt(10)
            run_day.font.color.rgb = CHARCOAL
            p_day.paragraph_format.space_after = Pt(0)
            cell_day.width = Inches(col_widths[0])
            
            # 2. Topic & Subtopics Column
            cell_topic = row.cells[1]
            set_cell_shading(cell_topic, bg_color)
            set_cell_borders(cell_topic, BORDER_GRAY, "4")
            set_cell_margins(cell_topic, 120, 120, 140, 140)
            
            p_topic = cell_topic.paragraphs[0]
            run_topic = p_topic.add_run(topic_title)
            run_topic.font.bold = True
            run_topic.font.size = Pt(10)
            run_topic.font.color.rgb = NAVY
            p_topic.paragraph_format.space_after = Pt(4)
            
            for sub in subtopics:
                p_sub = cell_topic.add_paragraph()
                run_sub = p_sub.add_run(f"▪  {sub}")
                run_sub.font.size = Pt(8.5)
                run_sub.font.color.rgb = BODY_TEXT
                p_sub.paragraph_format.space_after = Pt(2)
                p_sub.paragraph_format.left_indent = Inches(0.1)
                
            cell_topic.width = Inches(col_widths[1])
            
            # 3. Value Add Column
            cell_val = row.cells[2]
            set_cell_shading(cell_val, bg_color)
            set_cell_borders(cell_val, BORDER_GRAY, "4")
            set_cell_margins(cell_val, 120, 120, 140, 140)
            p_val = cell_val.paragraphs[0]
            run_val = p_val.add_run(value_add)
            run_val.font.size = Pt(9.5)
            run_val.font.color.rgb = BODY_TEXT
            p_val.paragraph_format.space_after = Pt(0)
            cell_val.width = Inches(col_widths[2])
            
    doc.add_paragraph()

def build_document():
    doc = Document()
    style_document(doc)
    
    # 1. Title Page
    add_cover_page(doc)
    
    # 2. Introduction & Philosophy
    add_section_header(doc, "Program Intro: The 'Vibe-Coding' Paradigm")
    
    p = doc.add_paragraph()
    r = p.add_run(
        "Software engineering is undergoing an unprecedented structural transition. The traditional model of "
        "manually writing boilerplate logic, configuration files, and basic HTML markup is rapidly giving way to "
        "'Vibe-Coding'—a collaborative workflow where engineers coordinate, direct, and audit autonomous AI agents "
        "to construct production-ready code. This 15-Day program is designed specifically for a diverse cohort of 57 "
        "students spanning AI, Computer Science (CSE), IoT, and Cybersecurity backgrounds. "
    )
    r.font.size = Pt(10.5)
    
    p2 = doc.add_paragraph()
    r2 = p2.add_run(
        "Organized into 14 collaborative teams of four (and one group of five), students are assigned a multi-dimensional product problem statement "
        "on Day 1 that merges UI, backend APIs, data storage, and AI agents. The daily lectures focus on core concept "
        "delivery and agentic pair-programming demonstrations, structured as 1 hour of teaching followed by 30 minutes of "
        "discussion and doubt resolution. Starting from Day 11, the cohort transitions to an Open Project Stage. Lectures "
        "conclude, and the session block becomes an open diagnostic space where teams receive hands-on mentoring, code-reviews, "
        "and architectural support to compile and launch their final POC."
    )
    r2.font.size = Pt(10.5)
    
    doc.add_paragraph()
    
    # 3. Schedule Table
    add_section_header(doc, "15-Day Daywise Schedule")
    
    curriculum_items = [
        # PART I: COHORT ONBOARDING & SETUP (DAYS 1-2)
        {"type": "category", "title": "PART I: COHORT ONBOARDING & SETUP (DAYS 1-2)"},
        {"type": "day", "day": "Day 1\n(2.0 Hours)", "title": "Program Intro & LaaS Portal Guide", "subtopics": [
            "Cohort introductions & team formations", "Signing up on the LaaS portal", "Launching remote desktops & booking slots", "ZFS home folder limits & hardware checks"
        ], "value": "Establishes immediate compute access. Interns learn how to launch and navigate isolated remote desktop sessions on local RTX 5090 machines."},
        {"type": "day", "day": "Day 2\n(2.0 Hours)", "title": "Vibe-Coding & Antigravity Setup", "subtopics": [
            "Understanding the 'Vibe-Coding' paradigm shift", "Downloading & installing the Antigravity developer environment", "Workspace configurations & environment variables", "Running basic agent commands & background checks"
        ], "value": "Prepares the local development environment, transitioning students from manual syntax writing to AI-driven command execution."},

        # PART II: PRODUCT WORKFLOWS & CONTEXT ENGINEERING (DAYS 3-6)
        {"type": "category", "title": "PART II: PRODUCT WORKFLOWS & CONTEXT ENGINEERING (DAYS 3-6)"},
        {"type": "day", "day": "Day 3\n(1.5 Hours)\n[1h Teach, 0.5h Doubt]", "title": "Product Dev Workflow & Harness (Part 1)", "subtopics": [
            "AI-assisted rapid prototyping models", "Setting up core repository directories", "Creating temporary code scratchpads", "Tracking file edits & compilation logs"
        ], "value": "Establishes the write-run-fix loop, showing how to coordinate AI outputs with real-time execution outputs to audit code changes."},
        {"type": "day", "day": "Day 4\n(1.5 Hours)\n[1h Teach, 0.5h Doubt]", "title": "Product Dev Workflow & Harness (Part 2)", "subtopics": [
            "Harness pipeline scripting & configs", "Running automated code verification steps", "Refactoring code blocks via agent commands", "Managing prompt-compile cycles"
        ], "value": "Enables test-driven agent logic, showing interns how to instruct agents to build, run, and self-correct files based on terminal execution errors."},
        {"type": "day", "day": "Day 5\n(1.5 Hours)\n[1h Teach, 0.5h Doubt]", "title": "Context Engineering: Rules & Skills (Part 1)", "subtopics": [
            "System prompts vs. user input structures", "Designing custom instructions files", "Controlling token boundaries", "Context size optimization techniques"
        ], "value": "Reduces model hallucinations. Interns learn to restrict and focus model contexts so the AI pair programmer outputs accurate results."},
        {"type": "day", "day": "Day 6\n(1.5 Hours)\n[1h Teach, 0.5h Doubt]", "title": "Context Engineering: Rules & Skills (Part 2)", "subtopics": [
            "Writing custom .qoder workspace rules", "Injecting team coding style guides", "Mapping specific workspace tools (skills)", "Enforcing API/schema coding boundaries"
        ], "value": "Ensures generated code aligns with project structures, database requirements, and repository parameters automatically."},

        # PART III: CORE WEB APIs, DATA & LLMs (DAYS 7-10)
        {"type": "category", "title": "PART III: CORE WEB APIs, DATA & LLMs (DAYS 7-10)"},
        {"type": "day", "day": "Day 7\n(1.5 Hours)\n[1h Teach, 0.5h Doubt]", "title": "Generative AI: LLM Fundamentals & APIs", "subtopics": [
            "Attention & Transformer architecture basics", "Temperature parameters & context constraints", "Securing API tokens & handling connection retries", "Enforcing outputs via JSON mode"
        ], "value": "Transitions students from standard chat apps to calling LLM APIs programmatically within Python or Node.js backend handlers."},
        {"type": "day", "day": "Day 8\n(1.5 Hours)\n[1h Teach, 0.5h Doubt]", "title": "Vector Databases, RAG & Prompting", "subtopics": [
            "Text embeddings & mathematical token distance", "Chroma DB & pgvector similarity search", "Implementing text chunking & overlap logic", "RAG query context injection pipelines"
        ], "value": "Enables semantic searches. Interns search local manuals, databases, and code directories semantically without expensive model fine-tuning costs."},
        {"type": "day", "day": "Day 9\n(1.5 Hours)\n[1h Teach, 0.5h Doubt]", "title": "Agentic Systems: Harness, Rules & Skills", "subtopics": [
            "Defining Agent autonomy: Reason-Act (ReAct) loop", "Function schema declaration to LLMs", "Managing states across recursive agent loops", "Setting sandbox tool execution parameters"
        ], "value": "Builds autonomous agents that can evaluate outputs, call external system tools, and iterate to solve complex multi-step tasks."},
        {"type": "day", "day": "Day 10\n(1.5 Hours)\n[1h Teach, 0.5h Doubt]", "title": "Multi-Agent Workflows & Orchestrations", "subtopics": [
            "Visual workflow automation using n8n", "Coordinating event-based triggers (webhooks, schedules)", "Chaining multi-agent inputs/outputs", "Orchestrating agent hierarchies"
        ], "value": "Teaches enterprise workflow integrations, showing how to connect databases, REST APIs, and agent decisions into clean automated pipelines."},

        # PART IV: OPEN PROJECT COACHING STAGE (DAYS 11-15)
        {"type": "category", "title": "PART IV: OPEN PROJECT COACHING STAGE (DAYS 11-15)"},
        {"type": "day", "day": "Days 11–15\n(Open stage daily)", "title": "Ongoing Project Mentoring & Doubt Clearing", "subtopics": [
            "Daily team progress & block checks", "Workspace debugging & environment fixes", "API routing & database connector audits", "Live code-reviews & optimization feedback", "Final project demo pitching rehearsals"
        ], "value": "Consolidates all program knowledge. Providing direct engineering support as teams build their POC asynchronously, culminating in the Day 15 showcase."}
    ]
    
    add_curriculum_table(doc, curriculum_items)
    
    # 4. 14 Unique Projects Section
    add_section_header(doc, "Collaborative Projects: Specifications for 14 Teams")
    
    p = doc.add_paragraph()
    r = p.add_run(
        "Each of the 14 teams is assigned a unique project on Day 1. The specifications below outline how each project "
        "is split across four specialized modules. Resource Allocation: 1 member gets a Blaze instance (4GB VRAM, 8GB RAM) "
        "for the core orchestrator/LLM, while the other 3 members get Spark instances (2GB VRAM, 4GB RAM) for localized UI/Backend/GPU inference:"
    )
    r.font.size = Pt(10.5)
    
    projects = [
        ("Team 1: AI Silk Yarn Grade Classifier & Quality Audit (Textile/Export)", 
         "Grading raw silk thread manually is slow and subjective. This POC standardizes yarn quality grading.\n"
         "• Module 1 (Blaze - Lead): Core Quality Auditor (Runs LLaMA-3.2-3B GGUF to synthesize thread metrics and write grade reports).\n"
         "• Module 2 (Spark): Thread Defect Detector (Runs YOLOv8-Nano on GPU to scan live thread feeds, counting thread slubs and breaks).\n"
         "• Module 3 (Spark): Luster & Sheen Matcher (Light PyTorch CNN on GPU analyzing reflective pixel intensity on thread reels).\n"
         "• Module 4 (Spark): Audit Ledger API & UI (FastAPI + PostgreSQL to store logs, batch details, and serve the inspector UI)."),
         
        ("Team 2: Automated Saree Weaving Pattern Auditor (Textile/Manufacturing)", 
         "Checks Jacquard embroidery for pattern mismatches in premium sarees before packaging for export.\n"
         "• Module 1 (Blaze - Lead): Pattern Compliance Orchestrator (Qwen-2.5-1.5B matching visual anomaly tags to design templates).\n"
         "• Module 2 (Spark): Visual Pattern Segmenter (Runs YOLOv8-Segmentation on GPU to isolate embroidery lines).\n"
         "• Module 3 (Spark): Design Similarity Indexer (Runs Sentence-Transformers on GPU to compare weave design features against templates).\n"
         "• Module 4 (Spark): Defect Ledger API & UI (MongoDB backend logging weaves and defect locations)."),
         
        ("Team 3: Intelligent Local Language Bill Generator (Voice-to-JSON for Local Retailers)", 
         "Local merchants don't know how to type on POS software. They speak queries to generate bills.\n"
         "• Module 1 (Blaze - Lead): Invoice Compiler (Qwen-2.5-1.5B on GPU parsing voice transcriptions and matching items to inventory).\n"
         "• Module 2 (Spark): Dialect Speech-to-Text (Runs Whisper-Tiny on GPU to transcribe local spoken language queries).\n"
         "• Module 3 (Spark): Inventory Database API (FastAPI + PostgreSQL database storing inventory items and sales logs).\n"
         "• Module 4 (Spark): Merchant Web UI (React web app displaying real-time bill additions and generating PDFs)."),
         
        ("Team 4: Export Invoice OCR & Trade Compliance Auditor (Finance/Export)", 
         "Parses handwritten invoices and cross-checks with export laws and trade sanction lists.\n"
         "• Module 1 (Blaze - Lead): Compliance Auditor Agent (Qwen-2.5-3B on GPU verifying parsed invoice items against trade restrictions).\n"
         "• Module 2 (Spark): Document Layout & OCR Parser (Runs EasyOCR on GPU to extract raw text blocks from invoice images).\n"
         "• Module 3 (Spark): Semantic Invoice Encoder (Sentence-Transformers clustering line items to flag pricing anomalies).\n"
         "• Module 4 (Spark): Compliance Audit API & UI (FastAPI + PostgreSQL logging compliance scores and flags)."),
         
        ("Team 5: B2B Agricultural Buyer Matching & Negotiator (Trade/Agriculture)", 
         "Bridges small farmers directly to B2B exporters with an automated negotiation assistant.\n"
         "• Module 1 (Blaze - Lead): Autonomous Negotiator Agent (LLaMA-3.2-3B on GPU negotiating price, quantity, and logistics terms).\n"
         "• Module 2 (Spark): Crop Quality Rating Predictor (Runs a PyTorch CNN on GPU grading crop quality from uploaded photos).\n"
         "• Module 3 (Spark): Market Price RAG (Chroma DB indexing daily state market bulletin reports for pricing context).\n"
         "• Module 4 (Spark): Contract Ledger API & UI (Express + MongoDB storing live negotiations and generating agreements)."),
         
        ("Team 6: AI-Driven Local Logistics Route Optimizer (Logistics/Trade)", 
         "Reduces transportation costs of moving silk and agricultural goods across regional depots.\n"
         "• Module 1 (Blaze - Lead): Central Logistics Coordinator (LLaMA-3.2-1B generating driver schedules and route recommendations).\n"
         "• Module 2 (Spark): GPU Path Solver (PyTorch calculations on GPU computing optimal routes for multiple stops).\n"
         "• Module 3 (Spark): Driver Safety Monitor (YOLOv8-Nano on 2GB VRAM tracking driver fatigue and distraction levels).\n"
         "• Module 4 (Spark): Telemetry API & UI (FastAPI + MongoDB logging GPS coordinates and truck diagnostics via WebSockets)."),
         
        ("Team 7: CCTV Safety & Safe-Workspace compliance monitor (Factory/Security)", 
         "Audits compliance with handloom safety codes (blocking fire exits, wearing safety gear).\n"
         "• Module 1 (Blaze - Lead): Safety Auditor Agent (LLaMA-3.2-1B writing compliance reports and analyzing alarms).\n"
         "• Module 2 (Spark): Fire & Smoke Visual Detector (YOLOv8-Nano on GPU scanning camera feeds for early fire indicators).\n"
         "• Module 3 (Spark): PPE Compliance Inspector (YOLOv8-Nano identifying missing helmet or mask zones at entry points).\n"
         "• Module 4 (Spark): Alarm Gateway API & UI (FastAPI + n8n triggers to sound alarms and send Slack alerts)."),
         
        ("Team 8: Smart Cold-Chain Cocoon Storage Monitor & Alarm (Sericulture/IoT)", 
         "Prevents spoilage of fresh exports by predicting temperature fluctuations during shipping.\n"
         "• Module 1 (Blaze - Lead): Cargo Health Analyzer Agent (Qwen-2.5-1.5B analyzing metrics and flagging decay).\n"
         "• Module 2 (Spark): GPU Time-Series Anomaly Detector (PyTorch LSTM running on GPU to predict temperature spikes).\n"
         "• Module 3 (Spark): Sensor Ingestion API (MQTT broker capturing telemetry streams, logging to MongoDB).\n"
         "• Module 4 (Spark): Cold-Chain Dashboard UI (React frontend tracking live sensor statistics)."),
         
        ("Team 9: Smart Ceramic & Clay Handicraft Defect Scanner (Manufacturing/Art)", 
         "Detects cracks and chips in regional pottery products post-firing to guarantee export quality.\n"
         "• Module 1 (Blaze - Lead): Batch Quality Evaluator (LLaMA-3.2-3B GGUF producing defect summaries and quality ratings).\n"
         "• Module 2 (Spark): Thermal Stress Predictor (PyTorch model on GPU evaluating kiln temperature patterns).\n"
         "• Module 3 (Spark): Structural Visual Scanner (YOLOv8-Nano locating cracks and chips in product photos).\n"
         "• Module 4 (Spark): Inventory Ledger API & UI (MongoDB + Express managing product listings and defect metrics)."),
         
        ("Team 10: Cybersecurity API Traffic Logger & Threat Shield (Cybersecurity)", 
         "Protects AI applications from prompt injections, data leaks, and API token abuse.\n"
         "• Module 1 (Blaze - Lead): Threat Response Planner (LLaMA-3.2-3B analyzing proxy logs and writing protection rules).\n"
         "• Module 2 (Spark): Prompt Injection Shield (Tiny BERT classifier on GPU blocking malicious prompts).\n"
         "• Module 3 (Spark): Traffic Anomaly Classifier (PyTorch autoencoder identifying DDoS and scrapers).\n"
         "• Module 4 (Spark): Proxy Gateway API & UI (FastAPI reverse proxy routing traffic and logging to PostgreSQL)."),
         
        ("Team 11: Regional E-Commerce Auto-Tagger & Metadata Extractor (Retail/Trade)", 
         "Helps local weavers write optimized product descriptions to increase e-commerce search visibility.\n"
         "• Module 1 (Blaze - Lead): SEO Listing Creator (LLaMA-3.2-3B writing optimized descriptions and SEO copy).\n"
         "• Module 2 (Spark): Feature Tag Extractor (ViT image classifier extracting saree colors, textures, and patterns).\n"
         "• Module 3 (Spark): Semantic Search Indexer (Chroma DB indexing listings to enable conversational search).\n"
         "• Module 4 (Spark): Catalog API & UI (Express + MongoDB managing product stock levels)."),
         
        ("Team 12: Autonomous Solar Panel Defect Detector (Energy/IoT)", 
         "Identifies micro-cracks and dust covers on solar arrays to optimize local power generation.\n"
         "• Module 1 (Blaze - Lead): Maintenance Dispatch Agent (LLaMA-3.2-3B generating repair work orders for technicians).\n"
         "• Module 2 (Spark): Infrared Hotspot Detector (PyTorch CNN flagging thermal anomalies in drone imagery).\n"
         "• Module 3 (Spark): Telemetry Ingestion (MQTT broker capturing telemetry metrics, logging to PostgreSQL).\n"
         "• Module 4 (Spark): Operations API & UI (FastAPI + PostgreSQL tracking panel health logs)."),
         
        ("Team 13: Local AI Technical Document Runbook Assistant (Business/Support)", 
         "Helps technicians query machine user manuals using natural language speech.\n"
         "• Module 1 (Blaze - Lead): Runbook Assistant (LLaMA-3.2-3B parsing manual data to answer queries).\n"
         "• Module 2 (Spark): RAG Document Parser (Chroma DB indexing text chunks and diagrams from PDF uploads).\n"
         "• Module 3 (Spark): Voice Query Transcriber (Whisper-Base transcribing spoken questions on the shop floor).\n"
         "• Module 4 (Spark): Search History API & UI (FastAPI + PostgreSQL tracking queries and document paths)."),
         
        ("Team 14: Handloom Weaving Tension Auditor & Quality Log (Manufacturing/IoT)", 
         "Detects tension anomalies and missing threads on handlooms to protect fabric quality.\n"
         "• Module 1 (Blaze - Lead): Weave Auditor (LLaMA-3.2-3B generating batch rating summaries and weaver tips).\n"
         "• Module 2 (Spark): Tension Anomaly Scanner (PyTorch autoencoder on GPU analyzing tension sensor logs).\n"
         "• Module 3 (Spark): Visual Weave Auditor (YOLOv8-Nano processing loom cam frames, flagging missing warps).\n"
         "• Module 4 (Spark): Weave Analytics API & UI (FastAPI + PostgreSQL logging thread logs and alerts).")
    ]
    
    for title, desc in projects:
        p_proj = doc.add_paragraph()
        run_title = p_proj.add_run(f"•  {title}\n")
        run_title.font.bold = True
        run_title.font.size = Pt(11)
        run_title.font.color.rgb = NAVY
        
        run_desc = p_proj.add_run(desc)
        run_desc.font.size = Pt(9.5)
        run_desc.font.color.rgb = BODY_TEXT
        p_proj.paragraph_format.left_indent = Inches(0.2)
        p_proj.paragraph_format.space_after = Pt(8)
        
    doc.add_paragraph()
    
    # Save document
    output_path = r"c:\Users\Punith\LaaS\CS_15_Day_VibeCoding_Curriculum.docx"
    try:
        doc.save(output_path)
        print(f"SUCCESS: Curriculum document generated at: {output_path}")
    except Exception as e:
        print(f"ERROR: Error saving document: {e}")

if __name__ == "__main__":
    build_document()
