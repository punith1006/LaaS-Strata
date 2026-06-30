import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Professional corporate colors: Navy & Charcoal theme
NAVY = RGBColor(26, 54, 93)       # #1A365D - Primary Accent
STEEL = RGBColor(43, 108, 176)    # #2B6CB0 - Secondary Accent
CHARCOAL = RGBColor(45, 55, 72)    # #2D3748 - Section Headers
BODY_TEXT = RGBColor(74, 85, 104)  # #4A5568 - Body Text
WHITE = RGBColor(255, 255, 255)
LIGHT_GRAY = "F7FAFC"             # #F7FAFC - Alternating rows
BORDER_GRAY = "E2E8F0"            # #E2E8F0 - Table borders
LINE_COLOR = "1A365D"             # Navy for accents

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_borders(cell, color_hex="E2E8F0", width="4"):
    """Set cell borders with specified color and width."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), width)
        border.set(qn('w:color'), color_hex)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    """Set cell padding in dxa. 140 dxa = 7pt, 180 dxa = 9pt."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        margin = OxmlElement(f'w:{margin_name}')
        margin.set(qn('w:w'), str(val))
        margin.set(qn('w:type'), 'dxa')
        tcMar.append(margin)
    tcPr.append(tcMar)

def style_document(doc):
    """Apply general document styling, margins, and body text formatting."""
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = BODY_TEXT
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(6)
    
    # 1-inch margins
    for section in doc.sections:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)

def add_cover_page(doc):
    """Create a modern cover page with navy accents and bold professional hierarchy."""
    for _ in range(3):
        doc.add_paragraph()
    
    # Elegant Navy accent bar
    accent_table = doc.add_table(rows=1, cols=1)
    accent_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    accent_cell = accent_table.rows[0].cells[0]
    accent_cell.width = Inches(6.5)
    set_cell_shading(accent_cell, LINE_COLOR)
    p = accent_cell.paragraphs[0]
    p.add_run(" ")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    
    doc.add_paragraph()
    
    # Program Tag
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("L&D INTERNSHIP COHORT - 15-DAY PROGRAM")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = STEEL
    p.paragraph_format.space_after = Pt(12)
    
    # Main Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AI-Powered Full-Stack Product Scaffolding")
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = NAVY
    p.paragraph_format.space_after = Pt(6)
    
    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("A 15-Day Vibe-Coding Curriculum and 14 GPU-Enabled Showcase Projects")
    run.font.size = Pt(14)
    run.font.italic = True
    run.font.color.rgb = CHARCOAL
    p.paragraph_format.space_after = Pt(24)
    
    for _ in range(4):
        doc.add_paragraph()
        
    # Metadata Block Table
    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Prepared For:", "B.Tech Cohort (57 Interns, 14 Teams: 13x4, 1x5)"),
        ("Program Duration:", "15 Days (Days 1–10: Lectures, Days 11–15: Open Coaching Stage)"),
        ("Session Structure:", "Lecture Days: 1.0 hr Core Concept + 0.5 hr Doubt Resolution"),
        ("Expected Outcome:", "Functional Product POC incorporating UI, DB, and GPU Models"),
        ("Document Version:", "v1.5.0 (Final project_ideas.md Alignments - June 2026)")
    ]
    
    for r_idx, (label, val) in enumerate(meta_data):
        row = meta_table.rows[r_idx]
        
        cell_lbl = row.cells[0]
        set_cell_shading(cell_lbl, "F8FAFC")
        set_cell_borders(cell_lbl, "E2E8F0", "2")
        p_lbl = cell_lbl.paragraphs[0]
        r_lbl = p_lbl.add_run(label)
        r_lbl.font.bold = True
        r_lbl.font.size = Pt(9.5)
        r_lbl.font.color.rgb = CHARCOAL
        cell_lbl.width = Inches(2.0)
        set_cell_margins(cell_lbl, 80, 80, 100, 100)
        
        cell_val = row.cells[1]
        set_cell_shading(cell_val, "FFFFFF")
        set_cell_borders(cell_val, "E2E8F0", "2")
        p_val = cell_val.paragraphs[0]
        r_val = p_val.add_run(val)
        r_val.font.size = Pt(9.5)
        r_val.font.color.rgb = BODY_TEXT
        cell_val.width = Inches(4.5)
        set_cell_margins(cell_val, 80, 80, 100, 100)
        
    doc.add_page_break()

def add_section_header(doc, text):
    """Add a professional section header with a navy left-accent indicator."""
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.allow_autofit = False
    
    # Accent indicator bar (Left)
    accent_cell = table.rows[0].cells[0]
    accent_cell.width = Inches(0.12)
    set_cell_shading(accent_cell, LINE_COLOR)
    accent_p = accent_cell.paragraphs[0]
    accent_p.add_run(" ")
    accent_p.paragraph_format.space_before = Pt(0)
    accent_p.paragraph_format.space_after = Pt(0)
    
    # Text cell (Right)
    text_cell = table.rows[0].cells[1]
    text_cell.width = Inches(6.38)
    p = text_cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    
    doc.add_paragraph() # Spacing below heading

def add_curriculum_table(doc, items_data):
    """Generate the structured grid table with merged category visual dividers."""
    table = doc.add_table(rows=1 + len(items_data), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    
    headers = ["Day & Duration", "Curated Topic & Sub-topics", "First Principles & Value Add"]
    col_widths = [1.0, 2.2, 3.3] # Total = 6.5 inches
    
    # Table Header Row
    header_row = table.rows[0]
    for i, title in enumerate(headers):
        cell = header_row.cells[i]
        set_cell_shading(cell, LINE_COLOR)
        set_cell_borders(cell, LINE_COLOR, "4")
        set_cell_margins(cell, 160, 160, 180, 180)
        p = cell.paragraphs[0]
        run = p.add_run(title)
        run.font.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = WHITE
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        cell.width = Inches(col_widths[i])
        
    day_row_count = 0
    
    # Fill Table Rows
    for r_idx, item in enumerate(items_data):
        row = table.rows[r_idx + 1]
        
        if item["type"] == "category":
            merged_cell = row.cells[0].merge(row.cells[1]).merge(row.cells[2])
            set_cell_shading(merged_cell, "2B6CB0")
            set_cell_borders(merged_cell, "2B6CB0", "4")
            set_cell_margins(merged_cell, 140, 140, 180, 180)
            
            p = merged_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(item["title"])
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = WHITE
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            merged_cell.width = Inches(6.5)
            
        else:
            bg_color = "FFFFFF" if day_row_count % 2 == 0 else LIGHT_GRAY
            day_row_count += 1
            
            day_str = item["day"]
            topic_title = item["title"]
            subtopics = item["subtopics"]
            value_add = item["value"]
            
            # 1. Day Column
            cell_day = row.cells[0]
            set_cell_shading(cell_day, bg_color)
            set_cell_borders(cell_day, BORDER_GRAY, "4")
            set_cell_margins(cell_day, 120, 120, 140, 140)
            p_day = cell_day.paragraphs[0]
            run_day = p_day.add_run(day_str)
            run_day.font.bold = True
            run_day.font.size = Pt(10)
            run_day.font.color.rgb = CHARCOAL
            p_day.paragraph_format.space_after = Pt(0)
            cell_day.width = Inches(col_widths[0])
            
            # 2. Topic Column
            cell_topic = row.cells[1]
            set_cell_shading(cell_topic, bg_color)
            set_cell_borders(cell_topic, BORDER_GRAY, "4")
            set_cell_margins(cell_topic, 120, 120, 140, 140)
            
            p_topic = cell_topic.paragraphs[0]
            run_topic = p_topic.add_run(topic_title)
            run_topic.font.bold = True
            run_topic.font.size = Pt(10)
            run_topic.font.color.rgb = NAVY
            p_topic.paragraph_format.space_after = Pt(4)
            
            for sub in subtopics:
                p_sub = cell_topic.add_paragraph()
                run_sub = p_sub.add_run(f"▪  {sub}")
                run_sub.font.size = Pt(8.5)
                run_sub.font.color.rgb = BODY_TEXT
                p_sub.paragraph_format.space_after = Pt(2)
                p_sub.paragraph_format.left_indent = Inches(0.1)
                
            cell_topic.width = Inches(col_widths[1])
            
            # 3. Value Column
            cell_val = row.cells[2]
            set_cell_shading(cell_val, bg_color)
            set_cell_borders(cell_val, BORDER_GRAY, "4")
            set_cell_margins(cell_val, 120, 120, 140, 140)
            p_val = cell_val.paragraphs[0]
            run_val = p_val.add_run(value_add)
            run_val.font.size = Pt(9.5)
            run_val.font.color.rgb = BODY_TEXT
            p_val.paragraph_format.space_after = Pt(0)
            cell_val.width = Inches(col_widths[2])
            
    doc.add_paragraph()

def add_project_modules_table(doc, modules):
    """Generate a table showing the modular breakdown of the project."""
    table = doc.add_table(rows=1 + len(modules), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    
    headers = ["Module", "Member", "GPU Model (VRAM)", "What They Build / Responsibility"]
    col_widths = [1.2, 0.8, 1.8, 2.7] # Total = 6.5 inches
    
    # Style Header
    header_row = table.rows[0]
    for i, title in enumerate(headers):
        cell = header_row.cells[i]
        set_cell_shading(cell, "2B6CB0") # Steel Blue
        set_cell_borders(cell, "2B6CB0", "4")
        set_cell_margins(cell, 100, 100, 120, 120)
        p = cell.paragraphs[0]
        run = p.add_run(title)
        run.font.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = WHITE
        cell.width = Inches(col_widths[i])
        
    for r_idx, mod in enumerate(modules):
        row = table.rows[r_idx + 1]
        bg_color = "FFFFFF" if r_idx % 2 == 0 else LIGHT_GRAY
        
        for c_idx, val in enumerate(mod):
            cell = row.cells[c_idx]
            set_cell_shading(cell, bg_color)
            set_cell_borders(cell, BORDER_GRAY, "2")
            set_cell_margins(cell, 80, 80, 100, 100)
            p = cell.paragraphs[0]
            
            run = p.add_run(val)
            run.font.size = Pt(8.5)
            run.font.color.rgb = BODY_TEXT
            
            if c_idx in [0, 1]:
                run.font.bold = True
                run.font.color.rgb = CHARCOAL
            elif c_idx == 2:
                run.font.bold = True
                run.font.color.rgb = NAVY
                
            cell.width = Inches(col_widths[c_idx])
            
    doc.add_paragraph()

def build_document():
    doc = Document()
    style_document(doc)
    
    # 1. Title Page
    add_cover_page(doc)
    
    # 2. Introduction & Philosophy
    add_section_header(doc, "Program Intro: The 'Vibe-Coding' Paradigm")
    
    p = doc.add_paragraph()
    r = p.add_run(
        "Software engineering is undergoing an architectural transition. The traditional model of "
        "manually writing boilerplate code is rapidly giving way to 'Vibe-Coding'—a collaborative workflow where engineers "
        "direct and audit autonomous AI agents to construct production-ready code. This 15-Day program is designed for a "
        "diverse cohort of 57 B.Tech students spanning AI, Computer Science (CSE), IoT, and Cybersecurity backgrounds. "
    )
    r.font.size = Pt(10.5)
    
    p2 = doc.add_paragraph()
    r2 = p2.add_run(
        "Organized into 14 collaborative teams (13 teams of four, and one team of five), students are assigned a GPU-enabled product problem statement "
        "on Day 1 that merges UI, backend APIs, data storage, and local machine learning models. Daily lectures focus on core concepts "
        "and agentic workflows, structured as 1 hour of teaching followed by 30 minutes of discussion. From Day 11 to Day 15, "
        "the cohort transitions to an Open Project Stage. Lectures conclude, and the session block becomes an open diagnostic space "
        "where teams receive code reviews and deployment support to launch their final POC on the LaaS platform."
    )
    r2.font.size = Pt(10.5)
    
    doc.add_paragraph()
    
    # 3. Schedule Table
    add_section_header(doc, "15-Day Daywise Schedule")
    
    curriculum_items = [
        {"type": "category", "title": "PART I: COHORT ONBOARDING & SETUP (DAYS 1-2)"},
        {"type": "day", "day": "Day 1\n(2.0 Hours)", "title": "Program Intro & LaaS Portal Guide", "subtopics": [
            "Cohort introductions & team formations", "Signing up on the LaaS portal", "Launching remote desktops & booking slots", "ZFS home folder limits & hardware checks"
        ], "value": "Establishes immediate compute access. Interns learn how to launch and navigate isolated remote desktop sessions on local RTX 5090 machines."},
        {"type": "day", "day": "Day 2\n(2.0 Hours)", "title": "Vibe-Coding & Antigravity Setup", "subtopics": [
            "Understanding the 'Vibe-Coding' paradigm shift", "Downloading & installing the Antigravity developer environment", "Workspace configurations & environment variables", "Running basic agent commands & background checks"
        ], "value": "Prepares the local development environment, transitioning students from manual syntax writing to AI-driven command execution."},

        {"type": "category", "title": "PART II: PRODUCT WORKFLOWS & CONTEXT ENGINEERING (DAYS 3-6)"},
        {"type": "day", "day": "Day 3\n(1.5 Hours)\n[1h Teach, 0.5h Doubt]", "title": "Product Dev Workflow & Harness (Part 1)", "subtopics": [
            "AI-assisted rapid prototyping models", "Setting up core repository directories", "Creating temporary code scratchpads", "Tracking file edits & compilation logs"
        ], "value": "Establishes the write-run-fix loop, showing how to coordinate AI outputs with real-time execution outputs to audit code changes."},
        {"type": "day", "day": "Day 4\n(1.5 Hours)\n[1h Teach, 0.5h Doubt]", "title": "Product Dev Workflow & Harness (Part 2)", "subtopics": [
            "Harness pipeline scripting & configs", "Running automated code verification steps", "Refactoring code blocks via agent commands", "Managing prompt-compile cycles"
        ], "value": "Enables test-driven agent logic, showing interns how to instruct agents to build, run, and self-correct files based on terminal execution errors."},
        {"type": "day", "day": "Day 5\n(1.5 Hours)\n[1h Teach, 0.5h Doubt]", "title": "Context Engineering: Rules & Skills (Part 1)", "subtopics": [
            "System prompts vs. user input structures", "Designing custom instructions files", "Controlling token boundaries", "Context size optimization techniques"
        ], "value": "Reduces model hallucinations. Interns learn to restrict and focus model contexts so the AI pair programmer outputs accurate results."},
        {"type": "day", "day": "Day 6\n(1.5 Hours)\n[1h Teach, 0.5h Doubt]", "title": "Context Engineering: Rules & Skills (Part 2)", "subtopics": [
            "Writing custom .qoder workspace rules", "Injecting team coding style guides", "Mapping specific workspace tools (skills)", "Enforcing API/schema coding boundaries"
        ], "value": "Ensures generated code aligns with project structures, database requirements, and repository parameters automatically."},

        {"type": "category", "title": "PART III: CORE WEB APIs, DATA & LLMs (DAYS 7-10)"},
        {"type": "day", "day": "Day 7\n(1.5 Hours)\n[1h Teach, 0.5h Doubt]", "title": "Generative AI: LLM Fundamentals & APIs", "subtopics": [
            "Attention & Transformer architecture basics", "Temperature parameters & context constraints", "Securing API tokens & handling connection retries", "Enforcing outputs via JSON mode"
        ], "value": "Transitions students from standard chat apps to calling LLM APIs programmatically within Python or Node.js backend handlers."},
        {"type": "day", "day": "Day 8\n(1.5 Hours)\n[1h Teach, 0.5h Doubt]", "title": "Vector Databases, RAG & Prompting", "subtopics": [
            "Text embeddings & mathematical token distance", "Chroma DB & pgvector similarity search", "Implementing text chunking & overlap logic", "RAG query context injection pipelines"
        ], "value": "Enables semantic searches. Interns search local manuals, databases, and code directories semantically without expensive model fine-tuning costs."},
        {"type": "day", "day": "Day 9\n(1.5 Hours)\n[1h Teach, 0.5h Doubt]", "title": "Agentic Systems: Harness, Rules & Skills", "subtopics": [
            "Defining Agent autonomy: Reason-Act (ReAct) loop", "Function schema declaration to LLMs", "Managing states across recursive agent loops", "Setting sandbox tool execution parameters"
        ], "value": "Builds autonomous agents that can evaluate outputs, call external system tools, and iterate to solve complex multi-step tasks."},
        {"type": "day", "day": "Day 10\n(1.5 Hours)\n[1h Teach, 0.5h Doubt]", "title": "Multi-Agent Workflows & Orchestrations", "subtopics": [
            "Visual workflow automation using n8n", "Coordinating event-based triggers (webhooks, schedules)", "Chaining multi-agent inputs/outputs", "Orchestrating agent hierarchies"
        ], "value": "Teaches enterprise workflow integrations, showing how to connect databases, REST APIs, and agent decisions into clean automated pipelines."},

        {"type": "category", "title": "PART IV: OPEN PROJECT COACHING STAGE (DAYS 11-15)"},
        {"type": "day", "day": "Days 11–15\n(Open stage daily)", "title": "Ongoing Project Mentoring & Doubt Clearing", "subtopics": [
            "Daily team progress & block checks", "Workspace debugging & environment fixes", "API routing & database connector audits", "Live code-reviews & optimization feedback", "Final project demo pitching rehearsals"
        ], "value": "Consolidates all program knowledge, providing direct engineering support as teams build their POC asynchronously, culminating in the Day 15 showcase."}
    ]
    
    add_curriculum_table(doc, curriculum_items)
    
    # 4. Resource Constraints and Data Input Methods
    add_section_header(doc, "Resource Constraints & Data Ingestion Scheme")
    
    p = doc.add_paragraph()
    r = p.add_run(
        "Each team is allocated a strict compute layout. One team member receives a Blaze instance (4GB VRAM) to host the core integrations, "
        "while the remaining members receive Spark instances (2GB VRAM) to develop their specific modules. Since containers run on central GPU "
        "hosts, there is no support for live webcam/mic feeds. All media feeds are simulated via pre-recorded files or public URLs:"
    )
    r.font.size = Pt(10.5)
    
    # Data Input Table
    input_table = doc.add_table(rows=5, cols=2)
    input_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    input_table.autofit = False
    input_table.allow_autofit = False
    
    input_headers = ["Ingestion Method", "Operational Workflow"]
    for i, title in enumerate(input_headers):
        cell = input_table.rows[0].cells[i]
        set_cell_shading(cell, LINE_COLOR)
        set_cell_borders(cell, LINE_COLOR, "4")
        p_cell = cell.paragraphs[0]
        run_h = p_cell.add_run(title)
        run_h.font.bold = True
        run_h.font.size = Pt(10)
        run_h.font.color.rgb = WHITE
        cell.width = Inches(3.25)
        
    inputs = [
        ("File Upload", "Intern records media on smartphone -> uploads via LaaS file transfer tool -> stored in /home/ubuntu ZFS volume."),
        ("URL Ingestion", "Containers download public links using yt-dlp, wget, or index RTSP streams from public URL inputs."),
        ("Text Input", "Students write structured descriptions or prompt text directly into apps running inside the Selkies desktop."),
        ("Pre-loaded Datasets", "Standard public datasets (such as COCO or custom image banks) are indexed directly on the ZFS storage dataset.")
    ]
    
    for r_idx, (method, workflow) in enumerate(inputs):
        row = input_table.rows[r_idx + 1]
        bg = "FFFFFF" if r_idx % 2 == 0 else LIGHT_GRAY
        
        cell_m = row.cells[0]
        set_cell_shading(cell_m, bg)
        set_cell_borders(cell_m, BORDER_GRAY, "2")
        set_cell_margins(cell_m, 80, 80, 100, 100)
        p_m = cell_m.paragraphs[0]
        run_m = p_m.add_run(method)
        run_m.font.bold = True
        run_m.font.size = Pt(9.5)
        run_m.font.color.rgb = CHARCOAL
        cell_m.width = Inches(2.0)
        
        cell_w = row.cells[1]
        set_cell_shading(cell_w, bg)
        set_cell_borders(cell_w, BORDER_GRAY, "2")
        set_cell_margins(cell_w, 80, 80, 100, 100)
        p_w = cell_w.paragraphs[0]
        run_w = p_w.add_run(workflow)
        run_w.font.size = Pt(9)
        run_w.font.color.rgb = BODY_TEXT
        cell_w.width = Inches(4.5)
        
    doc.add_paragraph()
    
    # 5. The 14 Projects Section
    add_section_header(doc, "Collaborative Projects: Specifications for 14 Teams")
    
    projects_data = [
        {
            "name": "1. PixelRevive — AI Photo Restoration Studio",
            "diff": "Standard",
            "one_liner": "Upload a damaged, old, black-and-white family photo -> get back a restored, colorized, 4K version.",
            "demo": "Show a torn 1970s family photo -> watch AI fix scratches, add color, upscale to HD. Before/after reveal.",
            "input": "Student uploads photo via LaaS web UI",
            "modules": [
                ["Scratch Removal", "Member 1", "U-Net-small (~800 MB)", "Inpainting model to fill scratches, tears, water damage"],
                ["Colorization", "Member 2", "DeOldify (~100 MB)", "Auto-colorize B&W photos with plausible colors"],
                ["Super-Resolution", "Member 3", "Real-ESRGAN x4 (~500 MB)", "Upscale restored photo to 4K quality"],
                ["Web Studio + AI Agent", "Member 4", "ChromaDB + Embeddings (~300 MB)", "Upload UI, side-by-side comparison, batch processing, AI describes changes"]
            ]
        },
        {
            "name": "2. SnapChef — Fridge-to-Recipe AI",
            "diff": "Standard",
            "one_liner": "Snap a photo of your fridge -> AI identifies every ingredient -> generates a personalized recipe.",
            "demo": "Open a fridge, take a photo, upload it -> watch ingredients get detected with bounding boxes -> full recipe appears.",
            "input": "Student uploads fridge/pantry photo",
            "modules": [
                ["Ingredient Detection", "Member 1", "YOLOv8 Small (~1.5 GB)", "Fine-tune to detect 50+ common food items"],
                ["Recipe Generation", "Member 2", "Qwen2-1.5B-Q4 on Blaze (~1.1 GB)", "Generate recipes from detected ingredients"],
                ["Nutrition + Preferences", "Member 3", "ChromaDB + Embeddings (~300 MB)", "Dietary filters, allergy checks, nutrition scoring"],
                ["Cook-Along UI", "Member 4", "Frontend UI (No GPU model)", "Step-by-step interface, timer, ingredient checklist"]
            ]
        },
        {
            "name": "3. SoundForge — AI Music & Beat Generator",
            "diff": "Standard",
            "one_liner": "Describe a mood in text -> AI generates original music. Layer tracks, add beats, export.",
            "demo": "Type 'chill lo-fi with Tamil folk influence' -> music starts playing. Layer a beat. Export the mix.",
            "input": "Text description (typed in desktop)",
            "modules": [
                ["Music Generation", "Member 1", "MusicGen Small on Blaze (~2.5 GB)", "Text-to-music, genre/mood conditioning"],
                ["Beat Detection & Mixing", "Member 2", "Audio DSP + PyTorch (~500 MB)", "BPM detection, beat alignment, audio mixing"],
                ["Audio Effects", "Member 3", "Neural audio effects (~400 MB)", "Reverb, EQ, style transfer on audio"],
                ["Studio UI", "Member 4", "Frontend (Web Audio API)", "Multi-track editor, waveform viz, export to MP3"]
            ]
        },
        {
            "name": "4. FormCheck — AI Exercise Form Analyzer",
            "diff": "Standard",
            "one_liner": "Upload a video of yourself exercising -> AI tracks your pose -> generates a form report with corrections.",
            "demo": "Upload a squat video -> watch skeleton overlay -> 'knee angle: 72 (should be 90)' at timestamps -> AI tips.",
            "input": "Student uploads exercise video (recorded on phone)",
            "modules": [
                ["Video Processing", "Member 1", "OpenCV frame extraction", "Extract frames, manage video pipeline"],
                ["Pose Estimation", "Member 2", "MediaPipe Pose (~300 MB)", "Real-time 33-point body tracking on GPU"],
                ["Form Analysis", "Member 3", "Custom angle classifier (~200 MB)", "Joint angles, correct-form comparison, rep counting"],
                ["Report UI", "Member 4", "Frontend Canvas overlay", "Video playback + skeleton overlay, angle graphs, AI tips"]
            ]
        },
        {
            "name": "5. HawkEye — AI Video Intelligence Engine",
            "diff": "Standard",
            "one_liner": "Paste a YouTube URL or upload video -> AI detects objects, counts people, finds anomalies -> searchable event timeline.",
            "demo": "Paste a traffic camera YouTube URL -> container pulls stream -> real-time detection overlay -> searchable timeline.",
            "input": "Paste URL (yt-dlp) or upload .mp4",
            "modules": [
                ["Video Ingestion", "Member 1", "yt-dlp + OpenCV", "URL download, frame extraction, video management"],
                ["Frame Detection", "Member 2", "YOLOv8 Nano (~1 GB)", "Per-frame object/person/vehicle detection on GPU"],
                ["Event Intelligence", "Member 3", "ChromaDB + Embeddings (~300 MB)", "Anomaly detection, object tracking, event classification"],
                ["Timeline Dashboard", "Member 4", "Frontend Timeline", "Searchable event log, video scrubber with annotations, export"]
            ]
        },
        {
            "name": "6. GameBrain — AI That Learns to Play Classic Games",
            "diff": "Challenging",
            "one_liner": "Pick a game (Snake, Pong, Flappy Bird) -> watch AI train from zero to superhuman in real-time.",
            "demo": "Start training on Snake -> AI crashes -> 2 minutes later it plays perfectly. AI vs. Human challenge mode.",
            "input": "Internal (no external input - game runs inside the container)",
            "modules": [
                ["RL Training Engine", "Member 1", "DQN/PPO on GPU (~800 MB)", "Train agents with Deep Q-Learning on GPU"],
                ["Game Environments", "Member 2", "PyTorch + Gymnasium (~500 MB)", "GPU-accelerated Snake, Pong, Flappy Bird"],
                ["Visualization", "Member 3", "State analysis embeddings (~300 MB)", "Neural net viz, learning curves, state-action heatmaps"],
                ["Arena UI", "Member 4", "Frontend Canvas/WebGL", "Game viewer, training controls, leaderboard, AI vs. Human"]
            ]
        },
        {
            "name": "7. ArtForge — AI Art & Style Studio",
            "diff": "Standard",
            "one_liner": "Upload any photo -> apply 10+ artistic styles (Van Gogh, Picasso, anime) -> adjust intensity -> export.",
            "demo": "Upload a selfie -> click through 10 styles -> renders in <2 seconds -> drag slider -> export.",
            "input": "Student uploads photo",
            "modules": [
                ["Style Transfer Engine", "Member 1", "Fast Neural Style (~1 GB)", "10+ pre-trained style models on GPU"],
                ["Batch Processing", "Member 2", "GPU pipeline (~500 MB)", "Multi-style simultaneously, adjustable intensity"],
                ["AI Art Curation", "Member 3", "ChromaDB + Embeddings (~300 MB)", "'Find similar styles', gallery, style recommendation"],
                ["Studio UI", "Member 4", "Frontend Gallery", "Drag-drop upload, style gallery, intensity slider, comparison"]
            ]
        },
        {
            "name": "8. EchoScribe — AI Video Transcription & Study Platform",
            "diff": "Standard",
            "one_liner": "Upload a lecture video or paste YouTube URL -> AI transcribes, generates chapters, flashcards, and quizzes.",
            "demo": "Paste lecture URL -> real-time transcription -> chapters auto-generate -> click for AI flashcards.",
            "input": "Upload video or paste YouTube URL (yt-dlp)",
            "modules": [
                ["Transcription", "Member 1", "Whisper Small (~2 GB)", "Speech-to-text with timestamps + speaker diarization"],
                ["Intelligence Layer", "Member 2", "ChromaDB + Embeddings (~300 MB)", "Topic segmentation, chapter detection, key moments"],
                ["Study Tools", "Member 3", "Qwen2-1.5B-Q4 on Blaze (~1.1 GB)", "Flashcards, summaries, quiz questions from transcript"],
                ["Video Player UI", "Member 4", "Frontend Video Player", "Synced video + transcript, chapter sidebar, flashcard viewer"]
            ]
        },
        {
            "name": "9. VoxLingua — Voice Translator",
            "diff": "Standard",
            "one_liner": "Upload audio in Tamil -> AI transcribes, translates to English, and speaks it back.",
            "demo": "Upload a Tamil voice recording -> see transcription appear -> see English translation -> hear it spoken aloud.",
            "input": "Student uploads audio file (.wav/.mp3, recorded on phone)",
            "modules": [
                ["Speech-to-Text", "Member 1", "Whisper Small (~2 GB)", "Tamil + English + Hindi transcription on GPU"],
                ["Translation Engine", "Member 2", "Qwen2-1.5B-Q4 on Blaze (~1.1 GB)", "Context-aware translation"],
                ["Text-to-Speech", "Member 3", "Piper TTS (GPU, ~500 MB)", "Natural voice output in target language"],
                ["Conversation UI", "Member 4", "Frontend Audio Player", "Audio player, transcript view, translation panel, playback"]
            ]
        },
        {
            "name": "10. MemeForge — Context-Aware AI Meme Generator",
            "diff": "Standard",
            "one_liner": "Upload any image -> AI understands the context -> generates genuinely funny captions.",
            "demo": "Upload a photo of the professor -> AI generates a meme about assignments. The room erupts.",
            "input": "Student uploads image or picks from trending templates",
            "modules": [
                ["Image Understanding", "Member 1", "CLIP ViT-B/32 (~400 MB) + YOLOv8 Nano (~1 GB)", "Detect objects, scene context in image"],
                ["Caption Generation", "Member 2", "Qwen2-1.5B-Q4 on Blaze (~1.1 GB)", "Contextually relevant, funny captions"],
                ["Template Engine", "Member 3", "ChromaDB + Embeddings (~300 MB)", "Match image to meme templates, text placement"],
                ["Meme Studio UI", "Member 4", "Frontend Canvas Workspace", "Upload/template picker, caption editor, text drag-drop"]
            ]
        },
        {
            "name": "11. Scene3D — Single-Image 3D Depth Scanner",
            "diff": "Challenging",
            "one_liner": "Upload any photo -> AI estimates depth for every pixel -> generates a 3D point cloud you can rotate and fly through.",
            "demo": "Upload a landscape photo -> watch it transform into a rotating 3D scene. Fly through the depth.",
            "input": "Student uploads photo",
            "modules": [
                ["Depth Estimation", "Member 1", "Depth Anything V2 Small (~500 MB)", "Per-pixel depth prediction from single image"],
                ["3D Point Cloud", "Member 2", "PyTorch3D + GPU (~500 MB)", "Depth map -> 3D point cloud, mesh generation"],
                ["Scene Enhancement", "Member 3", "Real-ESRGAN (~500 MB)", "Upscale input for better depth, semantic understanding"],
                ["3D Viewer UI", "Member 4", "Frontend (Three.js)", "Interactive 3D viewer, orbit/fly controls, depth slider"]
            ]
        },
        {
            "name": "12. StoryQuest — AI Interactive Adventure Game",
            "diff": "Challenging",
            "one_liner": "An AI Dungeon Master that creates interactive adventures with voice narration AND generates an image for every scene.",
            "demo": "'You are in a dark forest. Which way?' -> choose left -> AI narrates -> generates scene image.",
            "input": "Text choices (typed in desktop)",
            "modules": [
                ["Story Engine", "Member 1", "Qwen2-1.5B-Q4 on Blaze (~1.1 GB)", "Dynamic story, branching narratives, character consistency"],
                ["Scene Illustration", "Member 2", "SD 1.5 optimized on Blaze (~3.5 GB)", "Generate image for each scene description"],
                ["Voice Narration", "Member 3", "Whisper Base (~1 GB) + Piper TTS", "Voice input for choices + AI narration"],
                ["Game UI", "Member 4", "Frontend Gaming Board", "Text adventure, scene illustrations, choice buttons, story map"]
            ]
        },
        {
            "name": "13. SignReader — Video Sign Language Translator",
            "diff": "Challenging",
            "one_liner": "Upload a video of someone signing -> AI tracks hand poses -> translates gestures to text -> speaks the translation.",
            "demo": "Upload a 30-second signing video -> watch hand tracking overlay -> see words appear -> hear spoken translation.",
            "input": "Student uploads signing video (recorded on phone)",
            "modules": [
                ["Video Processing", "Member 1", "OpenCV Frame Ingestion", "Extract frames, manage video pipeline"],
                ["Hand Tracking", "Member 2", "MediaPipe Hands (~300 MB)", "Real-time 21-point hand landmark detection on GPU"],
                ["Gesture Recognition", "Member 3", "LSTM on pose sequences (~500 MB)", "Classify gesture sequences into words/phrases"],
                ["Translation UI", "Member 4", "Frontend Subtitle Player", "Video player + synced translation, vocabulary lookup, TTS"]
            ]
        },
        {
            "name": "14. PixelPlayground — Creative AI Toolkit (5-person team)",
            "diff": "Standard+",
            "one_liner": "A Swiss Army knife of AI visual tools: background remover, style transfer, upscaler, colorizer, depth mapper — all in one app.",
            "demo": "Upload photo -> remove background -> apply style -> upscale -> see depth map. Each tool = different GPU model.",
            "input": "Student uploads photos",
            "modules": [
                ["Background Removal", "Member 1", "TinySAM (~500 MB) + rembg", "One-click bg removal, object cutout, bg swap"],
                ["Style Transfer + Filters", "Member 2", "Fast Neural Style (~1 GB)", "10+ art styles, adjustable intensity, real-time preview"],
                ["Super-Res + Colorization", "Member 3", "Real-ESRGAN (~500 MB) + DeOldify (~100 MB)", "4x upscaling, B&W colorization, photo enhancement"],
                ["Depth + 3D Preview", "Member 4", "Depth Anything V2 (~500 MB)", "Depth map, 3D parallax effect, depth-based blur"],
                ["Unified UI + AI Agent", "Member 5", "ChromaDB + Embeddings (~300 MB)", "Drag-drop canvas, tool palette, batch processing, AI suggests tools"]
            ]
        }
    ]
    
    for proj in projects_data:
        # Heading 3
        p_name = doc.add_paragraph()
        run_name = p_name.add_run(proj["name"])
        run_name.font.bold = True
        run_name.font.size = Pt(12)
        run_name.font.color.rgb = NAVY
        p_name.paragraph_format.space_before = Pt(12)
        p_name.paragraph_format.space_after = Pt(4)
        
        # Details Block
        p_det = doc.add_paragraph()
        p_det.paragraph_format.left_indent = Inches(0.15)
        p_det.paragraph_format.space_after = Pt(6)
        
        r_diff = p_det.add_run(f"Difficulty: {proj['diff']}\n")
        r_diff.font.bold = True
        r_diff.font.color.rgb = STEEL
        r_diff.font.size = Pt(9.5)
        
        p_det.add_run("▪  One-liner: ").font.bold = True
        p_det.add_run(f"{proj['one_liner']}\n").font.size = Pt(9.5)
        
        p_det.add_run("▪  Demo Moment: ").font.bold = True
        p_det.add_run(f"{proj['demo']}\n").font.size = Pt(9.5)
        
        p_det.add_run("▪  Data Input: ").font.bold = True
        p_det.add_run(f"{proj['input']}").font.size = Pt(9.5)
        
        # Add Modules Table
        add_project_modules_table(doc, proj["modules"])
        
    doc.add_page_break()
    
    # 6. GPU Diversity Matrix & Verification
    add_section_header(doc, "GPU Model Matrix & Verification Plan")
    p_mat = doc.add_paragraph()
    r_mat = p_mat.add_run(
        "To ensure high resource efficiency, no two teams run identical GPU stacks. All lightweight models fit "
        "comfortably inside Spark's 2GB VRAM allocation using FP16 weights, while heavier generative models (Stable Diffusion, MusicGen) "
        "leverage Blaze's 4GB VRAM. This layout ensures students run high-speed local inference sessions on the LaaS node "
        "without exceeding memory namespaces."
    )
    r_mat.font.size = Pt(10.5)
    
    # Save document
    output_path = r"c:\Users\Punith\LaaS\CS_15_Day_VibeCoding_Curriculum.docx"
    try:
        doc.save(output_path)
        print(f"SUCCESS: Curriculum document generated at: {output_path}")
    except Exception as e:
        print(f"ERROR saving to primary path: {e}")
        fallback_path = r"c:\Users\Punith\LaaS\CS_15_Day_VibeCoding_Curriculum_Final.docx"
        try:
            doc.save(fallback_path)
            print(f"SUCCESS: Generated at fallback path due to file lock: {fallback_path}")
        except Exception as e2:
            print(f"ERROR saving to fallback path: {e2}")

if __name__ == "__main__":
    build_document()
