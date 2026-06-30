import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Professional corporate colors: Navy & Charcoal theme
NAVY = RGBColor(26, 54, 93)       # #1A365D - Primary Accent
STEEL = RGBColor(43, 108, 176)    # #2B6CB0 - Secondary Accent
CHARCOAL = RGBColor(45, 55, 72)    # #2D3748 - Section Headers
BODY_TEXT = RGBColor(74, 85, 104)  # #4A5568 - Body Text
WHITE = RGBColor(255, 255, 255)
LIGHT_GRAY = "F7FAFC"             # #F7FAFC - Alternating rows
BORDER_GRAY = "E2E8F0"            # #E2E8F0 - Table borders
LINE_COLOR = "1A365D"             # Navy for accents

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_borders(cell, color_hex="E2E8F0", width="4"):
    """Set cell borders with specified color and width."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), width)
        border.set(qn('w:color'), color_hex)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    """Set cell padding in dxa (twentieths of a point). 140 dxa = 7pt, 180 dxa = 9pt."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        margin = OxmlElement(f'w:{margin_name}')
        margin.set(qn('w:w'), str(val))
        margin.set(qn('w:type'), 'dxa')
        tcMar.append(margin)
    tcPr.append(tcMar)

def style_document(doc):
    """Apply general document styling, margins, and body text formatting."""
    # Base normal style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = BODY_TEXT
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(6)
    
    # 1-inch margins
    for section in doc.sections:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)

def add_cover_page(doc):
    """Create a modern cover page with navy accents and bold professional hierarchy."""
    # Top spacing
    for _ in range(3):
        doc.add_paragraph()
    
    # Elegant Navy accent bar (using table trick)
    accent_table = doc.add_table(rows=1, cols=1)
    accent_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    accent_cell = accent_table.rows[0].cells[0]
    accent_cell.width = Inches(6.5)
    set_cell_shading(accent_cell, LINE_COLOR)
    p = accent_cell.paragraphs[0]
    p.add_run(" ")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    
    doc.add_paragraph()
    
    # Program Code or Tag
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("L&D INTERNSHIP COHORT - 15-DAY PROGRAM")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = STEEL
    p.paragraph_format.space_after = Pt(12)
    
    # Main Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AI-Powered Full-Stack Product Scaffolding")
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = NAVY
    p.paragraph_format.space_after = Pt(6)
    
    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("A 15-Day Vibe-Coding Curriculum for Teams from AI, CSE, IoT, and Cybersecurity")
    run.font.size = Pt(14)
    run.font.italic = True
    run.font.color.rgb = CHARCOAL
    p.paragraph_format.space_after = Pt(24)
    
    # Add vertical spacing
    for _ in range(4):
        doc.add_paragraph()
        
    # Metadata Block Table (Centered)
    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Prepared For:", "B.Tech Cohort (57 Interns, 14 Teams of 4, 1 Team of 5)"),
        ("Program Duration:", "15 Days (Days 1–10: Lectures, Days 11–15: Open Stage Coaching)"),
        ("Session Structure:", "Lecture Days: 1.0 hr Core Concept + 0.5 hr Doubt Resolution"),
        ("Expected Outcome:", "Functional Product POC incorporating UI, DB, and AI Agents"),
        ("Document Version:", "v1.5.0 (Final GPU Allocations - June 2026)")
    ]
    
    for r_idx, (label, val) in enumerate(meta_data):
        row = meta_table.rows[r_idx]
        
        cell_lbl = row.cells[0]
        set_cell_shading(cell_lbl, "F8FAFC")
        set_cell_borders(cell_lbl, "E2E8F0", "2")
        p_lbl = cell_lbl.paragraphs[0]
        r_lbl = p_lbl.add_run(label)
        r_lbl.font.bold = True
        r_lbl.font.size = Pt(9.5)
        r_lbl.font.color.rgb = CHARCOAL
        cell_lbl.width = Inches(2.0)
        set_cell_margins(cell_lbl, 80, 80, 100, 100)
        
        cell_val = row.cells[1]
        set_cell_shading(cell_val, "FFFFFF")
        set_cell_borders(cell_val, "E2E8F0", "2")
        p_val = cell_val.paragraphs[0]
        r_val = p_val.add_run(val)
        r_val.font.size = Pt(9.5)
        r_val.font.color.rgb = BODY_TEXT
        cell_val.width = Inches(4.5)
        set_cell_margins(cell_val, 80, 80, 100, 100)
        
    doc.add_page_break()

def add_section_header(doc, text):
    """Add a professional section header with a navy left-accent indicator."""
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.allow_autofit = False
    
    # Accent indicator bar (Left)
    accent_cell = table.rows[0].cells[0]
    accent_cell.width = Inches(0.12)
    set_cell_shading(accent_cell, LINE_COLOR)
    accent_p = accent_cell.paragraphs[0]
    accent_p.add_run(" ")
    accent_p.paragraph_format.space_before = Pt(0)
    accent_p.paragraph_format.space_after = Pt(0)
    
    # Text cell (Right)
    text_cell = table.rows[0].cells[1]
    text_cell.width = Inches(6.38)
    p = text_cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    
    doc.add_paragraph() # Spacing below heading

def add_curriculum_table(doc, items_data):
    """Generate the structured grid table with merged category visual dividers."""
    table = doc.add_table(rows=1 + len(items_data), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    
    headers = ["Day & Duration", "Curated Topic & Sub-topics", "First Principles & Value Add"]
    col_widths = [1.0, 2.2, 3.3] # Total = 6.5 inches
    
    # Table Header Row
    header_row = table.rows[0]
    for i, title in enumerate(headers):
        cell = header_row.cells[i]
        set_cell_shading(cell, LINE_COLOR)
        set_cell_borders(cell, LINE_COLOR, "4")
        set_cell_margins(cell, 160, 160, 180, 180)
        p = cell.paragraphs[0]
        run = p.add_run(title)
        run.font.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = WHITE
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        cell.width = Inches(col_widths[i])
        
    day_row_count = 0 # To track alternating colors for day rows
    
    # Fill Table Rows
    for r_idx, item in enumerate(items_data):
        row = table.rows[r_idx + 1]
        
        if item["type"] == "category":
            # Visual Divider Row: Merge all columns
            merged_cell = row.cells[0].merge(row.cells[1]).merge(row.cells[2])
            set_cell_shading(merged_cell, "2B6CB0") # Steel Blue background
            set_cell_borders(merged_cell, "2B6CB0", "4")
            set_cell_margins(merged_cell, 140, 140, 180, 180)
            
            p = merged_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(item["title"])
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = WHITE
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            
            merged_cell.width = Inches(6.5)
            
        else:
            # Alternating background color for day rows
            bg_color = "FFFFFF" if day_row_count % 2 == 0 else LIGHT_GRAY
            day_row_count += 1
            
            day_str = item["day"]
            topic_title = item["title"]
            subtopics = item["subtopics"]
            value_add = item["value"]
            
            # 1. Day Column
            cell_day = row.cells[0]
            set_cell_shading(cell_day, bg_color)
            set_cell_borders(cell_day, BORDER_GRAY, "4")
            set_cell_margins(cell_day, 120, 120, 140, 140)
            p_day = cell_day.paragraphs[0]
            run_day = p_day.add_run(day_str)
            run_day.font.bold = True
            run_day.font.size = Pt(10)
            run_day.font.color.rgb = CHARCOAL
            p_day.paragraph_format.space_after = Pt(0)
            cell_day.width = Inches(col_widths[0])
            
            # 2. Topic & Subtopics Column
            cell_topic = row.cells[1]
            set_cell_shading(cell_topic, bg_color)
            set_cell_borders(cell_topic, BORDER_GRAY, "4")
            set_cell_margins(cell_topic, 120, 120, 140, 140)
            
            p_topic = cell_topic.paragraphs[0]
            run_topic = p_topic.add_run(topic_title)
            run_topic.font.bold = True
            run_topic.font.size = Pt(10)
            run_topic.font.color.rgb = NAVY
            p_topic.paragraph_format.space_after = Pt(4)
            
            for sub in subtopics:
                p_sub = cell_topic.add_paragraph()
                run_sub = p_sub.add_run(f"▪  {sub}")
                run_sub.font.size = Pt(8.5)
                run_sub.font.color.rgb = BODY_TEXT
                p_sub.paragraph_format.space_after = Pt(2)
                p_sub.paragraph_format.left_indent = Inches(0.1)
                
            cell_topic.width = Inches(col_widths[1])
            
            # 3. Value Add Column
            cell_val = row.cells[2]
            set_cell_shading(cell_val, bg_color)
            set_cell_borders(cell_val, BORDER_GRAY, "4")
            set_cell_margins(cell_val, 120, 120, 140, 140)
            p_val = cell_val.paragraphs[0]
            run_val = p_val.add_run(value_add)
            run_val.font.size = Pt(9.5)
            run_val.font.color.rgb = BODY_TEXT
            p_val.paragraph_format.space_after = Pt(0)
            cell_val.width = Inches(col_widths[2])
            
    doc.add_paragraph()

def build_document():
    doc = Document()
    style_document(doc)
    
    # 1. Title Page
    add_cover_page(doc)
    
    # 2. Introduction & Philosophy
    add_section_header(doc, "Program Intro: The 'Vibe-Coding' Paradigm")
    
    p = doc.add_paragraph()
    r = p.add_run(
        "Software engineering is undergoing an unprecedented structural transition. The traditional model of "
        "manually writing boilerplate logic, configuration files, and basic HTML markup is rapidly giving way to "
        "'Vibe-Coding'—a collaborative workflow where engineers coordinate, direct, and audit autonomous AI agents "
        "to construct production-ready code. This 15-Day program is designed specifically for a diverse cohort of 57 "
        "students spanning AI, Computer Science (CSE), IoT, and Cybersecurity backgrounds. "
    )
    r.font.size = Pt(10.5)
    
    p2 = doc.add_paragraph()
    r2 = p2.add_run(
        "Organized into 14 collaborative teams of four (and one group of five), students are assigned a multi-dimensional product problem statement "
        "on Day 1 that merges UI, backend APIs, data storage, and AI agents. The daily lectures focus on core concept "
        "delivery and agentic pair-programming demonstrations, structured as 1 hour of teaching followed by 30 minutes of "
        "discussion and doubt resolution. Starting from Day 11, the cohort transitions to an Open Project Stage. Lectures "
        "conclude, and the session block becomes an open diagnostic space where teams receive hands-on mentoring, code-reviews, "
        "and architectural support to compile and launch their final POC."
    )
    r2.font.size = Pt(10.5)
    
    doc.add_paragraph()
    
    # 3. Schedule Table
    add_section_header(doc, "15-Day Daywise Schedule")
    
    curriculum_items = [
        # PART I: COHORT ONBOARDING & SETUP (DAYS 1-2)
        {"type": "category", "title": "PART I: COHORT ONBOARDING & SETUP (DAYS 1-2)"},
        {"type": "day", "day": "Day 1\n(2.0 Hours)", "title": "Program Intro & LaaS Portal Guide", "subtopics": [
            "Cohort introductions & team formations", "Signing up on the LaaS portal", "Launching remote desktops & booking slots", "ZFS home folder limits & hardware checks"
        ], "value": "Establishes immediate compute access. Interns learn how to launch and navigate isolated remote desktop sessions on local RTX 5090 machines."},
        {"type": "day", "day": "Day 2\n(2.0 Hours)", "title": "Vibe-Coding & Antigravity Setup", "subtopics": [
            "Understanding the 'Vibe-Coding' paradigm shift", "Downloading & installing the Antigravity developer environment", "Workspace configurations & environment variables", "Running basic agent commands & background checks"
        ], "value": "Prepares the local development environment, transitioning students from manual syntax writing to AI-driven command execution."},

        # PART II: PRODUCT WORKFLOWS & CONTEXT ENGINEERING (DAYS 3-6)
        {"type": "category", "title": "PART II: PRODUCT WORKFLOWS & CONTEXT ENGINEERING (DAYS 3-6)"},
        {"type": "day", "day": "Day 3\n(1.5 Hours)\n[1h Teach, 0.5h Doubt]", "title": "Product Dev Workflow & Harness (Part 1)", "subtopics": [
            "AI-assisted rapid prototyping models", "Setting up core repository directories", "Creating temporary code scratchpads", "Tracking file edits & compilation logs"
        ], "value": "Establishes the write-run-fix loop, showing how to coordinate AI outputs with real-time execution outputs to audit code changes."},
        {"type": "day", "day": "Day 4\n(1.5 Hours)\n[1h Teach, 0.5h Doubt]", "title": "Product Dev Workflow & Harness (Part 2)", "subtopics": [
            "Harness pipeline scripting & configs", "Running automated code verification steps", "Refactoring code blocks via agent commands", "Managing prompt-compile cycles"
        ], "value": "Enables test-driven agent logic, showing interns how to instruct agents to build, run, and self-correct files based on terminal execution errors."},
        {"type": "day", "day": "Day 5\n(1.5 Hours)\n[1h Teach, 0.5h Doubt]", "title": "Context Engineering: Rules & Skills (Part 1)", "subtopics": [
            "System prompts vs. user input structures", "Designing custom instructions files", "Controlling token boundaries", "Context size optimization techniques"
        ], "value": "Reduces model hallucinations. Interns learn to restrict and focus model contexts so the AI pair programmer outputs accurate results."},
        {"type": "day", "day": "Day 6\n(1.5 Hours)\n[1h Teach, 0.5h Doubt]", "title": "Context Engineering: Rules & Skills (Part 2)", "subtopics": [
            "Writing custom .qoder workspace rules", "Injecting team coding style guides", "Mapping specific workspace tools (skills)", "Enforcing API/schema coding boundaries"
        ], "value": "Ensures generated code aligns with project structures, database requirements, and repository parameters automatically."},

        # PART III: CORE WEB APIs, DATA & LLMs (DAYS 7-10)
        {"type": "category", "title": "PART III: CORE WEB APIs, DATA & LLMs (DAYS 7-10)"},
        {"type": "day", "day": "Day 7\n(1.5 Hours)\n[1h Teach, 0.5h Doubt]", "title": "Generative AI: LLM Fundamentals & APIs", "subtopics": [
            "Attention & Transformer architecture basics", "Temperature parameters & context constraints", "Securing API tokens & handling connection retries", "Enforcing outputs via JSON mode"
        ], "value": "Transitions students from standard chat apps to calling LLM APIs programmatically within Python or Node.js backend handlers."},
        {"type": "day", "day": "Day 8\n(1.5 Hours)\n[1h Teach, 0.5h Doubt]", "title": "Vector Databases, RAG & Prompting", "subtopics": [
            "Text embeddings & mathematical token distance", "Chroma DB & pgvector similarity search", "Implementing text chunking & overlap logic", "RAG query context injection pipelines"
        ], "value": "Enables semantic searches. Interns search local manuals, databases, and code directories semantically without expensive model fine-tuning costs."},
        {"type": "day", "day": "Day 9\n(1.5 Hours)\n[1h Teach, 0.5h Doubt]", "title": "Agentic Systems: Harness, Rules & Skills", "subtopics": [
            "Defining Agent autonomy: Reason-Act (ReAct) loop", "Function schema declaration to LLMs", "Managing states across recursive agent loops", "Setting sandbox tool execution parameters"
        ], "value": "Builds autonomous agents that can evaluate outputs, call external system tools, and iterate to solve complex multi-step tasks."},
        {"type": "day", "day": "Day 10\n(1.5 Hours)\n[1h Teach, 0.5h Doubt]", "title": "Multi-Agent Workflows & Orchestrations", "subtopics": [
            "Visual workflow automation using n8n", "Coordinating event-based triggers (webhooks, schedules)", "Chaining multi-agent inputs/outputs", "Orchestrating agent hierarchies"
        ], "value": "Teaches enterprise workflow integrations, showing how to connect databases, REST APIs, and agent decisions into clean automated pipelines."},

        # PART IV: OPEN PROJECT COACHING STAGE (DAYS 11-15)
        {"type": "category", "title": "PART IV: OPEN PROJECT COACHING STAGE (DAYS 11-15)"},
        {"type": "day", "day": "Days 11–15\n(Open stage daily)", "title": "Ongoing Project Mentoring & Doubt Clearing", "subtopics": [
            "Daily team progress & block checks", "Workspace debugging & environment fixes", "API routing & database connector audits", "Live code-reviews & optimization feedback", "Final project demo pitching rehearsals"
        ], "value": "Consolidates all program knowledge. Providing direct engineering support as teams build their POC asynchronously, culminating in the Day 15 showcase."}
    ]
    
    add_curriculum_table(doc, curriculum_items)
    
    # 4. 14 Unique Projects Section
    add_section_header(doc, "Collaborative Projects: Specifications for 14 Teams")
    
    p = doc.add_paragraph()
    r = p.add_run(
        "Each of the 14 teams is assigned a unique project on Day 1. The specifications below outline how each project "
        "is split across specialized modules. Resource Allocation: 1 member gets a Blaze instance (4GB VRAM, 8GB RAM) "
        "for the core orchestrator/LLM, while the other members get Spark instances (2GB VRAM, 4GB RAM) for localized UI/Backend/GPU inference. "
        "All projects are designed to utilize pre-trained, low-resource GPU models running locally, ensuring compatibility with the LaaS environment:"
    )
    r.font.size = Pt(10.5)

    # 4b. WebCam & Audio capture guide
    add_section_header(doc, "Live Camera & Audio Feed Integration Workaround on Remote LaaS Containers")
    p_wc = doc.add_paragraph()
    r_wc = p_wc.add_run(
        "Since the LaaS instances are remote Docker containers on central GPU nodes without physical webcams or microphones, "
        "teams will use one of these standard workarounds to process live media feeds:\n\n"
        "1. Browser-based HTML5 Capture (Recommended Web Integration): The student's React/HTML frontend captures "
        "the local phone/laptop webcam or mic feed using browser standard APIs (navigator.mediaDevices.getUserMedia) "
        "and transmits frame snapshots (as base64 or binary blobs) via HTTP POST / WebSockets to the remote FastAPI backend "
        "for immediate GPU inference.\n"
        "2. Pre-recorded Media Fallback (Testing Pattern): Students record short 10-second videos (.mp4) or audio notes (.wav) "
        "on their smartphones, upload them to their remote home directory (/home/ubuntu) via the Selkies browser "
        "VNC file-transfer dashboard, and run local Python/OpenCV inference scripts directly on disk."
    )
    r_wc.font.size = Pt(10.0)

    doc.add_paragraph()
    
    projects = [
        ("Team 1: Smart City Parking Lot & License Plate Auditor (Smart City / Property Management)", 
         "Monitors parking occupancy and identifies authorized/blocked vehicles.\n"
         "• Module 1 (Blaze - Lead): Parking Coordinator Agent (Runs Qwen-2.5-1.5B on GPU to match plates, check reservations, and draft alert emails).\n"
         "• Module 2 (Spark): Space Occupancy Spotter (Runs YOLOv8-Nano on GPU to count empty vs. full spots from camera feeds).\n"
         "• Module 3 (Spark): License Plate Reader (Runs EasyOCR on GPU to extract license plate text from vehicle photos).\n"
         "• Module 4 (Spark): Reservation API & Web UI (FastAPI + PostgreSQL database tracking slots and displaying real-time occupancy)."),
         
        ("Team 2: Audio Support Helpdesk for Local Utilities (Public Services / Accessibility)", 
         "Allows citizens to record and submit voice complaints about local infrastructure issues.\n"
         "• Module 1 (Blaze - Lead): Ticket Dispatch Agent (Runs Llama-3.2-1B to summarize transcripts, assess urgency, and draft tasks).\n"
         "• Module 2 (Spark): Voice Transcriber (Runs Whisper-Tiny on GPU to transcribe incoming audio clips).\n"
         "• Module 3 (Spark): Duplicate Matcher (Runs Sentence-Transformers on GPU to check if an issue was already reported).\n"
         "• Module 4 (Spark): Support Ticket API & Dashboard (Express + MongoDB handling ticket statuses and audio playback)."),
         
        ("Team 3: Retail Storefront Traffic & Engagement Analyzer (Retail Analytics)", 
         "Analyzes visitor foot-traffic and attention responses to retail window displays.\n"
         "• Module 1 (Blaze - Lead): Foot-Traffic Analyst Agent (Runs Llama-3.2-1B to analyze weekly traffic charts and suggest layouts).\n"
         "• Module 2 (Spark): Visitor Counter (Runs YOLOv8-Nano on GPU to detect and count people crossing the entrance threshold).\n"
         "• Module 3 (Spark): Posture & Attention Classifier (Runs a simple pre-trained CNN on GPU to categorize visitor focus/looking).\n"
         "• Module 4 (Spark): Traffic Ledger API & UI (FastAPI + PostgreSQL storing visitor logs and timestamps)."),
         
        ("Team 4: Self-Service School Library & Document Archiver (Education / Corporate)", 
         "Digitizes, indexes, and queries physical books or flyers locally without external cloud dependencies.\n"
         "• Module 1 (Blaze - Lead): Runbook RAG Assistant (Runs Llama-3.2-1B with ChromaDB to answer student queries on scanned pages).\n"
         "• Module 2 (Spark): Scan Digitizer (Runs EasyOCR on GPU to extract raw text blocks from book page snapshots).\n"
         "• Module 3 (Spark): Text Embedder (Runs Sentence-Transformers on GPU to generate vectors for RAG indexing).\n"
         "• Module 4 (Spark): Library Portal UI (React app for document upload, index tracking, and semantic search queries)."),
         
        ("Team 5: Warehouse Cargo Loading & Volume Auditor (Logistics / Warehouse)", 
         "Verifies that packages loaded onto delivery trucks match their shipping manifests in size and quantity.\n"
         "• Module 1 (Blaze - Lead): Manifest Reconciler Agent (Runs Qwen-2.5-1.5B to compare visual box counts against digital invoices).\n"
         "• Module 2 (Spark): Package Counter (Runs YOLOv8-Nano on GPU to count and localize packages on the conveyor belt).\n"
         "• Module 3 (Spark): Package Feature Classifier (Runs CLIP on GPU to crosscheck package branding/labels against expected products).\n"
         "• Module 4 (Spark): Dispatch Logs API & UI (Express + MongoDB backend storing loading logs and showcasing discrepancies)."),
         
        ("Team 6: CCTV Safety Compliance Monitor for Warehouses (Smart Workplace)", 
         "Enforces warehouse safety compliance codes (such as high-visibility vests and clear exit routes).\n"
         "• Module 1 (Blaze - Lead): Safety Compliance Reporter (Runs Llama-3.2-1B to compile violation logs into formal daily reports).\n"
         "• Module 2 (Spark): Vest Detector (Runs YOLOv8-Nano on GPU to flag personnel not wearing safety vests).\n"
         "• Module 3 (Spark): Obstruction Watcher (Runs simple OpenCV masking on GPU to detect blockages in safety pathways).\n"
         "• Module 4 (Spark): Alarm Gateway & UI (FastAPI backend logging violations to PostgreSQL and playing alarms on the dashboard)."),
         
        ("Team 7: Solar Array Telemetry & Hotspot Auditor (Energy / Utilities)", 
         "Monitors solar panel voltage outputs and flags visual degradation or hotspots.\n"
         "• Module 1 (Blaze - Lead): Maintenance Dispatcher Agent (Runs Llama-3.2-1B to read panel faults and write technician schedules).\n"
         "• Module 2 (Spark): Telemetry Anomaly Predictor (Runs a simple PyTorch LSTM on GPU to predict future voltage drops based on logs).\n"
         "• Module 3 (Spark): Thermal Spot Scanner (Runs YOLOv8-Nano on infrared images to identify panel thermal anomalies).\n"
         "• Module 4 (Spark): Grid Operations API & UI (React UI + Node.js backend logging telemetry and rendering panel health)."),
         
        ("Team 8: Alternative Risk Assessor for Street Vendors (Fintech / Lending)", 
         "Evaluates creditworthiness for street vendors using alternative photo-receipt and asset verification.\n"
         "• Module 1 (Blaze - Lead): Underwriting Agent (Runs Qwen-2.5-1.5B to score credit based on receipt cash flow and peer logs).\n"
         "• Module 2 (Spark): Receipt Reader (Runs EasyOCR on GPU to extract transactional logs from photographed paper receipts).\n"
         "• Module 3 (Spark): Asset Classifier (Runs CLIP on GPU to verify if photos uploaded represent active vendor stands).\n"
         "• Module 4 (Spark): Application UI & Database (FastAPI + PostgreSQL managing vendor applications and loan histories)."),
         
        ("Team 9: Cafeteria Food Waste Counter & Menu Planner (Hospitality / Food Tech)", 
         "Tracks buffet plate waste categories to help catering businesses plan cost-optimized menus.\n"
         "• Module 1 (Blaze - Lead): Menu Planner Agent (Runs Qwen-2.5-1.5B to analyze plate waste lists and suggest cost-optimized recipe mixes).\n"
         "• Module 2 (Spark): Food Visual Auditor (Runs YOLOv8-Nano on GPU to recognize leftover food items on discarded plates).\n"
         "• Module 3 (Spark): Menu Card Reader (Runs EasyOCR on GPU to digitize weekly hand-written cafeteria menu schedules).\n"
         "• Module 4 (Spark): Waste Tracker API & UI (Express + MongoDB database logging waste items and timestamps)."),
         
        ("Team 10: Local Tourism Audio Guide Generator (Tourism / Cultural Heritage)", 
         "Generates localized, context-aware audio tours when visitors scan historical monument QR codes.\n"
         "• Module 1 (Blaze - Lead): Storyteller Agent (Runs Qwen-2.5-1.5B to generate site narratives based on scanned location tags).\n"
         "• Module 2 (Spark): Local Speech Generator (Runs a lightweight GPU TTS pipeline to convert text stories into audio files).\n"
         "• Module 3 (Spark): Semantic Query Matcher (Runs Sentence-Transformers on GPU to match visitor text queries to archive facts).\n"
         "• Module 4 (Spark): Tourist Web UI (React web app serving audio files, maps, and local site guides)."),
         
        ("Team 11: Smart Classroom Attendance & Posture Auditor (EdTech)", 
         "Tracks classroom attendance and student posture/attentiveness levels during lectures.\n"
         "• Module 1 (Blaze - Lead): Engagement Reporter Agent (Runs Llama-3.2-1B to analyze attendance tables and generate summary emails).\n"
         "• Module 2 (Spark): Attendance Counter (Runs YOLOv8-Nano on GPU to count faces present in classroom camera snapshots).\n"
         "• Module 3 (Spark): Posture Classifier (Runs a basic PyTorch CNN on GPU to flag students who are asleep or distracted).\n"
         "• Module 4 (Spark): Roster API & UI (FastAPI + PostgreSQL tracking schedules, student profiles, and roster data)."),
         
        ("Team 12: E-Commerce Product Catalog Auto-Tagger (E-Commerce)", 
         "Generates optimized SEO descriptions and product tags for small e-commerce merchants automatically.\n"
         "• Module 1 (Blaze - Lead): Copywriter Agent (Runs Llama-3.2-1B to generate product listings and titles from visual tag logs).\n"
         "• Module 2 (Spark): Feature Tagger (Runs CLIP on GPU to extract item colors, patterns, and categories from images).\n"
         "• Module 3 (Spark): Catalog Search Indexer (Runs Sentence-Transformers on GPU to enable semantic search over vendor items).\n"
         "• Module 4 (Spark): Seller API & Catalog UI (Node.js + MongoDB managing product listings and stocks)."),
         
        ("Team 13: API Guardrail & Prompt Injection Shield (Cybersecurity)", 
         "Protects local AI APIs against prompt injections, adversarial queries, and data scraping.\n"
         "• Module 1 (Blaze - Lead): Security Planner Agent (Runs Llama-3.2-1B to analyze threat logs and automatically generate IP blocking rules).\n"
         "• Module 2 (Spark): Prompt Injection Guard (Runs a tiny BERT model on GPU to classify and block injection prompts).\n"
         "• Module 3 (Spark): Rate Anomaly Classifier (Runs a PyTorch Autoencoder on GPU to identify rapid query scraping signatures).\n"
         "• Module 4 (Spark): Reverse Proxy API (FastAPI intercepting incoming queries, logging traffic to PostgreSQL)."),
         
        ("Team 14: Smart Warehouse Environmental & Pest Watch (Manufacturing / Food Storage) [5 Members]", 
         "Monitors grain/paper storage quality and checks for early pest infestation.\n"
         "• Module 1 (Blaze - Lead): Compliance Auditor Agent (Runs Qwen-2.5-1.5B to write safety compliance logs and recommend sensor adjustments).\n"
         "• Module 2 (Spark): Pest Visual Detector (Runs YOLOv8-Nano on GPU to scan camera frames for insects or rodents in bait stations).\n"
         "• Module 3 (Spark): Humidity Curve Predictor (Runs a PyTorch LSTM on GPU to forecast mold risks based on temperature curves).\n"
         "• Module 4 (Spark): Ingestion API (Node.js backend collecting sensor telemetry from MQTT, writing to MongoDB).\n"
         "• Module 5 (Spark): Operations Monitor (React dashboard tracking real-time status and triggering SMS warnings).")
    ]
    
    for title, desc in projects:
        p_proj = doc.add_paragraph()
        run_title = p_proj.add_run(f"•  {title}\n")
        run_title.font.bold = True
        run_title.font.size = Pt(11)
        run_title.font.color.rgb = NAVY
        
        run_desc = p_proj.add_run(desc)
        run_desc.font.size = Pt(9.5)
        run_desc.font.color.rgb = BODY_TEXT
        p_proj.paragraph_format.left_indent = Inches(0.2)
        p_proj.paragraph_format.space_after = Pt(8)
        
    doc.add_paragraph()
    
    # Save document
    output_path = r"c:\Users\Punith\LaaS\CS_15_Day_VibeCoding_Curriculum.docx"
    try:
        doc.save(output_path)
        print(f"SUCCESS: Curriculum document generated at: {output_path}")
    except Exception as e:
        print(f"ERROR: Error saving document: {e}")

if __name__ == "__main__":
    build_document()
