import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Professional corporate colors: Navy & Charcoal theme
NAVY = RGBColor(26, 54, 93)       # #1A365D - Primary Accent
STEEL = RGBColor(43, 108, 176)    # #2B6CB0 - Secondary Accent
CHARCOAL = RGBColor(45, 55, 72)    # #2D3748 - Section Headers
BODY_TEXT = RGBColor(74, 85, 104)  # #4A5568 - Body Text
WHITE = RGBColor(255, 255, 255)
LIGHT_GRAY = "F7FAFC"             # #F7FAFC - Alternating rows
BORDER_GRAY = "E2E8F0"            # #E2E8F0 - Table borders
LINE_COLOR = "1A365D"             # Navy for accents

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_borders(cell, color_hex="E2E8F0", width="4"):
    """Set cell borders with specified color and width."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), width)
        border.set(qn('w:color'), color_hex)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    """Set cell padding in dxa (twentieths of a point). 140 dxa = 7pt, 180 dxa = 9pt."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        margin = OxmlElement(f'w:{margin_name}')
        margin.set(qn('w:w'), str(val))
        margin.set(qn('w:type'), 'dxa')
        tcMar.append(margin)
    tcPr.append(tcMar)

def style_document(doc):
    """Apply general document styling, margins, and body text formatting."""
    # Base normal style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = BODY_TEXT
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(6)
    
    # 1-inch margins
    for section in doc.sections:
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)

def add_cover_page(doc):
    """Create a modern cover page with navy accents and bold professional hierarchy."""
    # Top spacing
    for _ in range(3):
        doc.add_paragraph()
    
    # Elegant Navy accent bar (using table trick)
    accent_table = doc.add_table(rows=1, cols=1)
    accent_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    accent_cell = accent_table.rows[0].cells[0]
    accent_cell.width = Inches(6.5)
    set_cell_shading(accent_cell, LINE_COLOR)
    p = accent_cell.paragraphs[0]
    p.add_run(" ")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    
    doc.add_paragraph()
    
    # Program Code or Tag
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("L&D INTERNSHIP COHORT - 15-DAY PROGRAM")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = STEEL
    p.paragraph_format.space_after = Pt(12)
    
    # Main Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AI-Powered Full-Stack Product Scaffolding")
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = NAVY
    p.paragraph_format.space_after = Pt(6)
    
    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("A 15-Day Vibe-Coding Curriculum for Teams from AI, CSE, IoT, and Cybersecurity")
    run.font.size = Pt(14)
    run.font.italic = True
    run.font.color.rgb = CHARCOAL
    p.paragraph_format.space_after = Pt(24)
    
    # Add vertical spacing
    for _ in range(4):
        doc.add_paragraph()
        
    # Metadata Block Table (Centered)
    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Prepared For:", "Diverse Cohort (57 Interns, 14 Teams of 4)"),
        ("Program Duration:", "15 Days (Daily 1.5 Hour Sessions)"),
        ("Instruction Mode:", "100% Concept Demos (1.5h) + Async Team Hacking"),
        ("Final Milestone:", "Working Full-Stack AI Product POC (UI+Backend+DB+Agents)"),
        ("Document Version:", "v1.0.0 (June 2026)")
    ]
    
    for r_idx, (label, val) in enumerate(meta_data):
        row = meta_table.rows[r_idx]
        
        cell_lbl = row.cells[0]
        set_cell_shading(cell_lbl, "F8FAFC")
        set_cell_borders(cell_lbl, "E2E8F0", "2")
        p_lbl = cell_lbl.paragraphs[0]
        r_lbl = p_lbl.add_run(label)
        r_lbl.font.bold = True
        r_lbl.font.size = Pt(9.5)
        r_lbl.font.color.rgb = CHARCOAL
        cell_lbl.width = Inches(2.0)
        set_cell_margins(cell_lbl, 80, 80, 100, 100)
        
        cell_val = row.cells[1]
        set_cell_shading(cell_val, "FFFFFF")
        set_cell_borders(cell_val, "E2E8F0", "2")
        p_val = cell_val.paragraphs[0]
        r_val = p_val.add_run(val)
        r_val.font.size = Pt(9.5)
        r_val.font.color.rgb = BODY_TEXT
        cell_val.width = Inches(4.5)
        set_cell_margins(cell_val, 80, 80, 100, 100)
        
    doc.add_page_break()

def add_section_header(doc, text):
    """Add a professional section header with a navy left-accent indicator."""
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.allow_autofit = False
    
    # Accent indicator bar (Left)
    accent_cell = table.rows[0].cells[0]
    accent_cell.width = Inches(0.12)
    set_cell_shading(accent_cell, LINE_COLOR)
    accent_p = accent_cell.paragraphs[0]
    accent_p.add_run(" ")
    accent_p.paragraph_format.space_before = Pt(0)
    accent_p.paragraph_format.space_after = Pt(0)
    
    # Text cell (Right)
    text_cell = table.rows[0].cells[1]
    text_cell.width = Inches(6.38)
    p = text_cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    
    doc.add_paragraph() # Spacing below heading

def add_curriculum_table(doc, items_data):
    """Generate the structured grid table with merged category visual dividers."""
    table = doc.add_table(rows=1 + len(items_data), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    
    headers = ["Day", "Curated Topic & Sub-topics", "First Principles & Value Add"]
    col_widths = [0.8, 2.2, 3.5] # Total = 6.5 inches
    
    # Table Header Row
    header_row = table.rows[0]
    for i, title in enumerate(headers):
        cell = header_row.cells[i]
        set_cell_shading(cell, LINE_COLOR)
        set_cell_borders(cell, LINE_COLOR, "4")
        set_cell_margins(cell, 160, 160, 180, 180)
        p = cell.paragraphs[0]
        run = p.add_run(title)
        run.font.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = WHITE
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        cell.width = Inches(col_widths[i])
        
    day_row_count = 0 # To track alternating colors for day rows
    
    # Fill Table Rows
    for r_idx, item in enumerate(items_data):
        row = table.rows[r_idx + 1]
        
        if item["type"] == "category":
            # Visual Divider Row: Merge all columns
            merged_cell = row.cells[0].merge(row.cells[1]).merge(row.cells[2])
            set_cell_shading(merged_cell, "2B6CB0") # Steel Blue background
            set_cell_borders(merged_cell, "2B6CB0", "4")
            set_cell_margins(merged_cell, 140, 140, 180, 180)
            
            p = merged_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(item["title"])
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = WHITE
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            
            merged_cell.width = Inches(6.5)
            
        else:
            # Alternating background color for day rows
            bg_color = "FFFFFF" if day_row_count % 2 == 0 else LIGHT_GRAY
            day_row_count += 1
            
            day_str = item["day"]
            topic_title = item["title"]
            subtopics = item["subtopics"]
            value_add = item["value"]
            
            # 1. Day Column
            cell_day = row.cells[0]
            set_cell_shading(cell_day, bg_color)
            set_cell_borders(cell_day, BORDER_GRAY, "4")
            set_cell_margins(cell_day, 120, 120, 140, 140)
            p_day = cell_day.paragraphs[0]
            run_day = p_day.add_run(day_str)
            run_day.font.bold = True
            run_day.font.size = Pt(10)
            run_day.font.color.rgb = CHARCOAL
            p_day.paragraph_format.space_after = Pt(0)
            cell_day.width = Inches(col_widths[0])
            
            # 2. Topic & Subtopics Column
            cell_topic = row.cells[1]
            set_cell_shading(cell_topic, bg_color)
            set_cell_borders(cell_topic, BORDER_GRAY, "4")
            set_cell_margins(cell_topic, 120, 120, 140, 140)
            
            p_topic = cell_topic.paragraphs[0]
            run_topic = p_topic.add_run(topic_title)
            run_topic.font.bold = True
            run_topic.font.size = Pt(10)
            run_topic.font.color.rgb = NAVY
            p_topic.paragraph_format.space_after = Pt(4)
            
            for sub in subtopics:
                p_sub = cell_topic.add_paragraph()
                run_sub = p_sub.add_run(f"▪  {sub}")
                run_sub.font.size = Pt(8.5)
                run_sub.font.color.rgb = BODY_TEXT
                p_sub.paragraph_format.space_after = Pt(2)
                p_sub.paragraph_format.left_indent = Inches(0.1)
                
            cell_topic.width = Inches(col_widths[1])
            
            # 3. Value Add Column
            cell_val = row.cells[2]
            set_cell_shading(cell_val, bg_color)
            set_cell_borders(cell_val, BORDER_GRAY, "4")
            set_cell_margins(cell_val, 120, 120, 140, 140)
            p_val = cell_val.paragraphs[0]
            run_val = p_val.add_run(value_add)
            run_val.font.size = Pt(9.5)
            run_val.font.color.rgb = BODY_TEXT
            p_val.paragraph_format.space_after = Pt(0)
            cell_val.width = Inches(col_widths[2])
            
    doc.add_paragraph()

def build_document():
    doc = Document()
    style_document(doc)
    
    # 1. Title Page
    add_cover_page(doc)
    
    # 2. Introduction & Paradigm Shift
    add_section_header(doc, "Program Intro: The 'Vibe-Coding' Paradigm")
    
    p = doc.add_paragraph()
    r = p.add_run(
        "Software engineering is undergoing an unprecedented structural transition. The traditional model of "
        "manually writing boilerplate logic, configuration files, and basic HTML markup is rapidly giving way to "
        "'Vibe-Coding'—a collaborative workflow where engineers coordinate, direct, and audit autonomous AI agents "
        "to construct production-ready code. This 15-Day program is designed specifically for a diverse cohort of 57 "
        "students spanning AI, Computer Science (CSE), IoT, and Cybersecurity backgrounds. "
    )
    r.font.size = Pt(10.5)
    
    p2 = doc.add_paragraph()
    r2 = p2.add_run(
        "Organized into 14 collaborative teams of four, students are assigned a multi-dimensional product problem statement "
        "on Day 1 that merges UI, backend APIs, data storage, and AI agents. The 1.5-hour daily sessions deliver core first-principles "
        "concepts and live agentic programming demonstrations. Teams hack asynchronously outside class hours, applying "
        "the lecture knowledge step-by-step to compile their product. This ensures they master system basics (OS, Linux, SQL, MongoDB) "
        "while using cutting-edge AI architectures (RAG, OpenClaw, Hermes, vector databases, n8n, secure gateways) to produce a working POC."
    )
    r2.font.size = Pt(10.5)
    
    doc.add_paragraph()
    
    # 3. Schedule Table
    add_section_header(doc, "15-Day Daywise Schedule")
    
    curriculum_items = [
        # PART I: CORE WORKSPACE & SYSTEMS (DAYS 1-3)
        {"type": "category", "title": "PART I: CORE WORKSPACE & SYSTEMS (DAYS 1-3)"},
        {"type": "day", "day": "Day 1", "title": "Antigravity Intro & Workspace Setup", "subtopics": [
            "Introduction to AI-Assisted Pair Programming", "Workspace environment variables & directory structures", "Triggering background tasks & logs", "Antigravity slash commands (/goal, /grill-me)"
        ], "value": "Establishes immediate developer velocity. Interns learn how to leverage agent workspaces to execute tasks, configure project settings, and automate coding iterations."},
        {"type": "day", "day": "Day 2", "title": "OS & Linux: Compute Arch on LaaS", "subtopics": [
            "Process execution, monitoring & thread states", "Resource allocation (CPU/RAM limits & GPU context)", "LaaS remote compute desktop environment", "Linux Filesystem Hierarchy Standard (FHS)"
        ], "value": "Removes GUI dependencies. Students learn how programs utilize CPU/GPU hardware on local supercomputing nodes, prepping them for deep backend debugging."},
        {"type": "day", "day": "Day 3", "title": "Shell CLI, Pipes & Automation Scripting", "subtopics": [
            "Terminal navigation & standard streams (stdin, stdout, stderr)", "Piping streams and filters (grep, find, xargs)", "Writing robust Bash shell scripts", "Task automation via Cron scheduling"
        ], "value": "Teaches command-line automation first-principles, enabling interns to write scripts to install dependencies, clean directories, and verify application processes."},

        # PART II: TEAM COLLABORATION & DATABASES (DAYS 4-6)
        {"type": "category", "title": "PART II: TEAM COLLABORATION & DATABASES (DAYS 4-6)"},
        {"type": "day", "day": "Day 4", "title": "Git Collaboration & Team Development", "subtopics": [
            "Git internals: commit graphs, staging & blobs", "Collaborative branching models (Feature branches, main)", "Opening Pull Requests (PRs) & code review etiquette", "Resolving merge conflicts and git rebase"
        ], "value": "Crucial for team alignment. Empowers 4-member teams to develop, review, merge, and pull code simultaneously without breaking the central codebase."},
        {"type": "day", "day": "Day 5", "title": "Relational Databases & SQL Schema Design", "subtopics": [
            "ACID properties & transactional reliability", "SQL schema design: PKs, FKs, and normalization", "Efficient lookups using B-Tree indexing", "PostgreSQL database configuration"
        ], "value": "Teaches structural data consistency. Interns learn to write clean relational schemas, preventing database locks and ensuring transactional integrity."},
        {"type": "day", "day": "Day 6", "title": "Non-Relational Databases: MongoDB", "subtopics": [
            "CAP Theorem: Consistency, Availability, Partition Tolerance", "Document-store model & collection modeling", "Handling unstructured datasets vs. relational data", "Connecting node/python clients to MongoDB"
        ], "value": "Explains horizontal scaling trade-offs, teaching students when to swap strict SQL constraints for fast, flexible JSON document storage."},

        # PART III: AI, VECTORS & RAG PIPELINES (DAYS 7-9)
        {"type": "category", "title": "PART III: AI, VECTORS & RAG PIPELINES (DAYS 7-9)"},
        {"type": "day", "day": "Day 7", "title": "Generative AI: LLM Fundamentals & APIs", "subtopics": [
            "LLM architecture overview (Transformers & attention)", "Tokens, context windows, and cost optimization", "System Prompts vs. User Prompts & JSON Mode output", "API request orchestration (handling timeouts & retries)"
        ], "value": "Bridges AI concepts with application logic, training interns to extract structured JSON data from raw LLM responses to drive backend business flows."},
        {"type": "day", "day": "Day 8", "title": "Vector Databases & Semantic Embeddings", "subtopics": [
            "Converting text data to high-dimensional vector embeddings", "Chroma DB and pgvector extensions", "Distance metrics: Cosine, Dot Product, L2", "Indexing vectors for sub-second similarity search"
        ], "value": "Enables unstructured semantic search. Students learn how search engines map mathematical meanings, setting up local context retrieval."},
        {"type": "day", "day": "Day 9", "title": "RAG Pipelines & Advanced Prompting", "subtopics": [
            "Retrieval-Augmented Generation (RAG) architecture", "Text chunking algorithms & overlap configurations", "Context injection and system prompt assembly", "Mitigating LLM hallucinations using ground truth data"
        ], "value": "Equips interns to feed private enterprise data to public LLMs, building secure search systems that answer based on local knowledge bases."},

        # PART IV: AUTONOMOUS AI AGENTS & ORCHESTRATION (DAYS 10-13)
        {"type": "category", "title": "PART IV: AUTONOMOUS AI AGENTS & ORCHESTRATION (DAYS 10-13)"},
        {"type": "day", "day": "Day 10", "title": "AI Agents: ReAct Loops & Function Calling", "subtopics": [
            "Defining Agent autonomy: Reason-Act (ReAct) frameworks", "Function calling: declaring API tool schemas to LLMs", "Parsing agent actions from text outputs", "Managing state across multi-turn agent loops"
        ], "value": "Transitions from static prompts to autonomous execution, showing how AI can select tools, write arguments, and act based on execution results."},
        {"type": "day", "day": "Day 11", "title": "Multi-Agent Systems & n8n Workflows", "subtopics": [
            "Low-code automation with n8n workflow designer", "Creating event triggers (webhooks, email, scheduled)", "n8n AI Nodes: injecting agents, tools, memory", "Chaining multi-agent inputs/outputs"
        ], "value": "Teaches enterprise-level system integration, allowing interns to automate complex pipelines (scrapers, triggers, notifications) with zero manual code."},
        {"type": "day", "day": "Day 12", "title": "OpenClaw Agent Orchestration Framework", "subtopics": [
            "OpenClaw core architecture & multi-agent managers", "Defining system instructions & skill definitions", "Writing custom execution tools in Python", "Tracing agent decision paths and debugging logs"
        ], "value": "Exposes interns to code-first agent design, teaching them to construct, customize, and orchestrate custom Python agent teams."},
        {"type": "day", "day": "Day 13", "title": "Local LLMs & Hermes Agent Execution", "subtopics": [
            "Ollama ecosystem: pulling & serving open-weights models", "Running Hermes / LLaMA 3 locally on LaaS GPU nodes", "Local hardware performance constraints (VRAM, memory)", "Local model security vs. public cloud APIs"
        ], "value": "Demystifies self-hosted AI, teaching students how to run models locally to guarantee data privacy, eliminate API costs, and utilize supercomputing nodes."},

        # PART V: BACKEND INTEGRATION & GenAI SECURITY (DAYS 14-15)
        {"type": "category", "title": "PART V: BACKEND INTEGRATION & GenAI SECURITY (DAYS 14-15)"},
        {"type": "day", "day": "Day 14", "title": "Backend API Development & Integration", "subtopics": [
            "RESTful API principles (Endpoints, HTTP methods, JSON payloads)", "Backend routing using FastAPI (Python) or Express (Node.js)", "Connecting endpoints to PostgreSQL/MongoDB clients", "Integrating agent execution loops into HTTP request handlers"
        ], "value": "Ties the product together, teaching interns how to expose databases and AI agent actions as secure APIs that any web frontend can call."},
        {"type": "day", "day": "Day 15", "title": "GenAI Security, Guardrails & Sandboxing", "subtopics": [
            "OWASP Top 10 for LLMs: Prompt Injection & Data Leakage", "API key management & securing environmental configurations", "Isolating code execution using Docker sandboxing", "Implementing output moderation & validation rules"
        ], "value": "Critical for cybersecurity and robustness. Ensures interns build production-grade, secure applications that protect api secrets and contain AI execution risks."}
    ]
    
    add_curriculum_table(doc, curriculum_items)
    
    # 4. 14 Unique Projects Section
    add_section_header(doc, "Collaborative Projects: Specifications for 14 Teams")
    
    p = doc.add_paragraph()
    r = p.add_run(
        "Each of the 14 teams is assigned a unique project on Day 1. The specifications below are structured "
        "to ensure a balanced workload where each member (UI, Backend, Database, AI/Security) contributes directly "
        "to the project's overall success:"
    )
    r.font.size = Pt(10.5)
    
    projects = [
        ("Team 1: Agentic Code-Execution Sandbox", 
         "A web terminal where users write natural language instructions to generate and execute Python scripts. "
         "The AI Agent (OpenClaw) translates instructions, writes code, executes it inside a secure, isolated Docker container on LaaS, "
         "and reports stdout/stderr back. UI: React Web Console. DB: PostgreSQL for script logs and execution audit records."),
         
        ("Team 2: AI-Driven Customer Ticket Router & Solver", 
         "A customer support portal that ingests incoming user queries. An agent automatically classifies query intent, "
         "queries a PostgreSQL database of historical solutions, and either resolves the ticket using a RAG agent or routes "
         "it to a human specialist. UI: Support dashboard. DB: PostgreSQL. AI: OpenClaw router agent."),
         
        ("Team 3: Smart IoT Energy Fleet Monitor & Alert System", 
         "A dashboard visualizing real-time telemetry from simulated IoT smart meters. Relies on an MQTT broker. "
         "An agent monitors incoming power spikes, logs anomalies in MongoDB, and triggers alert notifications "
         "via n8n email/Slack nodes. UI: Live chart monitor. DB: MongoDB. IoT: Virtual telemetry simulator."),
         
        ("Team 4: LLM Security Shield Middleware Proxy", 
         "A security proxy server sitting in front of public LLMs. It scans incoming prompts for known injection signatures, "
         "verifies API tokens, logs threats in MongoDB, and blocks malicious requests. AI: Guardrail classification model. "
         "DB: MongoDB. Security: Threat logger and rate-limiting middleware."),
         
        ("Team 5: RAG-Powered Academic Research Portal", 
         "A portal where students upload PDF textbooks and research papers. Papers are parsed, vectorized using pgvector, "
         "and stored. A local Hermes LLM runs RAG search loops to help students write literature summaries. "
         "UI: PDF uploader + chat UI. DB: PostgreSQL + pgvector. AI: Hermes RAG engine."),
         
        ("Team 6: Collaborative Block-Editor with Agent Summary", 
         "A block-based document editor (similar to Notion). Multiple users can collaborate. An agent runs "
         "periodically to read the block hierarchy, index document blocks into MongoDB, and generate automated summaries "
         "and tag recommendations. UI: React block editor. DB: MongoDB. AI: n8n text analyzer."),
         
        ("Team 7: Semantic Git Repository Explorer", 
         "A tool that scans local code repositories, extracts function docstrings and structures, and stores them "
         "in Chroma DB. A chat agent lets developers search code semantically (e.g. 'find the function handling user login'). "
         "UI: Code directory tree + Chat. DB: Chroma DB. AI: OpenClaw semantic code indexer."),
         
        ("Team 8: Automated Cybersecurity Vulnerability Logger", 
         "A security dashboard that ingests vulnerability scan reports (NMAP, OWASP). Reports are parsed, stored in "
         "Postgres, and an agent running local Hermes LLM recommends remediation scripts based on a RAG database of CVEs. "
         "UI: Vulnerability list. DB: PostgreSQL. AI: Hermes remediation recommender."),
         
        ("Team 9: Decentralized Vehicle Fleet Routing Dashboard", 
         "An administrative dashboard representing a vehicle fleet. Virtual vehicles send location coordinates. "
         "An agent monitors congestion data, runs route optimization checks, and saves coordinate logs in MongoDB. "
         "UI: Geospatial map dashboard. DB: MongoDB. IoT: Mock GPS coordinate publisher."),
         
        ("Team 10: Autonomous n8n Lead Generation Engine", 
         "A tool that scrapes public tech forums, vectorizes postings, and filters them based on interest templates. "
         "An agent writes a customized outreach draft and logs leads. UI: Campaign controller. "
         "DB: MongoDB. AI/Automation: n8n web-scraping workflow agent."),
         
        ("Team 11: Multi-Agent Patient Diagnostics Assistant", 
         "A medical helper that ingests patient symptoms. An agent maps symptoms to a PostgreSQL medical knowledge base, "
         "routes the file to an agent specialized in diagnosis, and generates a medical report brief. "
         "UI: Patient intake portal. DB: PostgreSQL. AI: Multi-agent classification loop."),
         
        ("Team 12: API Traffic Monitor & Agentic DDoS Guard", 
         "A middleware dashboard logging server requests. An anomaly agent reviews metrics in MongoDB, flags suspicious "
         "IP addresses using vector clustering, and updates rate-limiter configurations dynamically. "
         "UI: Traffic charts. DB: MongoDB. Security: Dynamic IP blocker and firewall configuration."),
         
        ("Team 13: Local AI Technical Documentation Search", 
         "A documentation portal containing system deployment guides. Vectorizes markdown pages into Chroma DB. "
         "A self-hosted Hermes LLM answers technical setup queries entirely offline on LaaS GPU nodes. "
         "UI: Search index + chat wrapper. DB: Chroma DB. AI: Local Hermes RAG."),
         
        ("Team 14: Personal AI Portfolio Scaffolder", 
         "A portfolio designer where developers write their bio. An agent automatically writes clean HTML, CSS, and JS, "
         "verifies syntax, saves config settings in MongoDB, and launches a mock preview. "
         "UI: Portfolio editor pane. DB: MongoDB. AI: Code generation agent.")
    ]
    
    for title, desc in projects:
        p_proj = doc.add_paragraph()
        run_title = p_proj.add_run(f"•  {title}\n")
        run_title.font.bold = True
        run_title.font.size = Pt(11)
        run_title.font.color.rgb = NAVY
        
        run_desc = p_proj.add_run(desc)
        run_desc.font.size = Pt(9.5)
        run_desc.font.color.rgb = BODY_TEXT
        p_proj.paragraph_format.left_indent = Inches(0.2)
        p_proj.paragraph_format.space_after = Pt(8)
        
    doc.add_paragraph()
    
    # Save document
    output_path = r"c:\Users\Punith\LaaS\CS_15_Day_Internship_Curriculum.docx"
    try:
        doc.save(output_path)
        print(f"SUCCESS: Curriculum document generated at: {output_path}")
    except Exception as e:
        print(f"ERROR: Error saving document: {e}")

if __name__ == "__main__":
    build_document()
